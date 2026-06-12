#!/usr/bin/env python3
"""
Generate AMIRA_PITCH_DECK.docx from AMIRA_PITCH_DECK.md.

Renders the markdown deck draft as a Word document with:
- One slide per page (page breaks between slides)
- Slide number indicator at the top of each slide
- Large slide titles (deck-style)
- Visual placeholders rendered as shaded callout boxes
- Tables, bullets, blockquotes preserved
- Cover slide with centered branding
- Final "Notes for finalization" appendix on its own page
"""

import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


ROOT = Path(__file__).resolve().parent
MD_PATH = ROOT / "AMIRA_PITCH_DECK.md"
_docx_name = sys.argv[1] if len(sys.argv) > 1 else "AMIRA_PITCH_DECK.docx"
DOCX_PATH = ROOT / _docx_name


# ---------- Color palette ----------

COLOR_TITLE = RGBColor(0x1F, 0x38, 0x64)      # Dark blue (slide titles)
COLOR_ACCENT = RGBColor(0x2E, 0x75, 0xB6)     # Medium blue (accents)
COLOR_MUTED = RGBColor(0x59, 0x59, 0x59)      # Dark gray (slide-N indicator)
COLOR_PLACEHOLDER_TEXT = RGBColor(0x9C, 0x6F, 0x00)  # Amber-brown
COLOR_HEADER_BG = "1F3864"                    # Table header shading
COLOR_PLACEHOLDER_BG = "FFF4CC"               # Light yellow (visual placeholders)
COLOR_TAGLINE = RGBColor(0x44, 0x55, 0x66)    # Slate gray


# ---------- Helpers ----------

def set_cell_shading(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def shade_paragraph(para, color_hex):
    """Add background shading to a paragraph."""
    p_pr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    p_pr.append(shd)


def add_paragraph_border(para, color_hex="2E75B6", size=6):
    """Add a border around a paragraph (used for visual placeholder boxes)."""
    p_pr = para._p.get_or_add_pPr()
    p_borders = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), str(size))
        border.set(qn("w:space"), "4")
        border.set(qn("w:color"), color_hex)
        p_borders.append(border)
    p_pr.append(p_borders)


def add_page_number(doc):
    """Add page numbering in footer."""
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.text = "PAGE"
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr)
    run._r.append(fld_char_2)
    run.font.size = Pt(9)
    run.font.color.rgb = COLOR_MUTED


def add_formatted_runs(paragraph, text):
    """Parse inline markdown (**bold**, *italic*, `code`) into runs."""
    i = 0
    n = len(text)
    while i < n:
        # Bold **...**
        if text[i:i + 2] == "**":
            end = text.find("**", i + 2)
            if end > i + 2:
                run = paragraph.add_run(text[i + 2:end])
                run.bold = True
                i = end + 2
                continue
        # Inline code `...`
        if text[i] == "`":
            end = text.find("`", i + 1)
            if end > i:
                run = paragraph.add_run(text[i + 1:end])
                run.font.name = "Consolas"
                run.font.size = Pt(10)
                i = end + 1
                continue
        # Italic *...*
        if (text[i] == "*"
                and (i + 1 < n)
                and text[i + 1] != " "
                and text[i + 1] != "*"):
            end = text.find("*", i + 1)
            if end > i + 1 and (end + 1 >= n or text[end + 1] != "*"):
                run = paragraph.add_run(text[i + 1:end])
                run.italic = True
                i = end + 1
                continue
        # Plain run — accumulate until next marker
        j = i
        while j < n:
            if text[j] == "`":
                break
            if text[j:j + 2] == "**":
                break
            if text[j] == "*" and (j + 1 < n) and text[j + 1] != " ":
                break
            j += 1
        paragraph.add_run(text[i:j])
        i = j


def add_table_from_rows(doc, header, data_rows):
    if not header:
        return
    cols = len(header)
    table = doc.add_table(rows=1, cols=cols)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    hdr_cells = table.rows[0].cells
    for idx, h in enumerate(header):
        hdr_cells[idx].text = ""
        p = hdr_cells[idx].paragraphs[0]
        run = p.add_run(h.strip())
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(hdr_cells[idx], COLOR_HEADER_BG)

    for row in data_rows:
        cells = table.add_row().cells
        for idx in range(cols):
            cell_text = row[idx].strip() if idx < len(row) else ""
            cells[idx].text = ""
            p = cells[idx].paragraphs[0]
            add_formatted_runs(p, cell_text)
    doc.add_paragraph()


def add_visual_placeholder(doc, text):
    """Render a `[VISUAL: ...]` callout as a shaded, bordered, italic block."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = COLOR_PLACEHOLDER_TEXT
    run.font.size = Pt(11)
    shade_paragraph(p, COLOR_PLACEHOLDER_BG)
    add_paragraph_border(p, color_hex="E5B85C", size=6)


def add_blockquote(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.right_indent = Cm(0.8)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = COLOR_MUTED


# ---------- Body renderer ----------

def render_body(doc, body):
    """Render a slide body (markdown text) into the doc."""
    lines = body.split("\n")
    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Visual placeholder: backtick-wrapped [VISUAL: ...]
        if stripped.startswith("`[VISUAL:") and stripped.endswith("]`"):
            text = stripped.strip("`")
            add_visual_placeholder(doc, text)
            i += 1
            continue

        # Tables
        if stripped.startswith("|") and "|" in stripped[1:]:
            tbl_lines = []
            while i < n and lines[i].strip().startswith("|"):
                tbl_lines.append(lines[i])
                i += 1
            if len(tbl_lines) >= 2 and re.search(r"\|[\s:\-|]+\|", tbl_lines[1]):
                header = [c.strip() for c in tbl_lines[0].strip().strip("|").split("|")]
                data_rows = []
                for row_line in tbl_lines[2:]:
                    cells = [c.strip() for c in row_line.strip().strip("|").split("|")]
                    data_rows.append(cells)
                add_table_from_rows(doc, header, data_rows)
            else:
                for tl in tbl_lines:
                    p = doc.add_paragraph()
                    add_formatted_runs(p, tl)
            continue

        # Blockquote
        if stripped.startswith(">"):
            quote_lines = []
            while i < n and lines[i].strip().startswith(">"):
                ql = lines[i].strip().lstrip(">").strip()
                if ql:
                    quote_lines.append(ql)
                i += 1
            if quote_lines:
                add_blockquote(doc, " ".join(quote_lines))
            continue

        # Bullets
        if stripped.startswith("- "):
            while i < n and lines[i].strip().startswith("- "):
                stripped_line = lines[i].strip()
                # determine indent level (root vs nested by leading spaces in original line)
                leading = len(lines[i]) - len(lines[i].lstrip(" "))
                style = "List Bullet 2" if leading >= 2 else "List Bullet"
                p = doc.add_paragraph(style=style)
                add_formatted_runs(p, stripped_line[2:])
                i += 1
                # continuation indented two spaces (no bullet)
                while i < n and lines[i].startswith("  ") and lines[i].strip() and not lines[i].lstrip().startswith("- "):
                    p.add_run("\n" + lines[i].strip())
                    i += 1
            continue

        # Numbered list
        if re.match(r"^\d+\.\s", stripped):
            while i < n and re.match(r"^\d+\.\s", lines[i].strip()):
                text = re.sub(r"^\d+\.\s", "", lines[i].strip())
                p = doc.add_paragraph(style="List Number")
                add_formatted_runs(p, text)
                i += 1
            continue

        # Bold standalone callout: line that is fully bold
        if stripped.startswith("**") and stripped.endswith("**") and stripped.count("**") == 2:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            run = p.add_run(stripped[2:-2])
            run.bold = True
            run.font.size = Pt(13)
            run.font.color.rgb = COLOR_ACCENT
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        add_formatted_runs(p, stripped)
        i += 1


# ---------- Slide-level rendering ----------

SLIDE_HEADER_RE = re.compile(r"^##\s+Slide\s+(\d+)\s+—\s+(.+?)\s*$", flags=re.MULTILINE)


def parse_slides(md_text):
    """Return list of (number, title, body) tuples."""
    matches = list(SLIDE_HEADER_RE.finditer(md_text))
    slides = []
    for idx, m in enumerate(matches):
        num = int(m.group(1))
        title = m.group(2).strip()
        body_start = m.end()
        body_end = matches[idx + 1].start() if idx + 1 < len(matches) else len(md_text)
        body = md_text[body_start:body_end]
        # Drop trailing horizontal-rule separators
        body = re.sub(r"\n---\s*\n", "\n\n", body).strip()
        slides.append((num, title, body))
    return slides


def render_cover_slide(doc, body, total_slides):
    """Slide 1 — cover with centered branding. Body is title + tagline content."""
    # Vertical spacing
    for _ in range(4):
        doc.add_paragraph()

    # Brand title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Amira")
    run.font.size = Pt(56)
    run.bold = True
    run.font.color.rgb = COLOR_TITLE

    # Body content from cover slide — render the bold lines and tagline
    # We render line-by-line, centering and styling specially
    lines = [ln.strip() for ln in body.split("\n") if ln.strip()]
    for line in lines:
        # Skip horizontal rules
        if line == "---":
            continue
        # Italic (already markdown italic)
        if line.startswith("*") and line.endswith("*") and not line.startswith("**"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line.strip("*"))
            run.italic = True
            run.font.size = Pt(18)
            run.font.color.rgb = COLOR_ACCENT
            continue
        # Visual placeholder
        if line.startswith("`[VISUAL:") and line.endswith("]`"):
            add_visual_placeholder(doc, line.strip("`"))
            continue
        # Blockquote
        if line.startswith(">"):
            text = line.lstrip(">").strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            run.italic = True
            run.font.size = Pt(11)
            run.font.color.rgb = COLOR_MUTED
            continue
        # Plain centered text (tagline)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_formatted_runs(p, line)
        for run in p.runs:
            run.font.size = Pt(13)
            run.font.color.rgb = COLOR_TAGLINE


def render_slide(doc, num, title, body, total):
    """Render a content slide on its own page."""
    doc.add_page_break()

    # "Slide N of M" indicator
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f"Slide {num} of {total}")
    run.font.size = Pt(9)
    run.font.color.rgb = COLOR_MUTED
    run.italic = True

    # Slide title
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = COLOR_TITLE

    # Body
    render_body(doc, body)


# ---------- Main ----------

def main():
    md = MD_PATH.read_text(encoding="utf-8")
    slides = parse_slides(md)

    if not slides:
        print("ERROR: no slides parsed from markdown.")
        sys.exit(1)

    doc = Document()

    # Page setup
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    total = len(slides)

    # Slide 1 = cover (special)
    cover_num, cover_title, cover_body = slides[0]
    render_cover_slide(doc, cover_body, total)

    # Slides 2..N = content
    for num, title, body in slides[1:]:
        render_slide(doc, num, title, body, total)

    # Final "Notes for finalization" appendix (if present in markdown)
    notes_match = re.search(
        r"^# Notes for finalization\s*\n(.*)\Z",
        md,
        flags=re.MULTILINE | re.DOTALL,
    )
    if notes_match:
        doc.add_page_break()
        p = doc.add_paragraph()
        run = p.add_run("Notes for finalization")
        run.bold = True
        run.font.size = Pt(20)
        run.font.color.rgb = COLOR_TITLE

        p = doc.add_paragraph()
        run = p.add_run("Internal working notes — to be removed before sending to the client.")
        run.italic = True
        run.font.size = Pt(10)
        run.font.color.rgb = COLOR_MUTED

        notes_body = notes_match.group(1).strip()
        # Strip subsection bold-headings out a bit so they render clearly
        render_body(doc, notes_body)

    add_page_number(doc)
    doc.save(DOCX_PATH)
    print(f"Wrote {DOCX_PATH}")
    print(f"Size: {DOCX_PATH.stat().st_size:,} bytes")
    print(f"Slides: {total}")


if __name__ == "__main__":
    main()
