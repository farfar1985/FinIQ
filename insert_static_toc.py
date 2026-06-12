"""
Insert a static Table of Contents into the INLINE and APPENDIX docx files.

Why static (not a Word TOC field): the user's Word is unlicensed, and saving
appears to mangle TOC field cached content. A static TOC (plain paragraphs
with tab-leader dots and explicit page numbers) is bulletproof.

INLINE page numbers come from the user's screenshot (Word actually rendered
those numbers, so they're correct).

APPENDIX page numbers are estimated from heading position + image distribution
in the Workflow Walkthrough — typically accurate within ±1 page.

The script also inserts a page break paragraph after the TOC so "1. Introduction"
starts on its own page.
"""

import os
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER

OUT_DIR = r"C:\Users\farza\Desktop\mars docs"
INLINE_PATH = os.path.join(OUT_DIR, "Amira_Proposal_for_Mars_2026-04-26_Polished_INLINE.docx")
APPENDIX_PATH = os.path.join(OUT_DIR, "Amira_Proposal_for_Mars_2026-04-26_Polished_APPENDIX.docx")


# (label, page, level)  — level 1 = main section, level 2 = subsection
INLINE_TOC = [
    ("1. Introduction",                          2,  1),
    ("1.1 The Challenge",                        2,  2),
    ("1.2 Amira at a Glance",                    2,  2),
    ("2. Platform Architecture",                 3,  1),
    ("2.1 The Three-Step Pipeline",              3,  2),
    ("2.2 Three Agents, One Workflow",           3,  2),
    ("3. Pipeline Phases",                       4,  1),
    ("3.1 Specifications Phase",                 4,  2),
    ("3.2 Development Phase (Canvas)",           8,  2),
    ("3.3 Artifacts Phase",                      10, 2),
    ("3.4 Reversibility and Versioning",         11, 2),
    ("4. Differentiators",                       12, 1),
    ("4.1 Proprietary APIs",                     12, 2),
    ("4.2 How Skills Connect to Specifications", 15, 2),
    ("4.3 Apps Become Agents",                   16, 2),
    ("5. Governance and Security",               19, 1),
    ("5.1 Human Governance and Audit",           19, 2),
    ("5.2 Knowledge Base and Secret Vault",      21, 2),
    ("6. Platform Features Summary",             22, 1),
    ("7. Validation and Replication",            23, 1),
    ("7.1 Proof Point: FinIQ",                   23, 2),
    ("7.2 Replication Roadmap",                  24, 2),
    ("8. Technical Implementation",              24, 1),
    ("8.1 Architecture Overview",                24, 2),
    ("8.2 Deployment Options",                   25, 2),
    ("8.3 Authentication and API Key Strategy",  28, 2),
    ("9. Amira Deployment Scope",                29, 1),
    ("10. Commercial Model",                     29, 1),
    ("10.1 Platform License",                    29, 2),
    ("10.2 Skill Development",                   30, 2),
    ("10.3 Application APIs",                    30, 2),
    ("10.4 Compounding Model",                   30, 2),
    ("11. Closing",                              30, 1),
]

APPENDIX_TOC = [
    ("1. Introduction",                          2,  1),
    ("1.1 The Challenge",                        2,  2),
    ("1.2 Amira at a Glance",                    2,  2),
    ("2. Platform Architecture",                 3,  1),
    ("2.1 The Three-Step Pipeline",              3,  2),
    ("2.2 Three Agents, One Workflow",           3,  2),
    ("3. Pipeline Phases",                       4,  1),
    ("3.1 Specifications Phase",                 4,  2),
    ("3.2 Development Phase (Canvas)",           4,  2),
    ("3.3 Artifacts Phase",                      5,  2),
    ("3.4 Reversibility and Versioning",         5,  2),
    ("4. Differentiators",                       6,  1),
    ("4.1 Proprietary APIs",                     6,  2),
    ("4.2 How Skills Connect to Specifications", 6,  2),
    ("4.3 Apps Become Agents",                   7,  2),
    ("5. Governance and Security",               7,  1),
    ("5.1 Human Governance and Audit",           7,  2),
    ("5.2 Knowledge Base and Secret Vault",      8,  2),
    ("6. Platform Features Summary",             9,  1),
    ("7. Validation and Replication",            11, 1),
    ("7.1 Proof Point: FinIQ",                   11, 2),
    ("7.2 Replication Roadmap",                  11, 2),
    ("8. Technical Implementation",              12, 1),
    ("8.1 Architecture Overview",                12, 2),
    ("8.2 Deployment Options",                   12, 2),
    ("8.3 Authentication and API Key Strategy",  13, 2),
    ("9. Amira Deployment Scope",                13, 1),
    ("10. Commercial Model",                     14, 1),
    ("10.1 Platform License",                    14, 2),
    ("10.2 Skill Development",                   14, 2),
    ("10.3 Application APIs",                    15, 2),
    ("10.4 Compounding Model",                   15, 2),
    ("11. Closing",                              15, 1),
    ("Appendix A — Workflow Walkthrough",        16, 1),
    ("Step 1 — Home portal",                     16, 2),
    ("Step 2 — Three ways to start a project",   17, 2),
    ("Step 3 — Decision Point",                  18, 2),
    ("Step 4 — Living spec, gaps, version history", 19, 2),
    ("Step 5 — Route for e-signature",           20, 2),
    ("Step 6 — Authorized approver view",        20, 2),
    ("Step 7 — Open Canvas",                     21, 2),
    ("Step 8 — Build Agent chat",                22, 2),
    ("Step 9 — Combine skills, knowledge, agents", 23, 2),
    ("Step 10 — Compliance matrix",              25, 2),
    ("Step 11 — Deploy modal",                   26, 2),
    ("Step 12 — Skills marketplace",             28, 2),
    ("Step 13 — Ask Amira",                      30, 2),
    ("Step 14 — Lineage",                        31, 2),
    ("Step 15 — Wrap",                           31, 2),
]


def find_intro_heading(doc):
    """Return the paragraph element of '1. Introduction' Heading 1."""
    for p in doc.paragraphs:
        if p.style.name == "Heading 1" and p.text.strip().startswith("1. Introduction"):
            return p
    raise LookupError("'1. Introduction' Heading 1 not found")


def remove_existing_page_break_in(paragraph):
    """If the paragraph contains a w:br[type=page], remove it (we'll re-add cleanly)."""
    runs = paragraph._element.findall(qn("w:r"))
    for r in runs:
        for br in r.findall(qn("w:br")):
            if br.get(qn("w:type")) == "page":
                r.remove(br)


def insert_static_toc(doc, toc_entries):
    """Insert TOC paragraphs immediately before '1. Introduction', followed by a page break."""
    intro_p = find_intro_heading(doc)
    intro_el = intro_p._element

    # Drop any leftover page-break run already inside "1. Introduction" (the user's earlier Ctrl+Enter
    # may have left a w:br in the heading run; we re-add the page break as its own paragraph below
    # so it's structurally clean).
    remove_existing_page_break_in(intro_p)

    # 1. TOC header paragraph (Normal, bold, 14pt, centered)
    header = doc.add_paragraph()
    header.alignment = 1  # WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run("Table of Contents")
    run.bold = True
    run.font.size = Pt(14)
    intro_el.addprevious(header._element)

    # 2. Spacer
    spacer = doc.add_paragraph()
    intro_el.addprevious(spacer._element)

    # 3. Each TOC entry: indent by level, label + tab + page number, with dot-leader tab stop
    for label, page, level in toc_entries:
        entry = doc.add_paragraph()

        # Indent by level
        if level == 2:
            entry.paragraph_format.left_indent = Inches(0.3)

        # Tab stop at right margin (6.3") with dot leader
        tab_stops = entry.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(6.3), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)

        # Label, then a tab character, then the page number
        entry.add_run(label)
        entry.add_run("\t")
        entry.add_run(str(page))

        intro_el.addprevious(entry._element)

    # 4. Spacer + explicit page break paragraph so Introduction starts on its own page
    intro_el.addprevious(doc.add_paragraph()._element)
    pb_xml = (
        '<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:r><w:br w:type="page"/></w:r>'
        '</w:p>'
    )
    intro_el.addprevious(parse_xml(pb_xml))


def main():
    for path, toc_entries, label in [
        (INLINE_PATH,   INLINE_TOC,   "INLINE"),
        (APPENDIX_PATH, APPENDIX_TOC, "APPENDIX"),
    ]:
        if not os.path.exists(path):
            print(f"  ** missing: {path}")
            continue
        print(f"Processing {label}: {path}")
        doc = Document(path)
        insert_static_toc(doc, toc_entries)
        doc.save(path)
        print(f"  saved ({len(toc_entries)} TOC entries inserted).")
    print("Done.")


if __name__ == "__main__":
    main()
