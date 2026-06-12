"""
Generate simplified FinIQ Project Plan — checklist style.
Clean status + what's done / in progress / upcoming.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# -- Page setup --
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# -- Styles --
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.15

NAVY = RGBColor(0x1B, 0x3A, 0x5C)
GRAY = RGBColor(0x77, 0x77, 0x77)
GREEN = RGBColor(0x1A, 0x7A, 0x2E)
ORANGE = RGBColor(0xCC, 0x7A, 0x00)
RED = RGBColor(0xCC, 0x22, 0x22)

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = NAVY
    return h

def add_check(text, status='done'):
    """Add a checklist item. status: done, progress, upcoming"""
    p = doc.add_paragraph()
    if status == 'done':
        icon = '\u2705  '
        color = GREEN
    elif status == 'progress':
        icon = '\U0001f7e1  '
        color = ORANGE
    else:
        icon = '\u2B1C  '
        color = GRAY

    run = p.add_run(icon)
    run.font.size = Pt(11)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = color if status != 'done' else RGBColor(0x33, 0x33, 0x33)
    return p

def add_table_simple(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shading = cell._element.get_or_add_tcPr()
        shading_elm = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): '1B3A5C', qn('w:val'): 'clear'
        })
        shading.append(shading_elm)

    for r, row_data in enumerate(rows):
        for c, cell_text in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(cell_text)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9.5)

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(width)
    return table

# ============================================================
# TITLE
# ============================================================
doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('FinIQ — Project Status & Plan')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = NAVY

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Unified Financial Analytics Hub')
run.font.size = Pt(14)
run.font.color.rgb = GRAY

doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run('April 1, 2026  |  Target: April 21 MLT Demo  |  Amira Technologies / QDT')
run.font.size = Pt(10)
run.font.color.rgb = GRAY

doc.add_paragraph()
doc.add_paragraph()

# ============================================================
# WHAT'S DONE
# ============================================================
add_heading_styled('Completed', level=1)

add_check('SRS v3.1 finalized (52 functional requirements)')
add_check('Frontend Design Guideline v1.0 (Bloomberg-inspired dark theme)')
add_check('Real Databricks schema documented (21 objects, 5.7B+ row tables)')
add_check('Semantic data layer created (7 YAML files, all tables/columns/metrics)')
add_check('Three parallel builds completed and tested independently')
add_check('Three-way merge into unified application (8 commits pushed to GitHub)')
add_check('Dashboard with 6 KPI cards and interactive Recharts visualizations')
add_check('Data Explorer with SQL query builder')
add_check('PES report generation (AI narratives, What\'s Working / Not Working)')
add_check('Natural language query engine (Anthropic Claude LLM \u2192 Databricks SQL)')
add_check('Competitive Intelligence module (10 tabs: SWOT, Porter\'s, ESG, Alerts, etc.)')
add_check('Real FMP API integration (10 competitors, live stock/financial data)')
add_check('Enterprise Job Board (submission, SLA routing, agent processing, WebSocket)')
add_check('XLSX export for job results')
add_check('Rate limiting and SQL injection protection')
add_check('Voice agent WebSocket proxy (OpenAI Realtime API)')
add_check('Suggested prompt library (18 curated prompts with variable resolution)')
add_check('Live testing with real Databricks production data \u2014 P&L, MAC, revenue queries confirmed')
add_check('Compliance: 85% (68/80 items PASS, 0 FAILs)')

doc.add_paragraph()

# ============================================================
# IN PROGRESS
# ============================================================
add_heading_styled('In Progress', level=1)

add_check('Unit alias resolution (fuzzy matching for 766 org units)', 'progress')
add_check('Production build test (next build \u2014 only dev mode verified so far)', 'progress')
add_check('PR merge: merged branch \u2192 main (team pulls main to get unified app)')
add_check('Multi-turn query context (LLM remembering entity/period across turns)', 'progress')
add_check('Chart follow-up detection ("show me a chart" re-renders existing data)', 'progress')
add_check('Voice agent browser microphone (needs HTTPS or AudioWorklet fix)', 'progress')
add_check('Budget variance query tuning (unit + date handling in replan tables)', 'progress')
add_check('A2A / Gemini documentation review (received from Mars, needs assessment)', 'progress')
add_check('Azure VM provisioning (Kumar/Marc provisioning resources)', 'progress')

doc.add_paragraph()

# ============================================================
# UPCOMING
# ============================================================
add_heading_styled('Upcoming (April 2\u201321)', level=1)

add_check('Deploy application to Azure VM', 'upcoming')
add_check('HTTPS certificate for voice agent', 'upcoming')
add_check('Semantic layer query routing (YAML-driven, Cesar\'s data agent)', 'upcoming')
add_check('MCP / A2A protocol decision for cross-platform integration', 'upcoming')
add_check('Dashboard and UI polish pass', 'upcoming')
add_check('CI module QA with live FMP data (all 10 tabs)', 'upcoming')
add_check('Compliance push from 85% \u2192 90%+', 'upcoming')
add_check('End-to-end integration testing (all pages, real data)', 'upcoming')
add_check('Demo script creation and narrative', 'upcoming')
add_check('Demo rehearsal #1 (April 15)', 'upcoming')
add_check('Demo rehearsal #2 (April 17)', 'upcoming')
add_check('Performance optimization (query response < 5 seconds)', 'upcoming')
add_check('Demo data curation (best-case queries pre-selected)', 'upcoming')
add_check('Fallback plan if Databricks is slow (cached snapshots)', 'upcoming')
add_check('MLT presentation and live demo (April 21)', 'upcoming')

doc.add_paragraph()

# ============================================================
# WEEKLY PLAN
# ============================================================
doc.add_page_break()
add_heading_styled('Week-by-Week Plan', level=1)

add_heading_styled('Week 1: Stabilization (April 1\u20134)', level=2)
doc.add_paragraph(
    'Fix core query issues (unit aliases, budget variance), get production build passing, '
    'verify CI module with live data. Goal: all standard queries working reliably.'
)

add_heading_styled('Week 2: Enhancement & Deployment (April 7\u201311)', level=2)
doc.add_paragraph(
    'Deploy to Azure VM, complete voice agent, integrate semantic layer routing, '
    'push compliance to 90%+. Goal: app running on cloud infrastructure, accessible to team.'
)

add_heading_styled('Week 3: Demo Prep & Polish (April 14\u201321)', level=2)
doc.add_paragraph(
    'Write demo script, two rehearsals, curate demo data, optimize performance, '
    'prepare fallback scenarios. Goal: polished, rehearsed demo for Mars MLT.'
)

doc.add_paragraph()

# ============================================================
# KEY MILESTONES
# ============================================================
add_heading_styled('Key Milestones', level=1)

add_table_simple(
    ['Date', 'Milestone', 'Status'],
    [
        ['Apr 1', 'Project plan delivered to team', '\u2705 Done'],
        ['Apr 4', 'All core NL queries working on real Databricks', '\U0001f7e1 In progress'],
        ['Apr 7', 'Semantic layer query routing live', '\u2B1C Upcoming'],
        ['Apr 11', 'App deployed on Azure VM', '\u2B1C Upcoming'],
        ['Apr 11', '90%+ compliance (72+/80 items)', '\u2B1C Upcoming'],
        ['Apr 15', 'Demo rehearsal #1', '\u2B1C Upcoming'],
        ['Apr 17', 'Demo rehearsal #2', '\u2B1C Upcoming'],
        ['Apr 21', 'MLT Presentation & Live Demo', '\u2B1C Upcoming'],
    ],
    col_widths=[0.8, 3.0, 1.2]
)

doc.add_paragraph()

# ============================================================
# PLATFORM VISION (brief)
# ============================================================
add_heading_styled('Platform Vision', level=1)

doc.add_paragraph(
    'FinIQ is the first proof-of-concept for a reusable AI agent platform. '
    'The architecture is built around four modular components\u2014orchestration, '
    'data fetching (text-to-SQL), data science, and visualization\u2014that can be '
    'applied across business units without rebuilding per use case.'
)

p = doc.add_paragraph()
run = p.add_run('Why this matters: ')
run.bold = True
p.add_run(
    'Previous internal builds took 12 weeks. FinIQ was built in under 2 weeks. '
    'Reusable components mean each new use case (health data, supply chain, '
    'marketing) requires only data ingestion and mapping\u2014not a full rebuild.'
)

doc.add_paragraph()

add_heading_styled('After April 21', level=2)

add_check('Embeddings & vectorization layer for semantic search', 'upcoming')
add_check('Amira Financial Forecasting API integration', 'upcoming')
add_check('Amira Marketing Analytics API integration', 'upcoming')
add_check('RBAC and role-based data access', 'upcoming')
add_check('Multi-tenant deployment (per-division agents)', 'upcoming')
add_check('A2A protocol finalization (MCP or custom)', 'upcoming')
add_check('Precomputed tables in Databricks for heavy query optimization', 'upcoming')
add_check('Cross-domain expansion (Digital Health, supply chain)', 'upcoming')

# ============================================================
# FOOTER
# ============================================================
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Questions? Reach out in the FinIQ GenAI group chat.')
run.font.color.rgb = GRAY
run.font.size = Pt(10)

# Save
output_path = r'C:\Users\farza\Desktop\FinIQ Project Plan (Simple) - April 2026.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
