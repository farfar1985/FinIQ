"""
Amira FinIQ - Software Requirements Specification (SRS) Generator
IEEE 830 / ISO/IEC/IEEE 29148 Format
Generates a professional Word document for presentation to Mars Inc.
Version 2.0
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

# ── Branding colors ──
NAVY = RGBColor(0x0A, 0x1F, 0x44)       # Dark navy for titles
BLUE = RGBColor(0x1A, 0x56, 0xDB)       # Amira blue for headings
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)  # Body text
LIGHT_BLUE = RGBColor(0xE8, 0xF0, 0xFE) # Table header bg
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
TABLE_BORDER_COLOR = "CCCCCC"
ACCENT = "1A56DB"

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# ── Style definitions ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = DARK_GRAY

for level, (size, color) in enumerate([(22, NAVY), (16, BLUE), (13, BLUE), (11, DARK_GRAY)], 1):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Calibri'
    h.font.size = Pt(size)
    h.font.color.rgb = color
    h.font.bold = True
    if level <= 2:
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(8)

# ── Helper functions ──

def add_title_page():
    """Create a professional cover page."""
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Amira FinIQ")
    run.font.size = Pt(36)
    run.font.color.rgb = NAVY
    run.font.bold = True
    run.font.name = 'Calibri'

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Unified Financial Analytics Hub")
    run.font.size = Pt(18)
    run.font.color.rgb = BLUE
    run.font.name = 'Calibri'

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Software Requirements Specification")
    run.font.size = Pt(16)
    run.font.color.rgb = DARK_GRAY
    run.font.name = 'Calibri'

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("IEEE 830 / ISO/IEC/IEEE 29148 Format")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.font.name = 'Calibri'

    doc.add_paragraph()
    doc.add_paragraph()

    # Metadata table
    meta = [
        ("Document Version", "2.0"),
        ("Date", datetime.date.today().strftime("%B %d, %Y")),
        ("Classification", "Confidential"),
        ("Prepared For", "Mars, Incorporated"),
        ("Prepared By", "Amira Technologies"),
        ("Standard", "IEEE 830-1998 / ISO/IEC/IEEE 29148:2018"),
    ]
    table = doc.add_table(rows=len(meta), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(meta):
        for j, text in enumerate([k, v]):
            cell = table.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            run.font.size = Pt(10)
            run.font.color.rgb = DARK_GRAY
            run.font.name = 'Calibri'
            if j == 0:
                run.font.bold = True

    doc.add_page_break()


def add_para(text, bold=False, italic=False, size=None, color=None, space_after=6):
    """Add a paragraph with formatting."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.font.bold = True
    if italic:
        run.font.italic = True
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_bullet(text, level=0):
    """Add a bullet point."""
    p = doc.add_paragraph(text, style='List Bullet')
    if level > 0:
        p.paragraph_format.left_indent = Inches(0.5 * (level + 1))
    p.paragraph_format.space_after = Pt(2)
    return p


def add_numbered(text):
    """Add a numbered item."""
    p = doc.add_paragraph(text, style='List Number')
    p.paragraph_format.space_after = Pt(2)
    return p


def add_table(headers, rows, col_widths=None):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE
        run.font.name = 'Calibri'
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{ACCENT}"/>')
        cell._element.get_or_add_tcPr().append(shading)

    # Data rows
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.cell(i + 1, j)
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.name = 'Calibri'
            run.font.color.rgb = DARK_GRAY
            if i % 2 == 1:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F6FC"/>')
                cell._element.get_or_add_tcPr().append(shading)

    # Borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="{TABLE_BORDER_COLOR}"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="{TABLE_BORDER_COLOR}"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{TABLE_BORDER_COLOR}"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="{TABLE_BORDER_COLOR}"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{TABLE_BORDER_COLOR}"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="{TABLE_BORDER_COLOR}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

    if col_widths:
        for i_row in range(len(rows) + 1):
            for j, w in enumerate(col_widths):
                table.cell(i_row, j).width = Inches(w)

    doc.add_paragraph()  # spacing
    return table


def add_req_table(req_id, title, description, priority="High", category="Functional"):
    """Add a requirement in structured format."""
    p = doc.add_paragraph()
    run = p.add_run(f"{req_id}: {title}")
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = BLUE
    p.paragraph_format.space_after = Pt(4)
    add_para(description, space_after=2)
    p2 = doc.add_paragraph()
    run2 = p2.add_run(f"Priority: {priority}  |  Category: {category}")
    run2.font.size = Pt(9)
    run2.font.italic = True
    run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    p2.paragraph_format.space_after = Pt(10)


# ══════════════════════════════════════════════════════════════════
#  BUILD THE DOCUMENT (IEEE 830 Format)
# ══════════════════════════════════════════════════════════════════

add_title_page()

# ── TABLE OF CONTENTS placeholder ──
doc.add_heading("Table of Contents", level=1)
add_para("[Table of Contents will auto-generate when opened in Word: References > Update Table]", italic=True, color=RGBColor(0x99, 0x99, 0x99))
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
#  1. INTRODUCTION (IEEE 830 Section 1)
# ══════════════════════════════════════════════════════════════════
doc.add_heading("1. Introduction", level=1)

# 1.1 Purpose
doc.add_heading("1.1 Purpose", level=2)
add_para(
    "This Software Requirements Specification (SRS) defines the functional and non-functional "
    "requirements for Amira FinIQ, a Unified Financial Analytics Hub proposed for Mars, Incorporated. "
    "This document is intended for use by Amira Technologies' engineering and product teams, Mars's "
    "IT and digital leadership, and key business stakeholders who will evaluate, approve, and ultimately "
    "use the platform."
)
add_para(
    "This SRS follows the IEEE 830-1998 and ISO/IEC/IEEE 29148:2018 standards for software "
    "requirements specifications."
)

# 1.2 Scope
doc.add_heading("1.2 Scope", level=2)
add_para(
    "Amira FinIQ is a unified financial analytics hub designed to consolidate and augment Mars, "
    "Incorporated's existing analytical capabilities, transitioning from a dispersed setup to a single, "
    "intelligent platform. The system will unify internal financial reporting and external competitive "
    "intelligence into one hub, enabling on-demand report generation, natural language queries, and "
    "enterprise-scale agent-based job processing."
)
add_para(
    "The platform will integrate with Amira's existing Financial Forecasting and Marketing Analytics "
    "services, enabling forward-looking recommendations alongside historical analysis."
)

doc.add_heading("1.2.1 In Scope", level=3)
add_bullet("Internal financial data ingestion and processing (P&L, Brand, Product, NCFO sheets)")
add_bullet("External competitor document ingestion (earnings releases, presentations, transcripts)")
add_bullet("Acquired research and third-party data integration")
add_bullet("Natural language query interface for ad-hoc and structured reporting")
add_bullet("Period End Summary generation (Summary, WWW, WNWW formats) with full configurability")
add_bullet("Peer-to-peer competitive benchmarking with quantitative comparison tables")
add_bullet("Enterprise job board with agent queue, priority routing, and SLA management")
add_bullet("Integration APIs for Amira Financial Forecasting and Marketing Analytics platforms")
add_bullet("Dynamic, configurable user interface with real-time updates")
add_bullet("Role-based access control (RBAC) aligned with Mars organizational hierarchy")
add_bullet("Audit trail and data lineage for all generated reports")
add_bullet("Caching and performance optimization for sub-second cached retrieval")

doc.add_heading("1.2.2 Out of Scope", level=3)
add_bullet("Direct SAP/ERP integration (data provided via preprocessed Excel exports)")
add_bullet("Real-time streaming data (batch processing model)")
add_bullet("Mobile native applications (responsive web application)")
add_bullet("Multi-language report generation (English only initially)")

# 1.3 Definitions, Acronyms, Abbreviations
doc.add_heading("1.3 Definitions, Acronyms, and Abbreviations", level=2)
add_table(
    ["Term", "Definition"],
    [
        ["GBU", "Global Business Unit (e.g., Mars Snacking, Petcare, Food Nutrition & Multisales)"],
        ["PES", "Period End Summary \u2014 the current Mars tool for financial narrative generation"],
        ["CI", "Competitive Intelligence \u2014 the current Mars tool for competitor analysis"],
        ["KPI", "Key Performance Indicator (e.g., Organic Growth, MAC Shape %)"],
        ["OG", "Organic Growth \u2014 top-line revenue growth excluding FX and M&A"],
        ["MAC Shape", "Margin After Conversion as a percentage of Net Sales"],
        ["A&CP Shape", "Advertising & Consumer Promotions spend as a percentage of sales"],
        ["CE Shape", "Controllable Earnings as a percentage of Net Sales"],
        ["NCFO", "Net Cash From Operations"],
        ["P2P", "Peer-to-Peer benchmarking \u2014 comparing performance across competitor peer group"],
        ["WWW / WNWW", "What's Working Well / What's Not Working Well \u2014 output formats for PES"],
        ["RAG", "Retrieval Augmented Generation \u2014 LLM technique combining search with generation"],
        ["SSE", "Server-Sent Events \u2014 protocol for real-time server-to-client streaming"],
        ["SLA", "Service Level Agreement \u2014 target response time for job processing"],
        ["RBAC", "Role-Based Access Control"],
        ["ROAS", "Return on Ad Spend \u2014 marketing efficiency metric"],
        ["LLM", "Large Language Model"],
    ],
    col_widths=[1.5, 5.0]
)

# 1.4 References
doc.add_heading("1.4 References", level=2)
add_bullet("IEEE 830-1998: IEEE Recommended Practice for Software Requirements Specifications")
add_bullet("ISO/IEC/IEEE 29148:2018: Systems and software engineering \u2014 Life cycle processes \u2014 Requirements engineering")
add_bullet("Mars Period End Summary Technical Documentation (internal, 11 pages)")
add_bullet("Mars Competitive Intelligence System Architecture (internal)")
add_bullet("Amira Financial Forecasting API Specification (internal)")
add_bullet("Amira Marketing Analytics API Specification (internal)")

# 1.5 Overview
doc.add_heading("1.5 Document Overview", level=2)
add_para(
    "Section 2 provides an overall description of the product including its context, capabilities, "
    "users, and constraints. Section 3 contains the detailed specific requirements organized by "
    "external interfaces, functional requirements, performance requirements, design constraints, and "
    "system attributes. Sections 4-7 cover data model, system architecture, deployment, and acceptance "
    "criteria. Appendices provide supplementary reference material."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
#  2. OVERALL DESCRIPTION (IEEE 830 Section 2)
# ══════════════════════════════════════════════════════════════════
doc.add_heading("2. Overall Description", level=1)

# 2.1 Product Perspective
doc.add_heading("2.1 Product Perspective", level=2)
add_para(
    "Mars currently operates two separate AI-powered analytical tools: a Period End Summary (PES) "
    "system for generating internal financial performance narratives, and a Competitive Intelligence (CI) "
    "system for analyzing competitor earnings documents. These systems operate independently, require "
    "engineering effort for even minor changes to report templates or data sources, and do not share a "
    "common data layer."
)
add_para(
    "Amira FinIQ consolidates these capabilities into a unified platform while augmenting them with "
    "new functionality: on-demand reporting via natural language queries, an enterprise agent job board "
    "for scalable analytics processing, cross-platform intelligence through integration with Amira's "
    "Forecasting and Marketing Analytics services, a dynamic configurable user interface, and self-service "
    "configuration for business users."
)

doc.add_heading("2.1.1 Current System: Period End Summary (PES)", level=3)
add_bullet("Function: Transforms raw financial Excel data into AI-generated executive performance summaries")
add_bullet("Input: Preprocessed Excel workbooks (P&L, Product, Brand, NCFO sheets) from Azure Blob Storage")
add_bullet("Processing: 6 parallel GPT-4.1 pipelines (one per KPI: Organic Growth, MAC Shape, A&CP Shape, CE Shape, Overhead Shape, NCFO)")
add_bullet("Output: Markdown narratives with trend taglines, sub-unit rankings, and HTML KPI tables")
add_bullet("Performance: 10-15 seconds first generation, < 1 second cached retrieval")
add_bullet("Formats: Summary (balanced), What's Working Well (positive), What's Not Working Well (negative)")

doc.add_heading("2.1.2 Current System: Competitive Intelligence (CI)", level=3)
add_bullet("Function: Ingests competitor earnings documents and generates structured competitive analysis")
add_bullet("Input: PDF documents (press releases, earnings presentations, prepared remarks)")
add_bullet("Pipeline: Upload > Parse/Extract metadata > Chunk/Embed to Azure AI Search > Generate themed summaries > Peer-to-peer benchmarking")
add_bullet("Output: Themed summaries (Organic Growth, Margins, Projections, Consumer, Products, etc.), P2P quantitative tables, Q&A chat")
add_bullet("Infrastructure: Azure Blob Storage, Azure Document Intelligence, Azure OpenAI, Azure AI Search, Cosmos DB")

doc.add_heading("2.1.3 Gap Analysis", level=3)
add_table(
    ["Gap", "Impact", "FinIQ Resolution"],
    [
        ["Separate systems", "Users switch between PES and CI tools; no cross-referencing", "Unified platform with single query interface"],
        ["No ad-hoc queries on internal data", "PES only generates predefined 6-KPI summaries", "Natural language queries across all data"],
        ["Template changes require engineering", "Business users cannot modify report structure", "Self-service template builder"],
        ["No forecast integration", "Reports are backward-looking only", "Amira Forecasting API integration"],
        ["No marketing analytics link", "Cannot correlate financial performance with marketing spend", "Amira Marketing Analytics integration"],
        ["Single-user architecture", "No queue or prioritization for multiple users", "Enterprise job board with agent queue"],
        ["Limited data sources", "Only internal Excel + competitor PDFs", "Extensible connector framework"],
        ["No data lineage", "Cannot trace generated insights back to source", "Full audit trail with source references"],
        ["Static user interface", "No dashboard customization or real-time updates", "Dynamic configurable UI with live updates"],
    ],
    col_widths=[1.8, 2.2, 2.5]
)

# 2.2 Product Functions
doc.add_heading("2.2 Product Functions", level=2)
add_para("Amira FinIQ provides the following high-level capabilities:")
add_bullet("Unified Platform: One hub consolidating two existing tools, with a single data layer and query interface")
add_bullet("On-Demand Reporting: Any financial report generated in seconds from natural language queries")
add_bullet("Enterprise Agent Queue: 100+ users submit queries simultaneously; AI agents process them autonomously with SLAs and priority routing")
add_bullet("Cross-Platform Intelligence: Recommendations enriched by Amira's Financial Forecasting and Marketing Analytics engines")
add_bullet("Dynamic User Interface: Configurable dashboards, real-time updates, adaptive query interface, and responsive design")
add_bullet("Self-Service Configuration: Business users modify report templates, KPI definitions, and data sources without code changes")
add_bullet("Comprehensive Data Coverage: Internal financials, competitor intelligence, commodity markets, acquired research, and third-party data in one place")

# 2.3 User Characteristics
doc.add_heading("2.3 User Characteristics", level=2)

doc.add_heading("2.3.1 Stakeholder Map", level=3)
add_table(
    ["Role", "Description", "Key Interests"],
    [
        ["Executive Leadership", "C-suite, GBU Presidents", "Strategic insights, competitive positioning, board-ready reports"],
        ["Finance Directors", "Regional/divisional finance leads", "Period-end summaries, KPI tracking, variance analysis"],
        ["Strategy & Analytics", "Competitive intelligence analysts", "Competitor benchmarking, market trends, peer comparisons"],
        ["Brand & Marketing", "Brand managers, marketing leads", "Brand performance, marketing ROI, consumer insights"],
        ["IT / Digital", "Platform engineering, data ops", "Architecture, security, integration, scalability"],
        ["Data Science", "Forecasting and modeling teams", "Data quality, model integration, feature requests"],
    ],
    col_widths=[1.5, 2.0, 3.0]
)

doc.add_heading("2.3.2 User Personas", level=3)

add_para("Persona 1: The Finance Director (Primary User)", bold=True)
add_bullet("Needs: Generate period-end summaries for their GBU, compare performance across divisions, drill into KPI drivers")
add_bullet("Pain today: Waits for reports to be generated, cannot customize output format, no cross-referencing with forecasts")
add_bullet("With FinIQ: Self-service report generation, configurable KPIs, integrated forecast comparisons")

add_para("Persona 2: The Strategy Analyst (Power User)", bold=True)
add_bullet("Needs: Deep competitive analysis, cross-competitor benchmarking, ad-hoc queries across all data sources")
add_bullet("Pain today: Manually downloads competitor PDFs, uses separate CI tool, no connection to internal data")
add_bullet("With FinIQ: Unified query across internal + external data, automated competitor monitoring, job queue for complex analyses")

add_para("Persona 3: The Executive (Consumer)", bold=True)
add_bullet("Needs: Board-ready summaries, key takeaways, trend indicators, strategic recommendations")
add_bullet("Pain today: Receives static reports, cannot ask follow-up questions, no forward-looking insights")
add_bullet("With FinIQ: Interactive Q&A, forecast-enriched summaries, one-click export to presentation format")

# 2.4 Constraints
doc.add_heading("2.4 Constraints", level=2)
add_bullet("All infrastructure must be deployed within Mars's Azure tenant (no data egress to public internet)")
add_bullet("LLM service must be Azure OpenAI with data processing agreement ensuring no training on Mars data")
add_bullet("Authentication must integrate with Mars's existing Azure Entra ID (SSO)")
add_bullet("Initial data ingestion limited to preprocessed Excel exports (no direct SAP/ERP integration)")
add_bullet("English-only report generation initially")

# 2.5 Assumptions and Dependencies
doc.add_heading("2.5 Assumptions and Dependencies", level=2)
add_bullet("Mars will provide access to Azure tenant with sufficient quota for Azure OpenAI, AI Search, and compute resources")
add_bullet("Preprocessed Excel workbooks will continue to follow the existing naming convention and schema")
add_bullet("Amira Financial Forecasting and Marketing Analytics APIs will be available and documented")
add_bullet("Mars's organizational hierarchy data will be provided and maintained by Mars IT")
add_bullet("Competitor earnings documents will be sourced via existing public filings and Mars's research subscriptions")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
#  3. SPECIFIC REQUIREMENTS (IEEE 830 Section 3)
# ══════════════════════════════════════════════════════════════════
doc.add_heading("3. Specific Requirements", level=1)

# ── 3.1 External Interface Requirements ──
doc.add_heading("3.1 External Interface Requirements", level=2)

doc.add_heading("3.1.1 User Interfaces", level=3)
add_para(
    "The system shall provide a React-based single-page application (SPA) with responsive design, "
    "accessible via modern web browsers (Chrome, Edge, Safari). The UI shall include:"
)
add_bullet("Interactive dashboards with configurable widget layout")
add_bullet("Conversational query interface (chat-style) with streaming responses")
add_bullet("Job board dashboard showing active jobs, statuses, and SLA compliance")
add_bullet("Administration console for templates, KPIs, peer groups, and user management")
add_bullet("Real-time updates via Server-Sent Events (SSE) for report generation progress and job status")
add_bullet("Embeddable widgets for integration into existing Mars portals")

doc.add_heading("3.1.2 Hardware Interfaces", level=3)
add_para("No direct hardware interfaces. All interaction through web browser on standard desktop/laptop devices.")

doc.add_heading("3.1.3 Software Interfaces", level=3)
add_para("Amira FinIQ integrates with the following external systems:")

add_para("Amira Financial Forecasting API", bold=True)
add_table(
    ["Endpoint", "Method", "Description", "Response"],
    [
        ["/api/v1/forecast/{unit_id}", "GET", "Get forecast for organizational unit", "KPI projections with confidence intervals"],
        ["/api/v1/forecast/{unit_id}/scenarios", "GET", "Get scenario analysis (bull/base/bear)", "Three scenarios with drivers"],
        ["/api/v1/forecast/compare", "POST", "Compare actuals vs forecast", "Variance analysis with explanations"],
        ["/api/v1/forecast/drivers/{kpi_id}", "GET", "Get forecast drivers for a KPI", "Input factors, sensitivities, risk factors"],
    ],
    col_widths=[2.2, 0.8, 2.0, 1.5]
)
add_para("Integration Pattern: Request-response with 30-second timeout. Responses cached for 1 hour. "
         "Fallback: if Forecasting API unavailable, reports generated without forecast data (clearly labeled).", size=10)

add_para("Amira Marketing Analytics API", bold=True)
add_table(
    ["Endpoint", "Method", "Description", "Response"],
    [
        ["/api/v1/marketing/{unit_id}/spend", "GET", "Marketing spend by channel/brand", "Spend breakdown with period comparison"],
        ["/api/v1/marketing/{unit_id}/roi", "GET", "Marketing ROI metrics", "ROAS, CAC, CLV by channel"],
        ["/api/v1/marketing/{unit_id}/attribution", "GET", "Revenue attribution to marketing", "Multi-touch attribution percentages"],
        ["/api/v1/marketing/recommendations", "POST", "Get budget optimization suggestions", "Recommended allocation shifts"],
    ],
    col_widths=[2.5, 0.8, 2.0, 1.2]
)
add_para("Integration Pattern: Same as Forecasting API. Marketing data enriches reports only when the user's "
         "organizational unit has marketing data available and the user has marketing-authorized access.", size=10)

doc.add_heading("3.1.4 Communication Interfaces", level=3)
add_para("Amira FinIQ exposes the following RESTful API for third-party consumers and programmatic access:")
add_table(
    ["Endpoint", "Method", "Description"],
    [
        ["/api/v1/reports/generate", "POST", "Generate a report (PES, competitive, custom)"],
        ["/api/v1/reports/{report_id}", "GET", "Retrieve a generated report"],
        ["/api/v1/query", "POST", "Submit a natural language query"],
        ["/api/v1/jobs", "POST", "Submit a job to the agent queue"],
        ["/api/v1/jobs/{job_id}", "GET", "Check job status and retrieve result"],
        ["/api/v1/jobs/{job_id}/cancel", "POST", "Cancel a queued or in-progress job"],
        ["/api/v1/catalog/units", "GET", "List organizational units (filtered by RBAC)"],
        ["/api/v1/catalog/kpis", "GET", "List available KPI definitions"],
        ["/api/v1/catalog/competitors", "GET", "List tracked competitors and peer groups"],
        ["/api/v1/webhooks", "POST", "Register a webhook for event notifications"],
    ],
    col_widths=[2.5, 0.8, 3.2]
)

doc.add_page_break()

# ── 3.2 Functional Requirements ──
doc.add_heading("3.2 Functional Requirements", level=2)

# ── FR1: Data Ingestion ──
doc.add_heading("3.2.1 FR1: Unified Data Ingestion Engine", level=3)
add_para(
    "The ingestion engine is the foundation of Amira FinIQ. It normalizes data from all sources "
    "into a unified data layer that the analytics engine queries against."
)

add_req_table("FR1.1", "Internal Financial Data Ingestion",
    "The system shall ingest preprocessed Excel workbooks containing P&L, Product, Brand, and NCFO "
    "sheets from Azure Blob Storage. Column renaming, KPI mapping, unit filtering (150+ predefined "
    "organizational units), and derived metric calculations shall be applied automatically. The system "
    "shall support the existing naming convention: preprocessed_output_{Period}_{YearShort}.xlsx.")

add_req_table("FR1.2", "Competitor Document Ingestion",
    "The system shall ingest PDF documents (earnings releases, presentations, transcripts) via upload "
    "or automated scraping. Documents shall be parsed using Azure Document Intelligence, chunked, "
    "embedded, and indexed in a vector search store. Metadata extraction (company, quarter, document type) "
    "shall be automatic.")

add_req_table("FR1.3", "Third-Party Data Connectors",
    "The system shall provide a connector framework for integrating third-party data sources including "
    "commodity price feeds, acquired research databases, and market data providers. Each connector shall "
    "normalize data into the unified schema with source lineage metadata.")

add_req_table("FR1.4", "Data Catalog & Lineage",
    "All ingested data shall be cataloged with metadata (source, ingestion timestamp, schema version, "
    "quality score). Every derived metric and generated insight shall maintain a provenance chain back "
    "to its source data.")

add_req_table("FR1.5", "Incremental & Scheduled Ingestion",
    "The system shall support both manual uploads and scheduled ingestion jobs. Incremental ingestion "
    "shall detect changes and update only affected records. Ingestion status and errors shall be visible "
    "in the admin dashboard.")

# ── FR2: Analytics & Reporting ──
doc.add_heading("3.2.2 FR2: Financial Analytics & Reporting Engine", level=3)
add_para(
    "The analytics engine generates structured financial reports from the unified data layer. "
    "It evolves the current PES system into a more flexible, configurable architecture."
)

add_req_table("FR2.1", "Period End Summary Generation",
    "The system shall generate executive performance summaries for any organizational unit, time period, "
    "and output format. Six KPI narratives (Organic Growth, MAC Shape, A&CP Shape, CE Shape, Controllable "
    "Overhead Shape, NCFO) shall be generated in parallel, with trend analysis taglines injected from "
    "multi-period historical data. Output formats: Summary, What's Working Well, What's Not Working Well.")

add_req_table("FR2.2", "Configurable KPI Framework",
    "Business users shall be able to define, modify, and add KPIs without code changes. Each KPI "
    "definition shall include: formula, data sources, display name, narrative prompt template, and "
    "threshold parameters for trend analysis. KPI changes shall take effect on the next report generation.")

add_req_table("FR2.3", "Sub-Unit Ranking & Driver Analysis",
    "For each KPI, the system shall automatically identify and rank sub-units (child business units) "
    "by performance for both Periodic and YTD timeframes. Driver analysis shall explain what is driving "
    "each sub-unit's performance, with explicit rank labels (RANK 1, TOP 3, BOTTOM 3).")

add_req_table("FR2.4", "Interactive KPI Tables",
    "Every generated summary shall include interactive HTML KPI tables with the underlying data, "
    "enabling users to cross-reference AI-generated narratives with source numbers. Tables shall "
    "support sorting, filtering, and drill-down.")

add_req_table("FR2.5", "Custom Report Builder",
    "Users shall be able to compose custom reports by selecting KPIs, organizational units, time periods, "
    "comparison bases (vs LY, vs Budget, vs Forecast), and output sections. Reports shall be saveable "
    "as templates for reuse.")

add_req_table("FR2.6", "Export Capabilities",
    "Reports shall be exportable to PDF, Word (.docx), PowerPoint (.pptx), Excel (.xlsx), and email. "
    "PowerPoint exports shall use Mars-branded templates with charts and tables pre-formatted for "
    "board presentation.")

# ── FR3: Competitive Intelligence ──
doc.add_heading("3.2.3 FR3: Competitive Intelligence Module", level=3)
add_para(
    "The CI module evolves the current standalone Competitive Intelligence tool into an integrated "
    "component of the analytics hub, with cross-referencing to internal data."
)

add_req_table("FR3.1", "Themed Competitor Summaries",
    "The system shall generate structured summaries of competitor earnings organized by themes: "
    "Organic Growth, Margins, Projections, Consumer Trends, Product Launches, and Miscellaneous. "
    "Each bullet shall include source references (clickable links to source document sections).")

add_req_table("FR3.2", "Peer-to-Peer Benchmarking",
    "The system shall generate quantitative P2P comparison tables across predefined competitor peer "
    "groups (e.g., Petcare segment: Mars, Nestle PetCare, Colgate-Palmolive, Freshpet, IDEXX, "
    "J.M. Smucker). Metrics shall include OG%, Price, Volume, Mix, Adj. Core Operating Profit %, "
    "with both Quarterly and YTD views. Narrative analysis shall identify leaders, laggards, and drivers.")

add_req_table("FR3.3", "Internal-External Cross-Reference",
    "The system shall enable queries that combine internal Mars data with competitor data. Example: "
    "'Compare our Petcare organic growth vs Nestle PetCare and Freshpet for Q2 2024, and show our "
    "forecast for Q3.' This requires joining the unified data layer across internal and external sources.",
    priority="Critical")

add_req_table("FR3.4", "Competitor Monitoring & Alerts",
    "The system shall support automated monitoring of competitor document sources. When new earnings "
    "documents are detected, the system shall automatically ingest, process, and notify relevant users. "
    "Users shall configure alert rules based on competitor, metric thresholds, or keywords.")

# ── FR4: NL Query Interface ──
doc.add_heading("3.2.4 FR4: Natural Language Query Interface", level=3)
add_para(
    "The query interface is the primary user interaction layer. Users ask questions in natural "
    "language and receive structured, sourced answers."
)

add_req_table("FR4.1", "Conversational Query Engine",
    "The system shall accept natural language queries and return structured answers with supporting "
    "data tables and source references. Examples: 'What was Nestle's organic growth in Q2 2024?', "
    "'Show me our Petcare MAC Shape vs last year by division', 'Which competitor had the highest "
    "margin improvement this quarter?'")

add_req_table("FR4.2", "Multi-Turn Conversations",
    "The system shall maintain conversation context across multiple turns, allowing follow-up questions "
    "('Break that down by region', 'What about YTD?', 'Compare that with our forecast'). Conversation "
    "history shall be persisted per user session.")

add_req_table("FR4.3", "Query Intent Classification",
    "The system shall classify queries into types: factual lookup, comparison, trend analysis, "
    "recommendation, report generation, or job submission. Each type shall route to the appropriate "
    "processing pipeline (direct answer, analytics engine, forecasting API, or agent queue).")

add_req_table("FR4.4", "Source Attribution",
    "Every answer shall include source attribution: document name, section, page number (for PDFs), "
    "or cell reference (for Excel data). Users shall be able to click through to the source material.")

# ── FR5: Job Board ──
doc.add_heading("3.2.5 FR5: Enterprise Job Board & Agent Queue", level=3)
add_para(
    "The job board is a critical differentiator. It transforms Amira FinIQ from a single-user query "
    "tool into an enterprise-scale analytics service that can handle hundreds of concurrent requests."
)

add_req_table("FR5.1", "Job Submission",
    "Users shall submit analytical jobs via: (a) the query interface (complex queries auto-routed to queue), "
    "(b) a dedicated job submission form with fields for query, priority, deadline, recipients, and output "
    "format, (c) email integration (forward a question to finiq@mars.com), (d) API for programmatic submission.",
    priority="Critical")

add_req_table("FR5.2", "Agent Pool & Specialization",
    "The system shall maintain a pool of specialized AI agents, each optimized for specific query types: "
    "Period End Summary Agent, Competitive Analysis Agent, Forecasting Agent, Marketing Analytics Agent, "
    "Ad-Hoc Query Agent, and Data Quality Agent. Agents shall be horizontally scalable based on queue depth.")

add_req_table("FR5.3", "Priority Routing & SLA Management",
    "Jobs shall be assigned priority levels (Critical, High, Medium, Low) with corresponding SLA targets. "
    "Critical jobs: < 2 minutes. High: < 10 minutes. Medium: < 30 minutes. Low: < 2 hours. The system "
    "shall monitor SLA compliance and escalate breaches to administrators.")

add_req_table("FR5.4", "Job Lifecycle Management",
    "Each job shall have a visible lifecycle: Submitted > Queued > Assigned (to agent) > Processing > "
    "Review > Completed/Failed. Users shall see real-time status updates. Failed jobs shall include "
    "error details and retry options.")

add_req_table("FR5.5", "Job Dashboard",
    "An enterprise dashboard shall show: active jobs by status, agent utilization, SLA compliance rates, "
    "average processing times by query type, top requestors, and historical trends. Filterable by "
    "department, priority, agent type, and date range.")

add_req_table("FR5.6", "Scheduled & Recurring Jobs",
    "Users shall schedule recurring jobs (e.g., 'Generate Petcare Period End Summary on the 5th of every "
    "month' or 'Run competitive benchmarking weekly for all peer groups'). The scheduler shall support "
    "cron-like expressions with timezone awareness.")

add_req_table("FR5.7", "Collaborative Review & Approval",
    "Jobs producing sensitive or high-visibility reports shall support a review workflow: the AI agent "
    "generates a draft, a designated reviewer approves/edits, and the final version is distributed. "
    "Review assignments shall be configurable per report type and organizational unit.")

# ── FR6: Integration Layer ──
doc.add_heading("3.2.6 FR6: Integration Layer", level=3)
add_para(
    "Amira FinIQ integrates with Amira's existing platform services to provide forward-looking "
    "recommendations alongside historical analysis."
)

add_req_table("FR6.1", "Amira Financial Forecasting Integration",
    "The system shall call the Amira Financial Forecasting API to enrich reports with forward-looking "
    "projections. When a user queries historical performance, the system shall optionally append forecast "
    "data: 'Your Q2 Organic Growth was 6.3%. Based on current trends, our forecast model projects 5.8% "
    "for Q3, with downside risk from cocoa input costs.' Forecasts shall be clearly labeled as projections "
    "with confidence intervals.",
    priority="Critical")

add_req_table("FR6.2", "Amira Marketing Analytics Integration",
    "The system shall call the Amira Marketing Analytics API to correlate financial performance with "
    "marketing spend and effectiveness. Example: 'Royal Canin posted 11.3% OG in Q2. Marketing analytics "
    "shows a 15% increase in digital ad spend with a 2.3x ROAS, suggesting the growth is partially "
    "marketing-driven.' Marketing data shall be available as a dimension in custom reports.",
    priority="Critical")

add_req_table("FR6.3", "Unified Recommendation Engine",
    "When both forecasting and marketing data are available, the system shall generate actionable "
    "recommendations: 'Based on forecasted cocoa price increases of 12% in Q3 and your current marketing "
    "ROI on Confectionery brands, consider shifting 5% of Confectionery marketing budget to Petcare, "
    "which shows stronger pricing power.' Recommendations shall be flagged as AI-generated suggestions.")

add_req_table("FR6.4", "External API Gateway",
    "The system shall expose a RESTful API gateway for third-party integrations. Endpoints shall include: "
    "report generation, job submission, query execution, data catalog browsing, and webhook registration "
    "for event notifications.")

# ── FR7: Admin ──
doc.add_heading("3.2.7 FR7: Administration & Configuration", level=3)

add_req_table("FR7.1", "Template Management",
    "Administrators shall manage report templates through a visual editor. Templates define: sections, "
    "KPI selection, narrative tone, data filters, and output format. Templates shall be versionable "
    "with rollback capability.")

add_req_table("FR7.2", "Organizational Hierarchy Management",
    "The system shall maintain Mars's organizational hierarchy (Mars Inc > GBUs > Divisions > Regions > "
    "Sub-units) with the ability to add, modify, or reorganize units. Hierarchy changes shall automatically "
    "propagate to report generation and access control.")

add_req_table("FR7.3", "Peer Group Configuration",
    "Administrators shall define and modify competitor peer groups for benchmarking. Each peer group "
    "specifies: segment name, included companies, relevant metrics, and data source mappings.")

add_req_table("FR7.4", "Prompt Management",
    "LLM prompts shall be managed as configurable assets, not hardcoded. Each prompt shall have: "
    "system message, dynamic sections (output format selection), user message template, and injected "
    "data variables. Prompt changes shall take effect immediately without deployment.")

add_req_table("FR7.5", "User & Access Management",
    "RBAC aligned with Mars organizational hierarchy. Roles: Admin, Analyst, Viewer, API Consumer. "
    "Data access scoped by organizational unit (a Petcare analyst cannot access Snacking financials "
    "unless explicitly granted). SSO integration with Mars identity provider.")

# ── FR8: Dynamic UI ──
doc.add_heading("3.2.8 FR8: Dynamic User Interface", level=3)
add_para(
    "The user interface shall be dynamic and configurable, enabling users to personalize their "
    "experience and interact with data in real time without page reloads or engineering intervention."
)

add_req_table("FR8.1", "Configurable Dashboard Layout",
    "Users shall be able to arrange, resize, and pin dashboard widgets (KPI cards, charts, tables, "
    "job queue status panels) via drag-and-drop. Dashboard configurations shall be saved per user and "
    "shareable across teams. Administrators shall define default layouts per role.")

add_req_table("FR8.2", "Dynamic Report Viewer",
    "Reports shall render dynamically based on content type. The viewer shall support interactive "
    "charts (zoom, hover tooltips, click-to-drill-down), expandable/collapsible sections, inline data "
    "tables with sorting and filtering, and side-by-side comparisons of multiple reports or time periods.")

add_req_table("FR8.3", "Real-Time UI Updates",
    "The interface shall update in real time via Server-Sent Events (SSE) without requiring page refresh. "
    "This includes: report generation progress (streaming tokens), job status transitions, new data "
    "ingestion notifications, alert triggers, and collaborative review status changes.")

add_req_table("FR8.4", "Adaptive Query Interface",
    "The query interface shall adapt based on user context: suggesting relevant filters and organizational "
    "units, auto-completing entity names (companies, KPIs, periods), displaying recent queries for quick "
    "re-execution, and recommending follow-up questions based on the current conversation context.")

add_req_table("FR8.5", "Theme & Branding Customization",
    "Administrators shall configure UI themes including Mars branding assets, color schemes per GBU or "
    "division, and custom logo placement. The system shall support light and dark display modes. "
    "Branding shall carry through to exported reports and presentations.")

add_req_table("FR8.6", "Responsive & Accessible Design",
    "The UI shall comply with WCAG 2.1 Level AA accessibility standards. The application shall be "
    "responsive across desktop and tablet form factors. Full keyboard navigation shall be supported. "
    "Screen reader compatibility shall be validated for all core workflows.",
    priority="High", category="Non-Functional / UI")

doc.add_page_break()

# ── 3.3 Performance Requirements ──
doc.add_heading("3.3 Performance Requirements", level=2)
add_table(
    ["Scenario", "Requirement"],
    [
        ["Period End Summary (first generation)", "< 15 seconds for 6 parallel KPI narratives + trend analysis"],
        ["Period End Summary (cached)", "< 1 second retrieval from cache"],
        ["Single KPI regeneration", "< 8 seconds"],
        ["Ad-hoc natural language query", "< 10 seconds for simple queries, < 30 seconds for complex multi-source queries"],
        ["Competitive intelligence document ingestion", "< 5 minutes per document (parse, chunk, embed, index)"],
        ["Job queue throughput", "50+ concurrent jobs processing simultaneously"],
        ["Dashboard page load", "< 2 seconds initial load, < 500ms subsequent navigation"],
        ["API response time (p95)", "< 3 seconds for report generation endpoints"],
        ["Dashboard widget update (SSE)", "< 500ms from event to UI update"],
    ],
    col_widths=[3.2, 3.3]
)

# ── 3.4 Design Constraints ──
doc.add_heading("3.4 Design Constraints", level=2)
add_bullet("Azure-native deployment: all services must run on Microsoft Azure within Mars's tenant")
add_bullet("LLM provider: Azure OpenAI (GPT-4.1 or latest available model)")
add_bullet("Orchestration framework: LangGraph for multi-step workflows, LangChain for agent runtime")
add_bullet("Frontend framework: React with TypeScript")
add_bullet("Data transport: RESTful APIs with JSON payloads; SSE for streaming")
add_bullet("Authentication: OAuth 2.0 / OpenID Connect via Azure Entra ID")

# ── 3.5 Software System Attributes ──
doc.add_heading("3.5 Software System Attributes", level=2)

doc.add_heading("3.5.1 Scalability", level=3)
add_bullet("Support 100+ concurrent users initially, scaling to 500+")
add_bullet("Horizontally scalable agent pool: auto-scale based on queue depth (min 5, max 50 agents)")
add_bullet("Data layer shall handle 10TB+ of financial data across all sources")
add_bullet("Vector search index shall support 1M+ document chunks with < 100ms retrieval")

doc.add_heading("3.5.2 Availability & Reliability", level=3)
add_bullet("99.9% uptime SLA during business hours (06:00-22:00 local time, all Mars regions)")
add_bullet("Graceful degradation: if LLM service is unavailable, serve cached reports and queue new requests")
add_bullet("Automated failover for all stateful components (database, cache, search index)")
add_bullet("Maximum Recovery Point Objective (RPO): 1 hour. Maximum Recovery Time Objective (RTO): 4 hours")

doc.add_heading("3.5.3 Security", level=3)
add_bullet("All data encrypted at rest (AES-256) and in transit (TLS 1.3)")
add_bullet("SSO integration with Mars identity provider (Azure Entra ID)")
add_bullet("Row-level security: users see only data for their authorized organizational units")
add_bullet("API authentication via OAuth 2.0 with JWT tokens; API keys for service-to-service")
add_bullet("All LLM interactions logged and auditable; no Mars data used for model training")
add_bullet("Sensitive financial data never leaves Mars's Azure tenant (private endpoints)")
add_bullet("SOC 2 Type II compliance for the Amira platform")

doc.add_heading("3.5.4 Compliance & Data Governance", level=3)
add_bullet("GDPR compliance for any personal data processed (employee names in reports)")
add_bullet("Data retention policies configurable per data source (default: 7 years for financial data)")
add_bullet("Audit log for all data access, report generation, and administrative actions")
add_bullet("Data classification labels (Public, Internal, Confidential, Restricted) enforced across all outputs")

doc.add_heading("3.5.5 Authentication & Authorization", level=3)
add_bullet("Single Sign-On (SSO) via Azure Entra ID (formerly Azure AD), integrated with Mars's existing identity provider")
add_bullet("OAuth 2.0 authorization code flow for web application; client credentials flow for API consumers")
add_bullet("JWT tokens with 1-hour expiry; refresh tokens with 24-hour expiry")
add_bullet("Role hierarchy: Global Admin > GBU Admin > Analyst > Viewer > API Consumer")
add_bullet("Data access scoped to organizational unit: a user authorized for 'GBU Petcare ex Russia' sees only Petcare data")

doc.add_heading("3.5.6 Data Protection", level=3)
add_bullet("All Azure services deployed within Mars's private virtual network (VNet) with private endpoints")
add_bullet("No data egress to public internet; all LLM calls routed through Azure Private Link")
add_bullet("Customer-managed encryption keys (CMK) via Azure Key Vault for data at rest")
add_bullet("Azure OpenAI deployed in Mars's tenant with data processing agreement ensuring no training on Mars data")
add_bullet("PII detection and redaction in generated reports (configurable per data classification level)")

doc.add_heading("3.5.7 Audit & Compliance", level=3)
add_bullet("Complete audit trail: who queried what data, when, and what was returned")
add_bullet("LLM prompt/response logging with configurable retention (default 90 days)")
add_bullet("Automated compliance reports for data access patterns and anomaly detection")
add_bullet("Annual penetration testing and security assessment by third-party auditor")
add_bullet("SOC 2 Type II attestation for Amira platform; Mars responsible for data governance policies")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
#  4. DATA MODEL
# ══════════════════════════════════════════════════════════════════
doc.add_heading("4. Data Model", level=1)

doc.add_heading("4.1 Core Entities", level=2)
add_table(
    ["Entity", "Description", "Key Attributes"],
    [
        ["OrganizationalUnit", "Mars hierarchy node (company, GBU, division, region)", "unit_id, name, parent_id, level, is_active"],
        ["FinancialPeriod", "Reporting period", "period_id, year, period_number, start_date, end_date, type (monthly/quarterly)"],
        ["KPIDefinition", "Configurable KPI specification", "kpi_id, name, formula, data_sources, prompt_template_id, thresholds"],
        ["FinancialFact", "Single financial data point", "fact_id, unit_id, period_id, kpi_id, value, currency, source_file"],
        ["CompetitorCompany", "Tracked competitor", "company_id, name, ticker, peer_groups, is_active"],
        ["CompetitorDocument", "Ingested competitor filing", "doc_id, company_id, period_id, doc_type, file_path, ingestion_status"],
        ["DocumentChunk", "Embedded document fragment", "chunk_id, doc_id, text, embedding_vector, section_title, page_number"],
        ["GeneratedReport", "Cached AI-generated output", "report_id, unit_id, period_id, format, kpi_summaries (JSON), created_at"],
        ["Job", "Agent queue work item", "job_id, user_id, query, priority, status, agent_type, sla_target, result"],
        ["Agent", "AI agent instance", "agent_id, type, status, current_job_id, jobs_completed, avg_processing_time"],
        ["PeerGroup", "Competitor benchmarking group", "group_id, segment_name, companies[], metrics[], is_active"],
        ["PromptTemplate", "Versioned LLM prompt", "template_id, name, version, system_msg, user_msg_template, variables"],
        ["AuditLog", "System activity record", "log_id, user_id, action, entity_type, entity_id, timestamp, details"],
        ["DashboardConfig", "User dashboard layout", "config_id, user_id, widgets[], layout, is_default, shared_with[]"],
    ],
    col_widths=[1.5, 2.0, 3.0]
)

doc.add_heading("4.2 Data Classification", level=2)
add_table(
    ["Data Category", "Classification", "Access Control", "Retention"],
    [
        ["Internal P&L data", "Restricted", "RBAC by organizational unit", "7 years"],
        ["Brand/Product performance", "Confidential", "RBAC by GBU", "7 years"],
        ["Competitor public filings", "Internal", "All authenticated users", "Indefinite"],
        ["Acquired research", "Confidential", "Licensed users only", "Per license"],
        ["Generated reports", "Inherits from source data", "Same as source data RBAC", "3 years"],
        ["Forecast projections", "Confidential", "Forecast-authorized users", "2 years"],
        ["Marketing analytics", "Confidential", "Marketing-authorized users", "3 years"],
        ["Audit logs", "Internal", "Admins only", "5 years"],
    ],
    col_widths=[1.8, 1.3, 2.0, 1.4]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
#  5. SYSTEM ARCHITECTURE
# ══════════════════════════════════════════════════════════════════
doc.add_heading("5. System Architecture", level=1)

doc.add_heading("5.1 High-Level Architecture", level=2)
add_para(
    "Amira FinIQ follows a layered microservices architecture deployed on Microsoft Azure, "
    "aligning with Mars's existing cloud infrastructure. The architecture consists of five layers:"
)

add_para("Layer 1: Presentation Layer", bold=True)
add_bullet("React-based single-page application (SPA) with responsive design")
add_bullet("Dynamic configurable dashboards with drag-and-drop widget layout")
add_bullet("Real-time updates via Server-Sent Events (SSE) for report generation progress")
add_bullet("Interactive dashboards built with a modern charting library")
add_bullet("Embeddable widgets for integration into existing Mars portals")

add_para("Layer 2: API Gateway & Orchestration", bold=True)
add_bullet("Azure API Management as the unified entry point")
add_bullet("LangGraph-based orchestration engine for multi-step analytical workflows")
add_bullet("Request routing: simple queries go direct, complex queries route to agent queue")
add_bullet("Rate limiting, authentication, and request logging")

add_para("Layer 3: Intelligence Layer", bold=True)
add_bullet("LLM Service: Azure OpenAI GPT-4.1 (or latest) for narrative generation")
add_bullet("Embedding Service: Azure OpenAI text-embedding-3-large for document chunking")
add_bullet("Agent Runtime: Specialized agents (PES, CI, Forecast, Marketing, Ad-Hoc, DQ)")
add_bullet("Prompt Registry: Versioned prompt templates stored in Azure Cosmos DB")
add_bullet("RAG Pipeline: Vector search (Azure AI Search) + reranking for document Q&A")

add_para("Layer 4: Data Layer", bold=True)
add_bullet("Unified Data Store: Azure SQL Database for structured financial data (P&L, KPIs)")
add_bullet("Document Store: Azure Blob Storage (raw documents) + Azure AI Search (indexed chunks)")
add_bullet("Cache Layer: Azure Redis Cache for generated reports and KPI summaries (JSON)")
add_bullet("Metadata Catalog: Azure Cosmos DB for data lineage, audit logs, and configuration")

add_para("Layer 5: Integration Layer", bold=True)
add_bullet("Amira Forecasting Connector: REST API client with response caching")
add_bullet("Amira Marketing Analytics Connector: REST API client with metric mapping")
add_bullet("Notification Service: Azure Logic Apps for email, Teams, and webhook notifications")
add_bullet("Export Service: Document generation (PDF, DOCX, PPTX, XLSX) via headless renderer")

doc.add_heading("5.2 Component Topology", level=2)
add_para(
    "The following describes the component topology. A visual diagram will be provided in the "
    "design phase."
)

add_table(
    ["Component", "Technology", "Purpose"],
    [
        ["Web Application", "React + TypeScript", "User interface, dashboards, query chat"],
        ["API Gateway", "Azure API Management", "Routing, auth, rate limiting, logging"],
        ["Orchestration Engine", "Python + LangGraph", "Multi-step workflow coordination"],
        ["Agent Runtime", "Python + LangChain", "Specialized AI agents for each query type"],
        ["LLM Service", "Azure OpenAI GPT-4.1", "Narrative generation, summarization, Q&A"],
        ["Embedding Service", "Azure OpenAI Embeddings", "Document chunking and vector encoding"],
        ["Vector Search", "Azure AI Search", "Semantic document retrieval (RAG)"],
        ["Structured Data", "Azure SQL Database", "Financial data, KPIs, organizational hierarchy"],
        ["Document Store", "Azure Blob Storage", "Raw PDFs, Excel files, generated reports"],
        ["Cache", "Azure Redis Cache", "Generated report caching (< 1s retrieval)"],
        ["Config & Metadata", "Azure Cosmos DB", "Prompts, templates, audit logs, lineage"],
        ["Job Queue", "Azure Service Bus", "Enterprise job board message broker"],
        ["Notifications", "Azure Logic Apps", "Email, Teams, webhook delivery"],
        ["Identity", "Azure Entra ID (AD)", "SSO, RBAC, token management"],
        ["Monitoring", "Azure Monitor + App Insights", "Performance, errors, usage analytics"],
        ["Forecasting API", "Amira Platform (external)", "Forward-looking projections"],
        ["Marketing API", "Amira Platform (external)", "Marketing spend and ROI data"],
    ],
    col_widths=[1.8, 2.0, 2.7]
)

doc.add_heading("5.3 Data Flows", level=2)

add_para("Flow 1: Period End Summary Generation", bold=True)
add_numbered("User selects organizational unit, period, year, and output format")
add_numbered("System checks Redis cache for existing summary")
add_numbered("If cache miss: retrieves Excel file from Blob Storage, loads into pandas DataFrames")
add_numbered("Data preprocessing: column renaming, KPI mapping, unit filtering, derived metric calculations")
add_numbered("Data converted to structured Markdown tables")
add_numbered("Six parallel LLM calls (one per KPI) via LangGraph, each with tailored prompt")
add_numbered("Trend analysis: multi-period data loaded, trend taglines generated")
add_numbered("Six KPI narratives + trend taglines combined into unified Executive Summary")
add_numbered("HTML KPI tables appended for data verification")
add_numbered("Result cached in Redis (JSON) and streamed to frontend via SSE")

add_para("Flow 2: Ad-Hoc Query with Forecast Enrichment", bold=True)
add_numbered("User submits natural language query via chat interface")
add_numbered("Query intent classifier determines type: factual, comparison, trend, recommendation")
add_numbered("For factual/comparison: relevant data retrieved from SQL (internal) and/or AI Search (external)")
add_numbered("LLM generates structured answer with source references")
add_numbered("If forecast enrichment enabled: Amira Forecasting API called with relevant parameters")
add_numbered("If marketing data relevant: Amira Marketing Analytics API called")
add_numbered("Combined response with historical data + forecast + marketing insights returned")
add_numbered("Sources, confidence levels, and data freshness indicators displayed")

add_para("Flow 3: Agent Job Processing", bold=True)
add_numbered("User submits job via UI, email, or API with priority and deadline")
add_numbered("Job enters Azure Service Bus queue with metadata (type, priority, SLA target)")
add_numbered("Scheduler assigns job to appropriate specialized agent based on query type")
add_numbered("Agent retrieves relevant data, executes analysis, generates draft output")
add_numbered("If review workflow enabled: draft routed to designated reviewer")
add_numbered("Completed job stored in database, notification sent to requestor")
add_numbered("Job metrics updated (processing time, SLA compliance, agent utilization)")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
#  6. DEPLOYMENT & INFRASTRUCTURE
# ══════════════════════════════════════════════════════════════════
doc.add_heading("6. Deployment & Infrastructure", level=1)

doc.add_heading("6.1 Infrastructure Requirements", level=2)
add_table(
    ["Component", "Azure Service", "SKU / Configuration"],
    [
        ["Web Application", "Azure App Service", "Premium v3 P2v3 (2 instances)"],
        ["API Gateway", "Azure API Management", "Standard tier"],
        ["Orchestration Workers", "Azure Container Apps", "4 vCPU, 8GB RAM (auto-scale 2-10)"],
        ["Agent Pool", "Azure Container Apps", "4 vCPU, 16GB RAM (auto-scale 5-50)"],
        ["LLM Service", "Azure OpenAI", "GPT-4.1 provisioned throughput (100K tokens/min)"],
        ["Embedding Service", "Azure OpenAI", "text-embedding-3-large"],
        ["Vector Search", "Azure AI Search", "Standard S1 (3 replicas)"],
        ["SQL Database", "Azure SQL Database", "Business Critical, 8 vCores"],
        ["Cache", "Azure Redis Cache", "Premium P1 (6GB)"],
        ["Document Store", "Azure Blob Storage", "Hot tier, 1TB"],
        ["Metadata Store", "Azure Cosmos DB", "Serverless, 1000 RU/s"],
        ["Message Queue", "Azure Service Bus", "Premium (1 messaging unit)"],
        ["Monitoring", "Azure Monitor + App Insights", "Standard"],
        ["Key Vault", "Azure Key Vault", "Standard"],
    ],
    col_widths=[1.5, 2.0, 3.0]
)

doc.add_heading("6.2 Deployment Model", level=2)
add_bullet("CI/CD via Azure DevOps or GitHub Actions")
add_bullet("Blue-green deployment for zero-downtime releases")
add_bullet("Infrastructure as Code (IaC) via Terraform or Bicep")
add_bullet("Staging environment mirroring production for pre-release validation")
add_bullet("Feature flags for gradual rollout of new capabilities")

doc.add_heading("6.3 Environments", level=2)
add_table(
    ["Environment", "Purpose", "Data", "Access"],
    [
        ["Development", "Active development and unit testing", "Synthetic / scrambled data", "Engineering team"],
        ["Staging", "Integration testing and UAT", "Anonymized production data subset", "Engineering + QA + key stakeholders"],
        ["Production", "Live system", "Real financial data", "All authorized users"],
    ],
    col_widths=[1.3, 2.0, 2.0, 1.2]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
#  7. PHASED ROLLOUT
# ══════════════════════════════════════════════════════════════════
doc.add_heading("7. Phased Rollout", level=1)

doc.add_heading("Phase 1: Foundation", level=2)
add_para("Goal: Consolidate current PES and CI tools into a unified platform with core query interface.", bold=True)
add_bullet("Unified data ingestion engine (internal Excel + competitor PDFs)")
add_bullet("Period End Summary generation (all 6 KPIs, 3 output formats)")
add_bullet("Competitive Intelligence module (themed summaries, P2P benchmarking)")
add_bullet("Natural language query interface (conversational, multi-turn)")
add_bullet("Basic job queue (submit, track, retrieve)")
add_bullet("RBAC and SSO integration")
add_bullet("Caching layer for sub-second cached retrieval")
add_bullet("Admin dashboard (templates, KPIs, peer groups)")
add_bullet("Dynamic dashboard with configurable widget layout")

add_para("Phase 1 Deliverables:", bold=True)
add_bullet("Web application deployed to Mars Azure tenant")
add_bullet("Migration of existing PES and CI users")
add_bullet("API documentation and developer portal")
add_bullet("User training materials and onboarding guides")

doc.add_heading("Phase 2: Intelligence", level=2)
add_para("Goal: Add forecasting integration, marketing analytics, and enterprise job board.", bold=True)
add_bullet("Amira Financial Forecasting API integration")
add_bullet("Amira Marketing Analytics API integration")
add_bullet("Unified Recommendation Engine")
add_bullet("Enterprise job board with priority routing, SLAs, and agent specialization")
add_bullet("Scheduled and recurring jobs")
add_bullet("Collaborative review and approval workflows")
add_bullet("Email-based job submission (finiq@mars.com)")
add_bullet("Third-party data connector framework")
add_bullet("Advanced export (PowerPoint with Mars templates, branded PDFs)")

add_para("Phase 2 Deliverables:", bold=True)
add_bullet("Forecast-enriched reports available for all GBUs")
add_bullet("Marketing correlation available for marketing-authorized users")
add_bullet("Job board live with SLA monitoring dashboard")
add_bullet("First third-party data connector (commodity prices or acquired research)")

doc.add_heading("Phase 3: Scale", level=2)
add_para("Goal: Enterprise scale, advanced analytics, and continuous improvement.", bold=True)
add_bullet("Scale to 500+ concurrent users")
add_bullet("Advanced analytics: what-if scenario modeling via query interface")
add_bullet("Custom dashboard builder (drag-and-drop KPI widgets)")
add_bullet("Multi-language report generation")
add_bullet("Automated competitor document scraping and monitoring")
add_bullet("Mobile-responsive optimization")
add_bullet("Advanced AI features: anomaly detection, proactive alerts, trend predictions")
add_bullet("API marketplace for internal Mars teams to build on FinIQ data")
add_bullet("Theme and branding customization per GBU")

add_para("Phase 3 Deliverables:", bold=True)
add_bullet("Full enterprise deployment across all Mars GBUs and regions")
add_bullet("Self-service analytics for 500+ users")
add_bullet("API marketplace with 3+ internal consumer applications")
add_bullet("Measurable reduction in time-to-insight across finance and strategy teams")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
#  8. ACCEPTANCE CRITERIA
# ══════════════════════════════════════════════════════════════════
doc.add_heading("8. Acceptance Criteria", level=1)

add_table(
    ["Requirement", "Acceptance Criterion", "Verification Method"],
    [
        ["FR2.1 Period End Summary", "System generates 6-KPI executive summary for any unit/period in < 15 seconds", "Automated load test with 10 concurrent requests"],
        ["FR3.2 P2P Benchmarking", "Quantitative comparison table matches manual analyst output within 2% tolerance", "Side-by-side validation with current CI tool output"],
        ["FR4.1 Query Engine", "90%+ of test queries return factually correct answers with source citations", "QA test suite of 200 representative queries"],
        ["FR5.3 SLA Compliance", "95%+ of jobs complete within SLA target times over a 30-day period", "SLA monitoring dashboard metrics"],
        ["FR6.1 Forecast Integration", "Forecast data appears in reports within 5 seconds of request", "Integration test with mock and live Forecasting API"],
        ["FR6.2 Marketing Integration", "Marketing metrics correlate correctly with financial periods and units", "Cross-validation with Marketing Analytics dashboard"],
        ["FR8.1 Dynamic Dashboard", "Users can configure, save, and reload custom dashboard layouts", "UI test suite with 10 representative dashboard configurations"],
        ["FR8.3 Real-Time Updates", "Job status and report progress update within 500ms via SSE", "Automated SSE latency test"],
        ["NFR Performance", "Cached report retrieval < 1 second (p99)", "Load test with 100 concurrent cached requests"],
        ["NFR Security", "Zero unauthorized data access in penetration test", "Third-party security audit"],
        ["NFR Availability", "99.9% uptime over 90-day measurement period", "Azure Monitor uptime tracking"],
    ],
    col_widths=[1.5, 3.0, 2.0]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
#  APPENDICES
# ══════════════════════════════════════════════════════════════════
doc.add_heading("Appendix A: KPI Definitions (from PES)", level=1)
add_table(
    ["KPI", "Formula", "What It Measures"],
    [
        ["Organic Growth %", "Net Sales growth excluding FX and M&A", "Top-line revenue momentum"],
        ["MAC Shape %", "Margin After Conversion / Net Sales", "Gross-to-net conversion efficiency"],
        ["A&CP Shape %", "Advertising & Consumer Promotions / GSV 3rd Party", "Promotional spend as % of gross sales"],
        ["CE Shape %", "Controllable Earnings / Net Sales Total", "Bottom-line efficiency including overheads"],
        ["Controllable Overhead Shape %", "Controllable Overhead Costs / Net Sales Total", "Overhead cost management"],
        ["Trade Shape %", "Trade Expenditures / GSV 3rd Party", "Trade spend as % of gross sales"],
        ["Total Growth Impact", "Net Sales Total x Growth % (3rd pty organic)", "Dollar impact of organic growth"],
        ["Periodic vs LY %", "(CY - LY) / LY", "Period-over-period growth rate"],
        ["YTD vs LY %", "(CY - LY) / LY", "Year-to-date growth rate"],
    ],
    col_widths=[1.8, 2.5, 2.2]
)

doc.add_heading("Appendix B: Current System Capabilities", level=1)
add_para(
    "This section documents the capabilities of Mars's existing PES and CI systems that Amira FinIQ "
    "will preserve and enhance. All current capabilities are maintained; enhancements are additive."
)

add_para("Period End Summary (Current)", bold=True)
add_bullet("Input: Preprocessed Excel workbooks with 4 sheets (P&L, Product, Brand, NCFO)")
add_bullet("Processing: 6 parallel GPT-4.1 pipelines with prompt engineering")
add_bullet("Output: Markdown narratives with trend taglines and HTML KPI tables")
add_bullet("Caching: Individual KPI JSONs in Azure Blob Storage")
add_bullet("Performance: 10-15s first generation, < 1s cached, 5-8s single KPI regeneration")

add_para("Competitive Intelligence (Current)", bold=True)
add_bullet("Input: PDF uploads (earnings releases, presentations, prepared remarks)")
add_bullet("Pipeline: Parse > Chunk > Embed > Index > Summarize > P2P Benchmark > Notify")
add_bullet("Output: Themed summaries with source links, P2P quantitative tables, Q&A chat")
add_bullet("Infrastructure: Azure Blob + Document Intelligence + OpenAI + AI Search + Cosmos DB")

# ── Document info footer ──
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("--- End of Document ---")
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.font.italic = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Amira FinIQ SRS v2.0 | Confidential | Prepared for Mars, Incorporated")
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)


# ══════════════════════════════════════════════════════════════════
#  SAVE
# ══════════════════════════════════════════════════════════════════
output_path = r"D:\Amira FinIQ\FinIQ SRS IEEE Format by Claude.docx"
doc.save(output_path)
print(f"SRS document saved to: {output_path}")
