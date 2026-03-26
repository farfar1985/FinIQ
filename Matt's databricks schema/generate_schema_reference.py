"""
Amira FinIQ — Databricks/FinSight Schema Reference Document Generator
Generates a comprehensive Word document cataloging every table, view, column,
data type, relationship, and SQL definition from Matt's FinIQ UC Documentation.

Source: FinIQ UC Documentation (46 pages)
Catalog: corporate_finance_analytics_dev
Schema: finsight_core_model_mvp3
Generated from screenshots taken 2026-03-25
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

# ── Branding colors (matching SRS style) ──
NAVY = RGBColor(0x0A, 0x1F, 0x44)
BLUE = RGBColor(0x1A, 0x56, 0xDB)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
LIGHT_BLUE = RGBColor(0xE8, 0xF0, 0xFE)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
TABLE_BORDER_COLOR = "CCCCCC"
ACCENT = "1A56DB"

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# ── Style definitions ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = DARK_GRAY

for level, (size, color) in enumerate([(22, NAVY), (16, BLUE), (13, BLUE), (12, BLUE)], 1):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Calibri'
    h.font.size = Pt(size)
    h.font.color.rgb = color
    h.font.bold = True
    if level <= 2:
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(8)
    else:
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(6)


# ── Helper functions ──

def add_para(text, bold=False, italic=False, size=None, color=None, space_after=None, alignment=None):
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(size or 11)
    run.font.color.rgb = color or DARK_GRAY
    run.bold = bold
    run.italic = italic
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p


def add_bullet(text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Pt(18 + level * 18)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        run.font.color.rgb = DARK_GRAY
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        run.font.color.rgb = DARK_GRAY
    else:
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        run.font.color.rgb = DARK_GRAY
    return p


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE
        set_cell_shading(cell, ACCENT)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
            run.font.color.rgb = DARK_GRAY
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F5F7FA")

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(width)

    doc.add_paragraph()
    return table


def add_sql_block(sql_text):
    """Add a monospaced SQL code block."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Pt(18)
    run = p.add_run(sql_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
    return p


def add_table_section(table_name, full_name, obj_type, overview, columns, properties, extra_info=None):
    """Add a complete table/view documentation section."""
    doc.add_heading(table_name, level=3)

    # Metadata box
    add_para(f"Full Name: {full_name}", bold=False, size=10, color=DARK_GRAY)
    add_para(f"Type: {obj_type}", bold=True, size=10, color=BLUE, space_after=4)

    # Overview
    doc.add_heading("Overview", level=4)
    add_para(overview, space_after=6)

    # Columns table
    doc.add_heading("Columns", level=4)
    add_table(
        ["Column Name", "Data Type", "Description"],
        columns,
        col_widths=[2.2, 1.5, 2.8]
    )

    # Properties
    doc.add_heading("Table Properties", level=4)
    for key, val in properties.items():
        p = doc.add_paragraph()
        run = p.add_run(f"{key}: ")
        run.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.color.rgb = DARK_GRAY
        run = p.add_run(str(val))
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.color.rgb = DARK_GRAY

    if extra_info:
        for heading, content in extra_info.items():
            doc.add_heading(heading, level=4)
            if isinstance(content, str):
                add_sql_block(content)
            else:
                add_para(content, size=10)

    doc.add_paragraph()  # spacer


# ═══════════════════════════════════════════════════════════════════════════════
# DOCUMENT CONTENT
# ═══════════════════════════════════════════════════════════════════════════════

# ── Cover Page ──
for _ in range(5):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Amira FinIQ")
run.font.size = Pt(36)
run.font.color.rgb = NAVY
run.font.bold = True
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Unified Financial Analytics Hub")
run.font.size = Pt(18)
run.font.color.rgb = BLUE
run.font.name = 'Calibri'

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Databricks / FinSight Schema Reference")
run.font.size = Pt(20)
run.font.color.rgb = NAVY
run.font.bold = True
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Complete Data Layer Documentation")
run.font.size = Pt(14)
run.font.color.rgb = BLUE
run.font.name = 'Calibri'

doc.add_paragraph()
doc.add_paragraph()

info_items = [
    ("Source Document", "FinIQ UC Documentation (46 pages)"),
    ("Source Author", "dipendra.das@effem.com (Mars / Effem)"),
    ("Source Generated", "2026-03-25 14:55:04"),
    ("Catalog", "corporate_finance_analytics_dev"),
    ("Schema", "finsight_core_model_mvp3"),
    ("Table/View Prefix", "finiq"),
    ("Storage Format", "Delta (Azure Blob Storage)"),
    ("Storage Location", "abfss://output@finsightmvp31218devsa.dfs.core.windows.net/"),
    ("Owner Group", "Finsight-Group-Mvp3"),
    ("Reference Compiled By", "Amira Technologies (QDT)"),
    ("Reference Date", datetime.datetime.now().strftime("%Y-%m-%d")),
    ("Classification", "Confidential — Mars / Amira Internal"),
]

t = doc.add_table(rows=len(info_items), cols=2)
t.style = 'Table Grid'
t.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (key, val) in enumerate(info_items):
    cell_k = t.rows[i].cells[0]
    cell_v = t.rows[i].cells[1]
    cell_k.text = ''
    cell_v.text = ''
    run = cell_k.paragraphs[0].add_run(key)
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.font.color.rgb = DARK_GRAY
    set_cell_shading(cell_k, "E8F0FE")
    run = cell_v.paragraphs[0].add_run(val)
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.font.color.rgb = DARK_GRAY
    cell_k.width = Inches(2.5)
    cell_v.width = Inches(4.0)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 1: EXECUTIVE SUMMARY
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading("1. Executive Summary", level=1)

add_para(
    "This document provides a complete reference of the FinIQ data layer as implemented in "
    "Mars's Databricks environment (FinSight Core Model MVP3). It catalogs every table, view, "
    "column, data type, storage property, and SQL definition from the FinIQ UC Documentation "
    "provided by Matt Hutton's team."
)

add_para(
    "The schema consists of 20 objects: 17 Delta tables and 3 precomputed SQL views, organized "
    "into a dimensional model that supports the full scope of Mars's financial analytics — "
    "from P&L reporting and NCFO analysis to brand/product breakdowns and actual-vs-replan "
    "variance analysis. This data layer is the target integration point for the Amira FinIQ "
    "platform, providing direct Databricks connectivity to augment the existing Excel-based "
    "Period End Summary (PES) pipeline."
)

add_para(
    "The schema uses a star/snowflake hybrid design with a fully denormalized wide fact table "
    "(finiq_financial, 39 columns) alongside normalized fact tables (finiq_financial_base, "
    "finiq_financial_cons) and dimension tables for entity hierarchy, account hierarchy, "
    "product taxonomy, customer data, and time. Three precomputed views provide ready-to-consume "
    "P&L and NCFO analytics with Year-to-Date and Periodic comparisons versus Last Year."
)


# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 2: SCHEMA OVERVIEW
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading("2. Schema Overview", level=1)

doc.add_heading("2.1 Connection Details", level=2)

add_table(
    ["Property", "Value"],
    [
        ["Databricks Catalog", "corporate_finance_analytics_dev"],
        ["Schema (Database)", "finsight_core_model_mvp3"],
        ["Table/View Prefix", "finiq"],
        ["Storage Account", "finsightmvp31218devsa"],
        ["Storage Protocol", "abfss:// (Azure Blob File System Secure)"],
        ["Base Path", "abfss://output@finsightmvp31218devsa.dfs.core.windows.net/finsight_core_model_mvp3/finiq/"],
        ["Table Format", "Delta Lake"],
        ["Compute Engine", "Spark"],
        ["Owner Group", "Finsight-Group-Mvp3"],
        ["Overview Generation", "AI with Rule-based Fallback"],
    ],
    col_widths=[2.5, 4.0]
)

doc.add_heading("2.2 Table of Contents — All Objects", level=2)

add_para("The schema contains 20 objects organized as follows:", space_after=6)

add_table(
    ["#", "Object Name", "Type", "Columns", "Category", "Created"],
    [
        ["1",  "finiq_account_formula",        "TABLE", "4",  "Dimension — Account Metadata",    "2026-01-09"],
        ["2",  "finiq_account_input",          "TABLE", "3",  "Dimension — Account Metadata",    "2025-11-04"],
        ["3",  "finiq_composite_item",         "TABLE", "12", "Dimension — Product Taxonomy",    "2025-11-04"],
        ["4",  "finiq_customer",               "TABLE", "11", "Dimension — Customer",            "2025-11-04"],
        ["5",  "finiq_customer_map",           "TABLE", "5",  "Dimension — Customer Hierarchy",  "2025-07-11"],
        ["6",  "finiq_date",                   "TABLE", "4",  "Dimension — Time",                "2025-11-04"],
        ["7",  "finiq_dim_account",            "TABLE", "6",  "Dimension — Account Hierarchy",   "2025-11-04"],
        ["8",  "finiq_dim_entity",             "TABLE", "5",  "Dimension — Org Hierarchy",       "2025-11-04"],
        ["9",  "finiq_economic_cell",          "TABLE", "3",  "Dimension — Economic Cell",       "2025-11-04"],
        ["10", "finiq_financial",              "TABLE", "39", "Fact — Denormalized (Wide)",       "2025-07-11"],
        ["11", "finiq_financial_base",         "TABLE", "7",  "Fact — Normalized (Base)",         "2025-07-11"],
        ["12", "finiq_financial_cons",         "TABLE", "9",  "Fact — Consolidated",              "2025-11-04"],
        ["13", "finiq_financial_replan",       "TABLE", "18", "Fact — Actual vs. Replan",         "2025-11-04"],
        ["14", "finiq_financial_replan_cons",  "TABLE", "6",  "Fact — Consolidated Replan",       "2025-11-04"],
        ["15", "finiq_item",                   "TABLE", "15", "Dimension — Product (Granular)",   "2025-07-11"],
        ["16", "finiq_item_composite_item",    "TABLE", "3",  "Dimension — Item Mapping",         "2025-11-04"],
        ["17", "finiq_rls_last_change",        "TABLE", "2",  "System — Row-Level Security",      "2026-01-08"],
        ["18", "finiq_vw_ncfo_entity",         "VIEW",  "7",  "View — NCFO by Entity",            "2026-03-04"],
        ["19", "finiq_vw_pl_brand_product",    "VIEW",  "8",  "View — P&L by Brand/Product",      "2026-03-16"],
        ["20", "finiq_vw_pl_entity",           "VIEW",  "7",  "View — P&L by Entity",             "2026-03-16"],
    ],
    col_widths=[0.3, 2.3, 0.6, 0.6, 1.8, 0.9]
)


doc.add_heading("2.3 Schema Layers", level=2)

add_para(
    "The data model follows a layered architecture with three distinct tiers:",
    space_after=6
)

add_bullet("Dimension Tables (11 tables): Reference data providing the \"what\" and \"who\" — "
           "organizational hierarchies, account structures, product taxonomy, customer data, "
           "time periods, and economic classifications.")
add_bullet("Fact Tables (5 tables): Transactional financial data providing the \"how much\" — "
           "ranging from a fully denormalized 39-column wide table to normalized base/consolidated "
           "variants, plus actual-vs-replan forecast tables.")
add_bullet("Precomputed Views (3 views): Ready-to-consume analytics providing the \"so what\" — "
           "P&L by entity, P&L by brand/product, and NCFO by entity, all with precomputed "
           "YTD and Periodic comparisons versus Last Year.")
add_bullet("System Table (1 table): Row-level security change tracking for access control.")


doc.add_heading("2.4 Delta Lake Properties", level=2)

add_para(
    "All tables use Delta Lake format with the following common properties. Two generations "
    "of Delta configurations exist in the schema:",
    space_after=6
)

add_table(
    ["Property Set", "Tables", "Configuration"],
    [
        ["Generation 1 (Jul 2025)", "finiq_financial, finiq_financial_base, finiq_customer_map, finiq_item",
         "delta.minReaderVersion=1, delta.minWriterVersion=2"],
        ["Generation 2 (Nov 2025+)", "All other tables",
         "delta.enableDeletionVectors=true, delta.feature.appendOnly=supported, "
         "delta.feature.deletionVectors=supported, delta.feature.invariants=supported, "
         "delta.minReaderVersion=3, delta.minWriterVersion=7"],
    ],
    col_widths=[1.5, 2.5, 2.5]
)

add_para(
    "Generation 2 tables support deletion vectors (enabling efficient row-level deletes/updates), "
    "append-only mode, and invariant enforcement — indicating a more mature data governance posture.",
    italic=True, size=10
)


doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 3: DIMENSION TABLES
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading("3. Dimension Tables", level=1)

add_para(
    "This section documents all 11 dimension tables that provide reference data, hierarchies, "
    "and classifications used by the fact tables and views."
)

# ── 3.1 finiq_date ──
add_table_section(
    "3.1 finiq_date",
    "corporate_finance_analytics_dev.finsight_core_model_mvp3.finiq_date",
    "TABLE",
    "Time dimension table providing the fiscal calendar grain. Date_ID is an integer key "
    "(likely YYYYPP or similar encoding). Year, Period, and Quarter define Mars's fiscal "
    "calendar structure. This table is referenced by all fact tables and views for "
    "time-based filtering and comparison logic.",
    [
        ["Date_ID",  "int",    "Primary key — fiscal date identifier"],
        ["Year",     "int",    "Fiscal year"],
        ["Period",   "string", "Fiscal period within the year"],
        ["Quarter",  "string", "Fiscal quarter"],
    ],
    {
        "Catalog": "corporate_finance_analytics_dev",
        "Database": "finsight_core_model_mvp3",
        "Table": "finiq_date",
        "Created Time": "Tue Nov 04 15:17:14 UTC 2025",
        "Created By": "Spark",
        "Type": "EXTERNAL",
        "Location": "abfss://output@finsightmvp31218devsa.dfs.core.windows.net/finsight_core_model_mvp3/finiq/date",
        "Provider": "delta",
        "Owner": "Finsight-Group-Mvp3",
    }
)

# ── 3.2 finiq_dim_entity ──
add_table_section(
    "3.2 finiq_dim_entity",
    "corporate_finance_analytics_dev.finsight_core_model_mvp3.finiq_dim_entity",
    "TABLE",
    "Organizational entity hierarchy table. Defines the full Mars organizational tree "
    "from Mars Inc down through GBUs, Divisions, Regions, and Sub-units (150+ units). "
    "Entity_Level indicates depth in the hierarchy (e.g., 1 = Mars Inc, 2 = GBU, etc.). "
    "Parent-child relationships enable roll-up reporting at any organizational level.",
    [
        ["Parent_Entity_ID", "string", "Parent organization unit identifier"],
        ["Parent_Entity",    "string", "Parent organization unit name"],
        ["Child_Entity_ID",  "string", "Child organization unit identifier"],
        ["Child_Entity",     "string", "Child organization unit name"],
        ["Entity_Level",     "int",    "Depth in the organizational hierarchy"],
    ],
    {
        "Catalog": "corporate_finance_analytics_dev",
        "Database": "finsight_core_model_mvp3",
        "Table": "finiq_dim_entity",
        "Created Time": "Tue Nov 04 15:17:36 UTC 2025",
        "Created By": "Spark",
        "Type": "EXTERNAL",
        "Location": "abfss://output@finsightmvp31218devsa.dfs.core.windows.net/finsight_core_model_mvp3/finiq/dim_entity",
        "Provider": "delta",
        "Owner": "Finsight-Group-Mvp3",
    }
)

# ── 3.3 finiq_dim_account ──
add_table_section(
    "3.3 finiq_dim_account",
    "corporate_finance_analytics_dev.finsight_core_model_mvp3.finiq_dim_account",
    "TABLE",
    "Account hierarchy dimension. Defines the chart of accounts structure with parent-child "
    "relationships. Parent_Account_ID and Parent_Account are arrays because a single child "
    "account can roll up to multiple parent accounts in the P&L structure. Sign_Conversion "
    "controls positive/negative treatment during financial roll-ups (e.g., expenses may need "
    "sign inversion). Statement is an array indicating which financial statement(s) an account "
    "belongs to (P&L, Balance Sheet, Cash Flow).",
    [
        ["Parent_Account_ID", "array<string>", "Parent account identifier(s) — array for multi-parent roll-up"],
        ["Parent_Account",    "array<string>", "Parent account name(s)"],
        ["Child_Account_ID",  "string",        "Child account identifier"],
        ["Child_Account",     "string",        "Child account name"],
        ["Sign_Conversion",   "int",           "Sign convention for roll-up calculations (+1 or -1)"],
        ["Statement",         "array<string>", "Financial statement(s) this account belongs to"],
    ],
    {
        "Catalog": "corporate_finance_analytics_dev",
        "Database": "finsight_core_model_mvp3",
        "Table": "finiq_dim_account",
        "Created Time": "Tue Nov 04 15:17:15 UTC 2025",
        "Created By": "Spark",
        "Type": "EXTERNAL",
        "Location": "abfss://output@finsightmvp31218devsa.dfs.core.windows.net/finsight_core_model_mvp3/finiq/dim_account",
        "Provider": "delta",
        "Owner": "Finsight-Group-Mvp3",
    }
)

# ── 3.4 finiq_account_formula ──
add_table_section(
    "3.4 finiq_account_formula",
    "corporate_finance_analytics_dev.finsight_core_model_mvp3.finiq_account_formula",
    "TABLE",
    "Account formula definitions. Maps derived/calculated accounts (KPIs) to their component "
    "formulas. The Components array lists the base accounts that feed into each calculated "
    "metric. Old_RL preserves legacy report line references. This table is essential for "
    "understanding how KPIs like Organic Growth, MAC Shape %, and CE Shape % are computed "
    "from base financial data.",
    [
        ["Account",    "string",        "Account name/identifier for the calculated metric"],
        ["Formula",    "string",        "Calculation formula definition"],
        ["Components", "array<string>", "List of base account IDs that feed this formula"],
        ["Old_RL",     "array<string>", "Legacy report line references"],
    ],
    {
        "Catalog": "corporate_finance_analytics_dev",
        "Database": "finsight_core_model_mvp3",
        "Table": "finiq_account_formula",
        "Created Time": "Fri Jan 09 10:28:05 UTC 2026",
        "Created By": "Spark",
        "Type": "EXTERNAL",
        "Location": "abfss://output@finsightmvp31218devsa.dfs.core.windows.net/finsight_core_model_mvp3/finiq/account_formula",
        "Provider": "delta",
        "Owner": "Finsight-Group-Mvp3",
    }
)

# ── 3.5 finiq_account_input ──
add_table_section(
    "3.5 finiq_account_input",
    "corporate_finance_analytics_dev.finsight_core_model_mvp3.finiq_account_input",
    "TABLE",
    "Account input mapping. Associates each Account_ID with its financial Statement type "
    "and Generation level within the account hierarchy. Generation indicates the depth "
    "of the account in the hierarchy tree (0 = root/summary, higher = more granular).",
    [
        ["Statement",  "string", "Financial statement type (e.g., P&L, Cash Flow)"],
        ["Account_ID", "string", "Account identifier"],
        ["Generation",  "int",   "Depth level in account hierarchy (0 = top-level summary)"],
    ],
    {
        "Catalog": "corporate_finance_analytics_dev",
        "Database": "finsight_core_model_mvp3",
        "Table": "finiq_account_input",
        "Created Time": "Tue Nov 04 15:16:54 UTC 2025",
        "Created By": "Spark",
        "Type": "EXTERNAL",
        "Location": "abfss://output@finsightmvp31218devsa.dfs.core.windows.net/finsight_core_model_mvp3/finiq/account_input",
        "Provider": "delta",
        "Owner": "Finsight-Group-Mvp3",
    }
)

# ── 3.6 finiq_composite_item ──
add_table_section(
    "3.6 finiq_composite_item",
    "corporate_finance_analytics_dev.finsight_core_model_mvp3.finiq_composite_item",
    "TABLE",
    "Product master / composite item dimension. Provides the full Mars product taxonomy with "
    "12 attributes covering economic group, technology classification, supply chain technology, "
    "business and market segmentation, brand identification, consumer pack format, and product "
    "categorization. This is the primary product dimension used by the financial fact tables.",
    [
        ["Composite_Item_ID",      "string", "Primary key — composite product identifier"],
        ["EC_Group",               "string", "Economic group classification"],
        ["Technology",             "string", "Technology classification"],
        ["Supply_Tech",            "string", "Supply chain technology"],
        ["Segment",                "string", "Business segment"],
        ["Business_Segment",       "string", "Business segment (detailed)"],
        ["Market_Segment",         "string", "Market segment"],
        ["Brand_ID",               "int",    "Brand numeric identifier"],
        ["Brand",                  "string", "Brand name"],
        ["Consumer_Pack_Format",   "string", "Consumer-facing pack format"],
        ["Product_Consolidation",  "string", "Product consolidation grouping"],
        ["Product_Category",       "string", "Product category"],
    ],
    {
        "Catalog": "corporate_finance_analytics_dev",
        "Database": "finsight_core_model_mvp3",
        "Table": "finiq_composite_item",
        "Created Time": "Tue Nov 04 15:16:31 UTC 2025",
        "Created By": "Spark",
        "Type": "EXTERNAL",
        "Location": "abfss://output@finsightmvp31218devsa.dfs.core.windows.net/finsight_core_model_mvp3/finiq/composite_item",
        "Provider": "delta",
        "Owner": "Finsight-Group-Mvp3",
    }
)

# ── 3.7 finiq_item ──
add_table_section(
    "3.7 finiq_item",
    "corporate_finance_analytics_dev.finsight_core_model_mvp3.finiq_item",
    "TABLE",
    "Granular product/item dimension with 15 attributes. Provides paired ID and Alias columns "
    "for each classification dimension, enabling both join-based lookups (via integer IDs) and "
    "human-readable display (via string aliases). Covers EC Group, Brand Flag, Financial Product "
    "Segment, Market Segment, Supply Technology, Business Segment, and Product Category "
    "Consolidation.",
    [
        ["Item_ID",                              "string", "Primary key — item identifier"],
        ["EC_Group_ID",                          "int",    "Economic group numeric ID"],
        ["EC_Group_Alias",                       "string", "Economic group display name"],
        ["Brand_Flag_ID",                        "int",    "Brand flag numeric ID"],
        ["Brand_Flag_Alias",                     "string", "Brand flag display name"],
        ["Financial_Product_Segment_ID",         "int",    "Financial product segment numeric ID"],
        ["Financial_Product_Segment_Alias",      "string", "Financial product segment display name"],
        ["Market_Segment_ID",                    "int",    "Market segment numeric ID"],
        ["Market_Segment_Alias",                 "string", "Market segment display name"],
        ["Supply_Tech_ID",                       "int",    "Supply technology numeric ID"],
        ["Supply_Tech_Alias",                    "string", "Supply technology display name"],
        ["Business_Segment_ID",                  "string", "Business segment identifier"],
        ["Business_Segment_Alias",               "string", "Business segment display name"],
        ["Product_Category_Consolidation_ID",    "string", "Product category consolidation ID"],
        ["Product_Category_Consolidation_Alias", "string", "Product category consolidation display name"],
    ],
    {
        "Catalog": "corporate_finance_analytics_dev",
        "Database": "finsight_core_model_mvp3",
        "Table": "finiq_item",
        "Created Time": "Fri Jul 11 07:52:24 UTC 2025",
        "Created By": "Spark",
        "Type": "EXTERNAL",
        "Location": "abfss://output@finsightmvp31218devsa.dfs.core.windows.net/finsight_core_model_mvp3/finiq/item",
        "Provider": "delta",
        "Owner": "Finsight-Group-Mvp3",
    }
)

# ── 3.8 finiq_item_composite_item ──
add_table_section(
    "3.8 finiq_item_composite_item",
    "corporate_finance_analytics_dev.finsight_core_model_mvp3.finiq_item_composite_item",
    "TABLE",
    "Junction/bridge table mapping granular items to composite items. Links the finiq_item "
    "dimension (granular product level) to the finiq_composite_item dimension (composite "
    "product level). IT_EC_Group_ID provides the economic group context for the mapping.",
    [
        ["Item_ID",           "string", "Foreign key to finiq_item"],
        ["Composite_Item_ID", "string", "Foreign key to finiq_composite_item"],
        ["IT_EC_Group_ID",    "int",    "Economic group identifier for this mapping"],
    ],
    {
        "Catalog": "corporate_finance_analytics_dev",
        "Database": "finsight_core_model_mvp3",
        "Table": "finiq_item_composite_item",
        "Created Time": "Tue Nov 04 15:19:20 UTC 2025",
        "Created By": "Spark",
        "Type": "EXTERNAL",
        "Location": "abfss://output@finsightmvp31218devsa.dfs.core.windows.net/finsight_core_model_mvp3/finiq/item_composite_item",
        "Provider": "delta",
        "Owner": "Finsight-Group-Mvp3",
    }
)

# ── 3.9 finiq_customer ──
add_table_section(
    "3.9 finiq_customer",
    "corporate_finance_analytics_dev.finsight_core_model_mvp3.finiq_customer",
    "TABLE",
    "Customer dimension with 11 attributes. Provides customer identification, geographic "
    "attribution (Country), a 3-level customer hierarchy (Customer_Level_1/2/3), and "
    "channel/format classification (Customer_Channel, Customer_Format, Customer_Subformat). "
    "SCM_ID links to supply chain management systems. Entity_Customer_ID is the compound "
    "key used by fact tables.",
    [
        ["Entity_Customer_ID",  "string", "Compound key — entity + customer identifier"],
        ["Customer_ID",         "string", "Customer identifier"],
        ["Country",             "string", "Customer country"],
        ["Customer_Name",       "string", "Customer name"],
        ["SCM_ID",              "string", "Supply chain management system ID"],
        ["Customer_Level_1",    "string", "Customer hierarchy level 1 (top)"],
        ["Customer_Level_2",    "string", "Customer hierarchy level 2"],
        ["Customer_Level_3",    "string", "Customer hierarchy level 3 (granular)"],
        ["Customer_Channel",    "string", "Sales channel classification"],
        ["Customer_Format",     "string", "Customer format (e.g., retail format)"],
        ["Customer_Subformat",  "string", "Customer sub-format (granular format)"],
    ],
    {
        "Catalog": "corporate_finance_analytics_dev",
        "Database": "finsight_core_model_mvp3",
        "Table": "finiq_customer",
        "Created Time": "Tue Nov 04 15:16:52 UTC 2025",
        "Created By": "Spark",
        "Type": "EXTERNAL",
        "Location": "abfss://output@finsightmvp31218devsa.dfs.core.windows.net/finsight_core_model_mvp3/finiq/customer",
        "Provider": "delta",
        "Owner": "Finsight-Group-Mvp3",
    }
)

# ── 3.10 finiq_customer_map ──
add_table_section(
    "3.10 finiq_customer_map",
    "corporate_finance_analytics_dev.finsight_core_model_mvp3.finiq_customer_map",
    "TABLE",
    "Customer hierarchy mapping table. Defines parent-child relationships between customers "
    "across different organizational entities. Enables customer-level roll-up reporting "
    "and cross-entity customer analysis. Entity_Customer_ID serves as the bridge key "
    "to the finiq_customer dimension.",
    [
        ["Child_Entity_ID",    "string", "Child entity identifier"],
        ["Child_Customer_ID",  "string", "Child customer identifier"],
        ["Parent_Entity_ID",   "string", "Parent entity identifier"],
        ["Parent_Customer_ID", "string", "Parent customer identifier"],
        ["Entity_Customer_ID", "string", "Combined entity-customer key (FK to finiq_customer)"],
    ],
    {
        "Catalog": "corporate_finance_analytics_dev",
        "Database": "finsight_core_model_mvp3",
        "Table": "finiq_customer_map",
        "Created Time": "Fri Jul 11 07:51:21 UTC 2025",
        "Created By": "Spark",
        "Type": "EXTERNAL",
        "Location": "abfss://output@finsightmvp31218devsa.dfs.core.windows.net/finsight_core_model_mvp3/finiq/customer_map",
        "Provider": "delta",
        "Owner": "Finsight-Group-Mvp3",
    }
)

# ── 3.11 finiq_economic_cell ──
add_table_section(
    "3.11 finiq_economic_cell",
    "corporate_finance_analytics_dev.finsight_core_model_mvp3.finiq_economic_cell",
    "TABLE",
    "Economic cell classification dimension. Provides the economic cell identifier, "
    "descriptive name, and archetype classification. Economic cells represent distinct "
    "economic groupings used in Mars's financial reporting structure.",
    [
        ["Economic_Cell_ID", "int",    "Primary key — economic cell identifier"],
        ["Economic_Cell",    "string", "Economic cell descriptive name"],
        ["Archetype",        "string", "Archetype classification"],
    ],
    {
        "Catalog": "corporate_finance_analytics_dev",
        "Database": "finsight_core_model_mvp3",
        "Table": "finiq_economic_cell",
        "Created Time": "Tue Nov 04 15:17:57 UTC 2025",
        "Created By": "Spark",
        "Type": "EXTERNAL",
        "Location": "abfss://output@finsightmvp31218devsa.dfs.core.windows.net/finsight_core_model_mvp3/finiq/economic_cell",
        "Provider": "delta",
        "Owner": "Finsight-Group-Mvp3",
    }
)


doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 4: FACT TABLES
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading("4. Fact Tables", level=1)

add_para(
    "This section documents all 5 fact tables containing financial transaction data. "
    "The schema provides multiple representations of the same financial data at different "
    "levels of normalization and aggregation."
)

# ── 4.1 finiq_financial ──
add_table_section(
    "4.1 finiq_financial (Denormalized Wide Table — 39 Columns)",
    "corporate_finance_analytics_dev.finsight_core_model_mvp3.finiq_financial",
    "TABLE",
    "The primary denormalized financial fact table. Contains 39 columns with all dimension "
    "attributes inlined — entity hierarchy, account/KPI details, product taxonomy, customer "
    "information, and currency-specific values. This is the most comprehensive single table "
    "in the schema, designed for direct analytical queries without requiring joins to "
    "dimension tables. Contains both USD_Value and Local_Value for multi-currency support.",
    [
        ["Date_ID",              "int",           "Foreign key to finiq_date"],
        ["Year",                 "int",           "Fiscal year (denormalized from date)"],
        ["Period",               "string",        "Fiscal period (denormalized from date)"],
        ["Quarter",              "string",        "Fiscal quarter (denormalized from date)"],
        ["Parent_Entity",        "string",        "Parent organization unit name"],
        ["Entity",               "string",        "Organization unit name"],
        ["Entity_Level",         "int",           "Depth in org hierarchy"],
        ["Parent_Account",       "array<string>", "Parent account name(s)"],
        ["Account_KPI",          "string",        "Account / KPI name"],
        ["Statement",            "array<string>", "Financial statement(s)"],
        ["EC_Group",             "string",        "Economic group"],
        ["Brand",                "string",        "Brand name"],
        ["Segment",              "string",        "Business segment"],
        ["Market_Segment",       "string",        "Market segment"],
        ["Technology",           "string",        "Technology classification"],
        ["Supply_Tech",          "string",        "Supply chain technology"],
        ["Product_Consolidation","string",        "Product consolidation group"],
        ["Product_Category",     "string",        "Product category"],
        ["Business_Segment",     "string",        "Business segment (detailed)"],
        ["Pack_Format",          "string",        "Consumer pack format"],
        ["Economic_Cell",        "string",        "Economic cell name"],
        ["Archetype",            "string",        "Archetype classification"],
        ["Customer_ID",          "string",        "Customer identifier"],
        ["Country",              "string",        "Customer country"],
        ["Customer_Name",        "string",        "Customer name"],
        ["SCM_ID",               "string",        "Supply chain management ID"],
        ["Customer_Level_1",     "string",        "Customer hierarchy level 1"],
        ["Customer_Level_2",     "string",        "Customer hierarchy level 2"],
        ["Customer_Level_3",     "string",        "Customer hierarchy level 3"],
        ["Customer_Channel",     "string",        "Customer sales channel"],
        ["Customer_Format",      "string",        "Customer retail format"],
        ["Customer_Subformat",   "string",        "Customer sub-format"],
        ["Currency",             "string",        "Currency code"],
        ["Sign_Conversion",      "int",           "Sign convention for roll-ups (+1/-1)"],
        ["Entity_ID",            "string",        "Entity identifier (FK)"],
        ["Account_ID",           "string",        "Account identifier (FK)"],
        ["Brand_ID",             "int",           "Brand identifier (FK)"],
        ["USD_Value",            "float",         "Financial value in USD"],
        ["Local_Value",          "float",         "Financial value in local currency"],
    ],
    {
        "Catalog": "corporate_finance_analytics_dev",
        "Database": "finsight_core_model_mvp3",
        "Table": "finiq_financial",
        "Created Time": "Fri Jul 11 07:51:41 UTC 2025",
        "Created By": "Spark",
        "Type": "EXTERNAL",
        "Location": "abfss://output@finsightmvp31218devsa.dfs.core.windows.net/finsight_core_model_mvp3/finiq/financial",
        "Provider": "delta",
        "Owner": "Finsight-Group-Mvp3",
    }
)

# ── 4.2 finiq_financial_base ──
add_table_section(
    "4.2 finiq_financial_base (Normalized Base Fact)",
    "corporate_finance_analytics_dev.finsight_core_model_mvp3.finiq_financial_base",
    "TABLE",
    "Normalized base financial fact table with 7 columns. Uses foreign keys to dimension "
    "tables rather than denormalized attributes. Contains only USD_Value (no local currency). "
    "This is the most storage-efficient representation, requiring joins to dimension tables "
    "for analytical queries.",
    [
        ["Date_ID",            "int",    "Foreign key to finiq_date"],
        ["Entity_ID",          "string", "Foreign key to finiq_dim_entity"],
        ["Account_ID",         "string", "Foreign key to finiq_dim_account"],
        ["Composite_Item_ID",  "string", "Foreign key to finiq_composite_item"],
        ["Economic_Cell_ID",   "int",    "Foreign key to finiq_economic_cell"],
        ["Entity_Customer_ID", "string", "Foreign key to finiq_customer"],
        ["USD_Value",          "float",  "Financial value in USD"],
    ],
    {
        "Catalog": "corporate_finance_analytics_dev",
        "Database": "finsight_core_model_mvp3",
        "Table": "finiq_financial_base",
        "Created Time": "Fri Jul 11 07:52:03 UTC 2025",
        "Created By": "Spark",
        "Type": "EXTERNAL",
        "Location": "abfss://output@finsightmvp31218devsa.dfs.core.windows.net/finsight_core_model_mvp3/finiq/financial_base",
        "Provider": "delta",
        "Owner": "Finsight-Group-Mvp3",
    }
)

# ── 4.3 finiq_financial_cons ──
add_table_section(
    "4.3 finiq_financial_cons (Consolidated Fact)",
    "corporate_finance_analytics_dev.finsight_core_model_mvp3.finiq_financial_cons",
    "TABLE",
    "Consolidated financial fact table with 9 columns. Extends the base fact with Currency_ID "
    "and Local_Value for multi-currency support. This is the primary fact table used by the "
    "precomputed views (vw_ncfo_entity, vw_pl_brand_product, vw_pl_entity) for their SQL "
    "definitions.",
    [
        ["Date_ID",            "int",    "Foreign key to finiq_date"],
        ["Entity_ID",          "string", "Foreign key to finiq_dim_entity"],
        ["Account_ID",         "string", "Foreign key to finiq_dim_account"],
        ["Composite_Item_ID",  "string", "Foreign key to finiq_composite_item"],
        ["Economic_Cell_ID",   "int",    "Foreign key to finiq_economic_cell"],
        ["Entity_Customer_ID", "string", "Foreign key to finiq_customer"],
        ["Currency_ID",        "string", "Currency identifier"],
        ["USD_Value",          "float",  "Financial value in USD"],
        ["Local_Value",        "float",  "Financial value in local currency"],
    ],
    {
        "Catalog": "corporate_finance_analytics_dev",
        "Database": "finsight_core_model_mvp3",
        "Table": "finiq_financial_cons",
        "Created Time": "Tue Nov 04 15:18:17 UTC 2025",
        "Created By": "Spark",
        "Type": "EXTERNAL",
        "Location": "abfss://output@finsightmvp31218devsa.dfs.core.windows.net/finsight_core_model_mvp3/finiq/financial_cons",
        "Provider": "delta",
        "Owner": "Finsight-Group-Mvp3",
    }
)

# ── 4.4 finiq_financial_replan ──
add_table_section(
    "4.4 finiq_financial_replan (Actual vs. Replan)",
    "corporate_finance_analytics_dev.finsight_core_model_mvp3.finiq_financial_replan",
    "TABLE",
    "Forecast/budget comparison table with 18 columns. Contains both Actual and Replan "
    "(budget/forecast) values side-by-side for variance analysis. Includes Submission_Type_ID "
    "to distinguish between different forecast submission types. This table enables "
    "actual-vs-plan variance analysis — a capability not available in the current PES system. "
    "Both USD and Local currency values are provided for actuals and replans.",
    [
        ["Submission_Type_ID", "int",           "Forecast submission type identifier"],
        ["Date_ID",            "int",           "Foreign key to finiq_date"],
        ["Year",               "int",           "Fiscal year"],
        ["Quarter",            "string",        "Fiscal quarter"],
        ["Parent_Entity",      "string",        "Parent organization unit"],
        ["Entity",             "string",        "Organization unit"],
        ["Entity_Level",       "int",           "Depth in org hierarchy"],
        ["Parent_Account",     "array<string>", "Parent account(s)"],
        ["Account_KPI",        "string",        "Account / KPI name"],
        ["Statement",          "array<string>", "Financial statement(s)"],
        ["Currency",           "string",        "Currency code"],
        ["Sign_Conversion",    "int",           "Sign convention for roll-ups"],
        ["Entity_ID",          "string",        "Entity identifier"],
        ["Account_ID",         "string",        "Account identifier"],
        ["Actual_USD_Value",   "float",         "Actual financial value in USD"],
        ["Actual_Local_Value", "float",         "Actual financial value in local currency"],
        ["Replan_USD_Value",   "float",         "Replan/budget value in USD"],
        ["Replan_Local_Value", "float",         "Replan/budget value in local currency"],
    ],
    {
        "Catalog": "corporate_finance_analytics_dev",
        "Database": "finsight_core_model_mvp3",
        "Table": "finiq_financial_replan",
        "Created Time": "Tue Nov 04 15:18:38 UTC 2025",
        "Created By": "Spark",
        "Type": "EXTERNAL",
        "Location": "abfss://output@finsightmvp31218devsa.dfs.core.windows.net/finsight_core_model_mvp3/finiq/financial_replan",
        "Provider": "delta",
        "Owner": "Finsight-Group-Mvp3",
    }
)

# ── 4.5 finiq_financial_replan_cons ──
add_table_section(
    "4.5 finiq_financial_replan_cons (Consolidated Replan)",
    "corporate_finance_analytics_dev.finsight_core_model_mvp3.finiq_financial_replan_cons",
    "TABLE",
    "Consolidated replan fact table with 6 columns. Normalized version of the replan data "
    "using foreign keys to dimension tables. Contains Currency_ID for multi-currency support "
    "and both USD and Local currency values.",
    [
        ["Date_ID",      "int",    "Foreign key to finiq_date"],
        ["Entity_ID",    "string", "Foreign key to finiq_dim_entity"],
        ["Account_ID",   "string", "Foreign key to finiq_dim_account"],
        ["Currency_ID",  "string", "Currency identifier"],
        ["USD_Value",    "float",  "Replan value in USD"],
        ["Local_Value",  "float",  "Replan value in local currency"],
    ],
    {
        "Catalog": "corporate_finance_analytics_dev",
        "Database": "finsight_core_model_mvp3",
        "Table": "finiq_financial_replan_cons",
        "Created Time": "Tue Nov 04 15:18:59 UTC 2025",
        "Created By": "Spark",
        "Type": "EXTERNAL",
        "Location": "abfss://output@finsightmvp31218devsa.dfs.core.windows.net/finsight_core_model_mvp3/finiq/financial_replan_cons",
        "Provider": "delta",
        "Owner": "Finsight-Group-Mvp3",
    }
)


doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 5: PRECOMPUTED VIEWS
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading("5. Precomputed Views", level=1)

add_para(
    "This section documents the 3 SQL views that provide ready-to-consume analytics. "
    "These views precompute YTD (Year-to-Date) and Periodic comparisons versus Last Year "
    "for P&L and NCFO metrics. They directly correspond to the Excel sheets currently "
    "used as input by the Period End Summary (PES) system."
)

add_para(
    "All three views share a common computational pattern:", bold=True, space_after=4
)
add_bullet("They query finiq_financial_cons as the base fact table")
add_bullet("They join to Dimensions_View_Date_Map and Dimensions_Date for time offset logic")
add_bullet("They use Date_Offset = 100 for Last Year and Date_Offset = 0 for Current Year")
add_bullet("They use View_ID = 1 for Periodic values and View_ID = 2 for YTD values")
add_bullet("They join to Dimensions_Entity and Dimensions_Account for alias resolution")
add_bullet("They ROUND all output values to 4 decimal places")
add_bullet("They use GROUP BY ALL for aggregation")
add_bullet("View Schema Mode is COMPENSATION for all three views")

# ── 5.1 finiq_vw_ncfo_entity ──
doc.add_heading("5.1 finiq_vw_ncfo_entity", level=2)

add_para("Full Name: corporate_finance_analytics_dev.finsight_core_model_mvp3.finiq_vw_ncfo_entity", size=10)
add_para("Type: VIEW", bold=True, size=10, color=BLUE, space_after=4)

doc.add_heading("Overview", level=4)
add_para(
    "NCFO (Net Cash From Operations) by entity view. Provides precomputed YTD and Periodic "
    "comparisons versus Last Year for NCFO-related accounts. This view corresponds to the "
    "NCFO sheet in the current PES input Excel workbooks."
)

doc.add_heading("Output Columns", level=4)
add_table(
    ["Column Name", "Data Type", "Description"],
    [
        ["Date_ID",           "int",    "Fiscal date identifier"],
        ["Entity_Alias",      "string", "Organization unit display name (from Dimensions_Entity)"],
        ["Account_Alias",     "string", "Account display name (from Dimensions_Account)"],
        ["YTD_LY_Value",      "double", "Year-to-Date value for Last Year (Date_Offset=100, View_ID=2)"],
        ["YTD_CY_Value",      "double", "Year-to-Date value for Current Year (Date_Offset=0, View_ID=2)"],
        ["Periodic_LY_Value", "double", "Periodic value for Last Year (Date_Offset=100, View_ID=1)"],
        ["Periodic_CY_Value", "double", "Periodic value for Current Year (Date_Offset=0, View_ID=1)"],
    ],
    col_widths=[1.8, 0.8, 3.9]
)

doc.add_heading("NCFO Account Filter", level=4)
add_para(
    "The view filters on the following Account_IDs, which represent Mars's NCFO chart of accounts codes:",
    space_after=4
)
add_table(
    ["Account_ID", "Type"],
    [
        ["CF8129",  "Cash Flow account"],
        ["CF8133",  "Cash Flow account"],
        ["MC8136",  "Management Control account"],
        ["MC8149",  "Management Control account"],
        ["CF8147",  "Cash Flow account"],
        ["MC8913",  "Management Control account"],
        ["S900147", "Summary account"],
        ["MC8100",  "Management Control account"],
        ["MC8902",  "Management Control account"],
        ["672899",  "Numeric account code"],
        ["507499",  "Numeric account code"],
        ["507899",  "Numeric account code"],
        ["508099",  "Numeric account code"],
        ["490099",  "Numeric account code"],
        ["209699",  "Numeric account code"],
    ],
    col_widths=[1.5, 3.0]
)

doc.add_heading("SQL Definition", level=4)
add_sql_block(
    "WITH T1 AS (\n"
    "  SELECT D.Date_ID,\n"
    "         D.Date_ID - V.Source_Date_ID Date_Offset,\n"
    "         V.View_ID,\n"
    "         F.Entity_ID,\n"
    "         F.Account_ID,\n"
    "         SUM(F.USD_Value) Value\n"
    "    FROM FinIQ_Financial_Cons F\n"
    "         INNER JOIN Dimensions_View_Date_Map V ON\n"
    "           V.Target_Date_ID = F.Date_ID AND V.View_ID IN (1, 2)\n"
    "         INNER JOIN Dimensions_Date D ON\n"
    "           D.Date_ID = V.Source_Date_ID\n"
    "           OR D.Date_ID = V.Source_Date_ID + 100\n"
    "    WHERE F.Account_ID IN ('CF8129', 'CF8133', 'MC8136', 'MC8149',\n"
    "          'CF8147', 'MC8913', 'S900147', 'MC8100', 'MC8902')\n"
    "          -- 672899, 507499, 507899, 508099, 490099, 209699\n"
    "    GROUP BY ALL)\n"
    "SELECT T1.Date_ID,\n"
    "       E.Entity_Alias,\n"
    "       A.Account_Alias,\n"
    "       ROUND(SUM(CASE WHEN T1.Date_Offset = 100 AND T1.View_ID = 2\n"
    "                      THEN Value END), 4) YTD_LY_Value,\n"
    "       ROUND(SUM(CASE WHEN T1.Date_Offset = 0 AND T1.View_ID = 2\n"
    "                      THEN Value END), 4) YTD_CY_Value,\n"
    "       ROUND(SUM(CASE WHEN T1.Date_Offset = 100 AND T1.View_ID = 1\n"
    "                      THEN Value END), 4) Periodic_LY_Value,\n"
    "       ROUND(SUM(CASE WHEN T1.Date_Offset = 0 AND T1.View_ID = 1\n"
    "                      THEN Value END), 4) Periodic_CY_Value\n"
    "  FROM T1 INNER JOIN Dimensions_Entity E USING (Entity_ID)\n"
    "          LEFT JOIN Dimensions_Account A USING (Account_ID)\n"
    "  GROUP BY ALL"
)

add_para("View Schema Mode: COMPENSATION", size=10, bold=True, space_after=2)
add_para("View Catalog and Namespace: corporate_finance_analytics_dev.finsight_core_model_mvp3", size=10, space_after=2)
add_para("View Query Output Columns: [Date_ID, Entity_Alias, Account_Alias, YTD_LY_Value, YTD_CY_Value, Periodic_LY_Value, Periodic_CY_Value]", size=10)
add_para("Owner: Finsight-Group-Mvp3 | Created: Wed Mar 04 15:51:05 UTC 2026 | Created By: Spark | Type: VIEW", size=9, italic=True)

doc.add_page_break()

# ── 5.2 finiq_vw_pl_entity ──
doc.add_heading("5.2 finiq_vw_pl_entity", level=2)

add_para("Full Name: corporate_finance_analytics_dev.finsight_core_model_mvp3.finiq_vw_pl_entity", size=10)
add_para("Type: VIEW", bold=True, size=10, color=BLUE, space_after=4)

doc.add_heading("Overview", level=4)
add_para(
    "P&L (Profit & Loss) by entity view. Provides precomputed YTD and Periodic comparisons "
    "versus Last Year for P&L accounts at the entity level. This view corresponds to the "
    "P&L sheet in the current PES input Excel workbooks. Includes a growth calculation "
    "pattern that derives parent account values from child accounts using a hardcoded "
    "growth mapping table."
)

doc.add_heading("Output Columns", level=4)
add_table(
    ["Column Name", "Data Type", "Description"],
    [
        ["Date_ID",           "int",    "Fiscal date identifier"],
        ["Entity_Alias",      "string", "Organization unit display name"],
        ["Account_Alias",     "string", "Account display name"],
        ["YTD_LY_Value",      "double", "Year-to-Date value for Last Year"],
        ["YTD_CY_Value",      "double", "Year-to-Date value for Current Year"],
        ["Periodic_LY_Value", "double", "Periodic value for Last Year"],
        ["Periodic_CY_Value", "double", "Periodic value for Current Year"],
    ],
    col_widths=[1.8, 0.8, 3.9]
)

doc.add_heading("P&L Account Filter", level=4)
add_para("The view filters on the following P&L Account_IDs:", space_after=4)
add_table(
    ["Account_ID", "Likely KPI Area"],
    [
        ["S900123",  "Summary / Total Revenue"],
        ["FR4100",   "Revenue line"],
        ["FR4200",   "Revenue line"],
        ["SR4001",   "Sub-revenue"],
        ["FR4300",   "Revenue line"],
        ["FR4000",   "Revenue line"],
        ["MR5200",   "Margin / Cost line"],
        ["MR5100",   "Margin / Cost line"],
        ["SR5101",   "Sub-margin"],
        ["SR5103",   "Sub-margin"],
        ["MR8005",   "Margin / Overhead"],
        ["SR6102",   "Sub-overhead"],
        ["SR6153",   "Sub-overhead"],
        ["MR8004",   "Margin / Overhead"],
        ["MR6300",   "Margin line"],
        ["MR6359",   "Margin line"],
        ["MR8003",   "Margin line"],
        ["S900144",  "Summary account"],
        ["S900233",  "Summary account"],
        ["S900130",  "Summary account"],
        ["S900227",  "Summary / Growth account"],
        ["S900077",  "Summary / Special account"],
        ["S900067",  "Summary account"],
        ["S900070",  "Summary account"],
        ["S900069",  "Summary account"],
        ["ST9020",   "Statistical account"],
    ],
    col_widths=[1.5, 3.0]
)

doc.add_heading("Growth Calculation Logic", level=4)
add_para(
    "The view includes a growth derivation pattern using a hardcoded mapping table "
    "called 'Growth' that defines parent-child account relationships for growth KPIs. "
    "The Growth table uses three columns: Parent_Account_ID (col1), Child_Account_ID (col2), "
    "and Pos (col3, where 0 = numerator and 1 = denominator).",
    space_after=4
)
add_para("Growth account mappings:", bold=True, space_after=4)
add_table(
    ["Parent_Account_ID", "Child_Account_ID", "Pos (0=Num, 1=Den)", "Meaning"],
    [
        ["S900083", "S900227", "0", "Growth numerator"],
        ["S900071", "S900067", "0", "Growth numerator"],
        ["S900073", "S900070", "0", "Growth numerator"],
        ["S900072", "S900069", "0", "Growth numerator"],
        ["S900083", "S900077", "1", "Growth denominator"],
        ["S900071", "S900077", "1", "Growth denominator"],
        ["S900073", "S900077", "1", "Growth denominator"],
        ["S900072", "S900077", "1", "Growth denominator"],
    ],
    col_widths=[1.5, 1.5, 1.5, 2.0]
)

add_para(
    "The growth value is calculated as: SUM(numerator values) / NULLIF(SUM(denominator values), 0) "
    "minus a CASE adjustment when Parent_Account_ID = 'S900083'. Accounts S900227, S900067, "
    "S900070, S900069, and S900077 are excluded from the direct pass-through to avoid double-counting.",
    italic=True, size=10
)

doc.add_heading("SQL Definition", level=4)
add_sql_block(
    "WITH T1 AS (\n"
    "  SELECT D.Date_ID,\n"
    "         D.Date_ID - V.Source_Date_ID - CASE WHEN F.Account_ID\n"
    "           = 'S900077' THEN 100 ELSE 0 END Date_Offset,\n"
    "         V.View_ID,\n"
    "         F.Entity_ID,\n"
    "         F.Account_ID,\n"
    "         SUM(F.USD_Value) Value\n"
    "    FROM FinIQ_Financial_Cons F\n"
    "         INNER JOIN Dimensions_View_Date_Map V ON\n"
    "           V.Target_Date_ID = F.Date_ID AND V.View_ID IN (1, 2)\n"
    "         INNER JOIN Dimensions_Date D ON\n"
    "           D.Date_ID = V.Source_Date_ID\n"
    "           AND F.Account_ID <> 'S900077'\n"
    "           OR D.Date_ID = V.Source_Date_ID + 100\n"
    "           OR D.Date_ID = V.Source_Date_ID + 200\n"
    "             AND F.Account_ID = 'S900077'\n"
    "    WHERE Account_ID IN ('S900123', 'FR4100', 'FR4200', 'SR4001',\n"
    "          'FR4300', 'FR4000', 'MR5200', 'MR5100', 'SR5101',\n"
    "          'SR5103', 'MR8005', 'SR6102', 'SR6153', 'MR8004',\n"
    "          'MR6300', 'MR6359', 'MR8003', 'S900144', 'S900233',\n"
    "          'S900130', 'S900227', 'S900077', 'S900067', 'S900070',\n"
    "          'S900069', 'ST9020')\n"
    "    GROUP BY ALL),\n"
    "Growth AS (\n"
    "  SELECT col1 Parent_Account_ID, col2 Child_Account_ID, col3 Pos\n"
    "    FROM VALUES\n"
    "      ('S900083','S900227',0), ('S900071','S900067',0),\n"
    "      ('S900073','S900070',0), ('S900072','S900069',0),\n"
    "      ('S900083','S900077',1), ('S900071','S900077',1),\n"
    "      ('S900073','S900077',1), ('S900072','S900077',1)),\n"
    "T2 AS (\n"
    "  SELECT T1.Date_ID, T1.Date_Offset, T1.View_ID, T1.Entity_ID,\n"
    "         Growth.Parent_Account_ID Account_ID, T1.Item,\n"
    "         SUM(CASE WHEN Growth.Pos = 0 THEN T1.Value END) /\n"
    "           NULLIF(SUM(CASE WHEN Growth.Pos = 1 THEN T1.Value END), 0)\n"
    "           - CASE WHEN Growth.Parent_Account_ID = 'S900083'\n"
    "             THEN 1 ELSE 0 END Value\n"
    "    FROM T1 INNER JOIN Growth ON Growth.Child_Account_ID = T1.Account_ID\n"
    "    GROUP BY ALL\n"
    "  UNION ALL\n"
    "  SELECT * FROM T1\n"
    "    WHERE Account_ID NOT IN ('S900227','S900067','S900070','S900069','S900077'))\n"
    "SELECT T2.Date_ID,\n"
    "       E.Entity_Alias, A.Account_Alias,\n"
    "       ROUND(SUM(CASE WHEN T2.Date_Offset=100 AND T2.View_ID=2\n"
    "                      THEN Value END), 4) YTD_LY_Value,\n"
    "       ROUND(SUM(CASE WHEN T2.Date_Offset=0 AND T2.View_ID=2\n"
    "                      THEN Value END), 4) YTD_CY_Value,\n"
    "       ROUND(SUM(CASE WHEN T2.Date_Offset=100 AND T2.View_ID=1\n"
    "                      THEN Value END), 4) Periodic_LY_Value,\n"
    "       ROUND(SUM(CASE WHEN T2.Date_Offset=0 AND T2.View_ID=1\n"
    "                      THEN Value END), 4) Periodic_CY_Value\n"
    "  FROM T2 INNER JOIN Dimensions_Entity E USING (Entity_ID)\n"
    "          LEFT JOIN Dimensions_Account A USING (Account_ID)\n"
    "  GROUP BY ALL"
)

add_para("View Schema Mode: COMPENSATION", size=10, bold=True, space_after=2)
add_para("View Catalog and Namespace: corporate_finance_analytics_dev.finsight_core_model_mvp3", size=10, space_after=2)
add_para("View Query Output Columns: [Date_ID, Entity_Alias, Account_Alias, YTD_LY_Value, YTD_CY_Value, Periodic_LY_Value, Periodic_CY_Value]", size=10)
add_para("Owner: Finsight-Group-Mvp3 | Created: Mon Mar 16 08:51:48 UTC 2026 | Created By: Spark | Type: VIEW", size=9, italic=True)

doc.add_page_break()

# ── 5.3 finiq_vw_pl_brand_product ──
doc.add_heading("5.3 finiq_vw_pl_brand_product", level=2)

add_para("Full Name: corporate_finance_analytics_dev.finsight_core_model_mvp3.finiq_vw_pl_brand_product", size=10)
add_para("Type: VIEW", bold=True, size=10, color=BLUE, space_after=4)

doc.add_heading("Overview", level=4)
add_para(
    "P&L by Brand/Product view. The most complex view in the schema. Extends the entity-level "
    "P&L with a product/brand dimension by joining to finiq_composite_item. Creates three "
    "parallel breakdowns (by Brand, by Product_Category, by Product_Consolidation) using "
    "UNION ALL with COALESCE fallbacks for empty values. Includes the same growth derivation "
    "pattern as vw_pl_entity. This view corresponds to the Product and Brand sheets in the "
    "current PES input Excel workbooks."
)

doc.add_heading("Output Columns", level=4)
add_table(
    ["Column Name", "Data Type", "Description"],
    [
        ["Date_ID",           "int",    "Fiscal date identifier"],
        ["Entity_Alias",      "string", "Organization unit display name"],
        ["Account_Alias",     "string", "Account display name"],
        ["Item",              "string", "Brand name, Product Category, or Product Consolidation"],
        ["YTD_LY_Value",      "double", "Year-to-Date value for Last Year"],
        ["YTD_CY_Value",      "double", "Year-to-Date value for Current Year"],
        ["Periodic_LY_Value", "double", "Periodic value for Last Year"],
        ["Periodic_CY_Value", "double", "Periodic value for Current Year"],
    ],
    col_widths=[1.8, 0.8, 3.9]
)

doc.add_heading("Item Dimension Logic", level=4)
add_para(
    "The view creates the Item column through a 3-way UNION ALL from the base data (T0):",
    space_after=4
)
add_bullet("Pass 1 (Brand): ", bold_prefix="COALESCE(Brand, 'EMPTY BRAND') as Item — ")
add_para("  Joins finiq_financial_cons with finiq_composite_item on Composite_Item_ID to get Brand", size=10)
add_bullet("Pass 2 (Product Category): ", bold_prefix="COALESCE(Product_Category, 'EMPTY PCAT') as Item — ")
add_para("  Same base data, using Product_Category instead", size=10)
add_bullet("Pass 3 (Product Consolidation): ", bold_prefix="COALESCE(Product_Consolidation, 'EMPTY PCONS') as Item — ")
add_para("  Same base data, using Product_Consolidation instead", size=10)

add_para(
    "This means each financial record appears three times in T1 — once for each product "
    "breakdown dimension. The growth calculation and final aggregation then operate on "
    "this tripled dataset.",
    italic=True, size=10, space_after=6
)

doc.add_heading("SQL Definition", level=4)
add_sql_block(
    "WITH T0 AS (\n"
    "  SELECT D.Date_ID,\n"
    "         D.Date_ID - V.Source_Date_ID - CASE WHEN F.Account_ID\n"
    "           = 'S900077' THEN 100 ELSE 0 END Date_Offset,\n"
    "         V.View_ID, F.Entity_ID, F.Account_ID,\n"
    "         I.Brand, I.Product_Category, I.Product_Consolidation,\n"
    "         SUM(F.USD_Value) Value\n"
    "    FROM FinIQ_Financial_Cons F\n"
    "         INNER JOIN FinIQ_Composite_Item I ON\n"
    "           F.Composite_Item_ID = I.Composite_Item_ID\n"
    "         INNER JOIN Dimensions_View_Date_Map V ON\n"
    "           V.Target_Date_ID = F.Date_ID AND V.View_ID IN (1, 2)\n"
    "         INNER JOIN Dimensions_Date D ON\n"
    "           D.Date_ID = V.Source_Date_ID\n"
    "           AND F.Account_ID <> 'S900077'\n"
    "           OR D.Date_ID = V.Source_Date_ID + 100\n"
    "           OR D.Date_ID = V.Source_Date_ID + 200\n"
    "             AND F.Account_ID = 'S900077'\n"
    "    WHERE Account_ID IN (<same 26 P&L account IDs as vw_pl_entity>)\n"
    "    GROUP BY ALL),\n"
    "T1 AS (\n"
    "  SELECT Date_ID, Date_Offset, View_ID, Entity_ID, Account_ID,\n"
    "         COALESCE(Brand, 'EMPTY BRAND') Item, Value\n"
    "    FROM T0\n"
    "  UNION ALL\n"
    "  SELECT Date_ID, Date_Offset, View_ID, Entity_ID, Account_ID,\n"
    "         COALESCE(Product_Category, 'EMPTY PCAT') Item, Value\n"
    "    FROM T0\n"
    "  UNION ALL\n"
    "  SELECT Date_ID, Date_Offset, View_ID, Entity_ID, Account_ID,\n"
    "         COALESCE(Product_Consolidation, 'EMPTY PCONS') Item, Value\n"
    "    FROM T0),\n"
    "Growth AS (<same growth mapping as vw_pl_entity>),\n"
    "T2 AS (\n"
    "  SELECT T1.Date_ID, T1.Date_Offset, T1.View_ID, T1.Entity_ID,\n"
    "         Growth.Parent_Account_ID Account_ID, T1.Item,\n"
    "         SUM(CASE WHEN Growth.Pos=0 THEN T1.Value END) /\n"
    "           NULLIF(SUM(CASE WHEN Growth.Pos=1 THEN T1.Value END), 0)\n"
    "           - CASE WHEN Growth.Parent_Account_ID='S900083'\n"
    "             THEN 1 ELSE 0 END Value\n"
    "    FROM T1 INNER JOIN Growth ON Growth.Child_Account_ID = T1.Account_ID\n"
    "    GROUP BY ALL\n"
    "  UNION ALL\n"
    "  SELECT * FROM T1\n"
    "    WHERE Account_ID NOT IN ('S900227','S900067','S900070',\n"
    "          'S900069','S900077'))\n"
    "SELECT T2.Date_ID, E.Entity_Alias, A.Account_Alias, T2.Item,\n"
    "       ROUND(SUM(CASE WHEN T2.Date_Offset=100 AND T2.View_ID=2\n"
    "                      THEN Value END), 4) YTD_LY_Value,\n"
    "       ROUND(SUM(CASE WHEN T2.Date_Offset=0 AND T2.View_ID=2\n"
    "                      THEN Value END), 4) YTD_CY_Value,\n"
    "       ROUND(SUM(CASE WHEN T2.Date_Offset=100 AND T2.View_ID=1\n"
    "                      THEN Value END), 4) Periodic_LY_Value,\n"
    "       ROUND(SUM(CASE WHEN T2.Date_Offset=0 AND T2.View_ID=1\n"
    "                      THEN Value END), 4) Periodic_CY_Value\n"
    "  FROM T2 INNER JOIN Dimensions_Entity E USING (Entity_ID)\n"
    "          LEFT JOIN Dimensions_Account A USING (Account_ID)\n"
    "  GROUP BY ALL"
)

add_para("View Schema Mode: COMPENSATION", size=10, bold=True, space_after=2)
add_para("View Catalog and Namespace: corporate_finance_analytics_dev.finsight_core_model_mvp3", size=10, space_after=2)
add_para("View Query Output Columns: [Date_ID, Entity_Alias, Account_Alias, Item, YTD_LY_Value, YTD_CY_Value, Periodic_LY_Value, Periodic_CY_Value]", size=10)
add_para("Owner: Finsight-Group-Mvp3 | Created: Mon Mar 16 08:51:28 UTC 2026 | Created By: Spark | Type: VIEW", size=9, italic=True)


doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 6: SYSTEM TABLE
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading("6. System Table", level=1)

add_table_section(
    "6.1 finiq_rls_last_change",
    "corporate_finance_analytics_dev.finsight_core_model_mvp3.finiq_rls_last_change",
    "TABLE",
    "Row-level security (RLS) change tracking table. Records the last modification timestamp "
    "and version number for RLS policy changes. Used by the access control layer to determine "
    "when user permissions were last updated and ensure security policies are current.",
    [
        ["Last_Change", "timestamp", "Timestamp of the most recent RLS policy change"],
        ["Version",     "bigint",    "Version counter for RLS policy changes"],
    ],
    {
        "Catalog": "corporate_finance_analytics_dev",
        "Database": "finsight_core_model_mvp3",
        "Table": "finiq_rls_last_change",
        "Created Time": "Thu Jan 08 15:48:33 UTC 2026",
        "Created By": "Spark",
        "Type": "EXTERNAL",
        "Location": "abfss://output@finsightmvp31218devsa.dfs.core.windows.net/finsight_core_model_mvp3/finiq/finiq_rls_last_change",
        "Provider": "delta",
        "Owner": "Finsight-Group-Mvp3",
    }
)


doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 7: EXTERNAL DEPENDENCIES
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading("7. External Dependencies (Referenced in Views)", level=1)

add_para(
    "The three precomputed views reference the following objects that are NOT part of the "
    "finiq_ prefixed schema but exist elsewhere in the finsight_core_model_mvp3 database. "
    "These are dimension tables used for alias resolution and date offset calculations:",
    space_after=6
)

add_table(
    ["Object Name", "Referenced By", "Purpose", "Join Key(s)"],
    [
        ["Dimensions_View_Date_Map", "All 3 views",
         "Maps Target_Date_ID to Source_Date_ID with View_ID (1=Periodic, 2=YTD). "
         "Enables the Last Year vs Current Year offset logic.",
         "Target_Date_ID = F.Date_ID, View_ID IN (1,2)"],
        ["Dimensions_Date", "All 3 views",
         "Provides Date_ID resolution. Used with Source_Date_ID offsets (+100 for LY, "
         "+200 for special S900077 account).",
         "D.Date_ID = V.Source_Date_ID (and offset variants)"],
        ["Dimensions_Entity", "All 3 views",
         "Provides Entity_Alias (display name) for entity IDs.",
         "USING (Entity_ID)"],
        ["Dimensions_Account", "All 3 views",
         "Provides Account_Alias (display name) for account IDs.",
         "USING (Account_ID) — LEFT JOIN"],
        ["FinIQ_Composite_Item", "vw_pl_brand_product only",
         "Provides Brand, Product_Category, Product_Consolidation for item breakdowns.",
         "F.Composite_Item_ID = I.Composite_Item_ID"],
        ["FinIQ_Financial_Cons", "All 3 views",
         "The base fact table (aliased without underscore prefix in view SQL).",
         "Base FROM table"],
    ],
    col_widths=[1.8, 1.2, 2.0, 1.5]
)

add_para(
    "Note: The view SQL references these objects without the finiq_ prefix (e.g., "
    "'FinIQ_Financial_Cons' instead of 'finiq_financial_cons', 'Dimensions_Entity' instead "
    "of 'finiq_dim_entity'). This suggests either table aliases or a naming convention "
    "difference between the documented schema and the actual view definitions.",
    italic=True, size=10
)


doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 8: RELATIONSHIP MAP
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading("8. Data Relationships", level=1)

add_para(
    "This section documents the foreign key relationships between tables, inferred from "
    "column naming conventions and view SQL join patterns.",
    space_after=6
)

doc.add_heading("8.1 Fact-to-Dimension Relationships", level=2)

add_table(
    ["Fact Table", "Dimension Table", "Join Column(s)", "Cardinality"],
    [
        ["finiq_financial_base", "finiq_date",           "Date_ID",            "Many-to-One"],
        ["finiq_financial_base", "finiq_dim_entity",     "Entity_ID",          "Many-to-One"],
        ["finiq_financial_base", "finiq_dim_account",    "Account_ID",         "Many-to-One"],
        ["finiq_financial_base", "finiq_composite_item", "Composite_Item_ID",  "Many-to-One"],
        ["finiq_financial_base", "finiq_economic_cell",  "Economic_Cell_ID",   "Many-to-One"],
        ["finiq_financial_base", "finiq_customer",       "Entity_Customer_ID", "Many-to-One"],
        ["finiq_financial_cons", "(same as base)",       "(same as base + Currency_ID)", "Many-to-One"],
        ["finiq_financial_replan", "finiq_date",         "Date_ID",            "Many-to-One"],
        ["finiq_financial_replan", "finiq_dim_entity",   "Entity_ID",          "Many-to-One"],
        ["finiq_financial_replan", "finiq_dim_account",  "Account_ID",         "Many-to-One"],
        ["finiq_financial_replan_cons", "(same keys)",   "Date_ID, Entity_ID, Account_ID, Currency_ID", "Many-to-One"],
    ],
    col_widths=[1.8, 1.5, 1.7, 1.0]
)

doc.add_heading("8.2 Dimension-to-Dimension Relationships", level=2)

add_table(
    ["From Table", "To Table", "Join Column(s)", "Relationship"],
    [
        ["finiq_customer_map",        "finiq_customer",       "Entity_Customer_ID", "Many-to-One"],
        ["finiq_item_composite_item", "finiq_item",           "Item_ID",            "Many-to-One"],
        ["finiq_item_composite_item", "finiq_composite_item", "Composite_Item_ID",  "Many-to-One"],
        ["finiq_dim_entity",          "(self-referencing)",    "Parent_Entity_ID → Child_Entity_ID", "Hierarchy"],
        ["finiq_dim_account",         "(self-referencing)",    "Parent_Account_ID → Child_Account_ID", "Hierarchy"],
        ["finiq_customer_map",        "(self-referencing)",    "Parent → Child (Entity + Customer)", "Hierarchy"],
    ],
    col_widths=[2.0, 1.5, 2.0, 1.0]
)


doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 9: PES SYSTEM MAPPING
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading("9. Mapping to Current PES System", level=1)

add_para(
    "This section maps the Databricks schema objects to the current Period End Summary (PES) "
    "system's input and output structures, demonstrating how FinIQ can transition from "
    "Excel-based ingestion to direct Databricks connectivity.",
    space_after=6
)

doc.add_heading("9.1 PES Input Excel Sheets → Databricks Views", level=2)

add_table(
    ["PES Excel Sheet", "Databricks View/Table", "Notes"],
    [
        ["P&L (Entity level)", "finiq_vw_pl_entity",
         "Direct replacement. View provides YTD_LY, YTD_CY, Periodic_LY, Periodic_CY by Entity and Account. "
         "Same comparison structure as PES Excel."],
        ["Product / Brand", "finiq_vw_pl_brand_product",
         "Direct replacement. View adds Item dimension (Brand, Product_Category, Product_Consolidation) "
         "to the P&L structure. Covers both Product and Brand sheets."],
        ["NCFO", "finiq_vw_ncfo_entity",
         "Direct replacement. View provides NCFO metrics by Entity with same YTD/Periodic LY/CY structure."],
        ["(4th sheet — varies)", "finiq_financial + dimensions",
         "The 4th PES input sheet varies by report. Ad-hoc queries against the full fact table "
         "with dimension joins can replicate any additional sheet."],
    ],
    col_widths=[1.5, 2.0, 3.0]
)

doc.add_heading("9.2 PES KPIs → Databricks Account Codes", level=2)

add_para(
    "The 6 KPIs generated by the current PES system map to account hierarchies defined "
    "in finiq_dim_account and the account codes filtered in the views:",
    space_after=6
)

add_table(
    ["PES KPI", "Relevant Account Codes", "Derivation"],
    [
        ["Organic Growth", "S900083, S900227, S900067, S900070, S900069, S900077",
         "Calculated via Growth mapping in vw_pl_entity (numerator/denominator pattern)"],
        ["MAC Shape %", "FR4100, FR4200, FR4300, FR4000, MR5200, MR5100",
         "Revenue and margin accounts in P&L view"],
        ["A&CP Shape %", "SR5101, SR5103, MR8005",
         "Advertising & consumer promotion accounts"],
        ["CE Shape %", "SR6102, SR6153, MR8004",
         "Controllable expense / overhead accounts"],
        ["Controllable Overhead Shape %", "MR6300, MR6359, MR8003",
         "Overhead margin accounts"],
        ["NCFO", "CF8129, CF8133, MC8136, MC8149, CF8147, MC8913, S900147, MC8100, MC8902",
         "Cash flow and management control accounts in vw_ncfo_entity"],
    ],
    col_widths=[1.8, 2.5, 2.2]
)

doc.add_heading("9.3 New Capabilities (Not in PES)", level=2)

add_para(
    "The Databricks schema provides capabilities beyond what the current PES system offers:",
    space_after=4
)

add_bullet("Actual vs. Replan Variance: ", bold_prefix="finiq_financial_replan — ")
add_para("  Contains Actual_USD/Local_Value alongside Replan_USD/Local_Value for budget variance analysis", size=10)
add_bullet("Customer-Level Analytics: ", bold_prefix="finiq_customer + finiq_customer_map — ")
add_para("  11-attribute customer dimension with 3-level hierarchy and channel/format classification", size=10)
add_bullet("Multi-Currency Native: ", bold_prefix="finiq_financial_cons — ")
add_para("  Both USD_Value and Local_Value with Currency_ID, enabling currency-specific reporting", size=10)
add_bullet("Full Product Taxonomy: ", bold_prefix="finiq_composite_item + finiq_item — ")
add_para("  27 combined product attributes (12 + 15) vs. PES's limited Brand/Product breakdown", size=10)
add_bullet("Account Formula Definitions: ", bold_prefix="finiq_account_formula — ")
add_para("  Programmatic access to KPI calculation logic, enabling dynamic KPI configuration", size=10)
add_bullet("Row-Level Security: ", bold_prefix="finiq_rls_last_change — ")
add_para("  Built-in RLS tracking for access control governance", size=10)


doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 10: OBSERVATIONS AND NOTES
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading("10. Observations and Technical Notes", level=1)

doc.add_heading("10.1 Schema Evolution Timeline", level=2)

add_table(
    ["Date", "Objects Created", "Significance"],
    [
        ["Jul 11, 2025", "finiq_financial, finiq_financial_base, finiq_customer_map, finiq_item",
         "Initial schema — core wide fact table, base fact, and early dimensions. Delta v1 properties."],
        ["Nov 04, 2025", "finiq_account_input, finiq_composite_item, finiq_customer, finiq_date, "
         "finiq_dim_account, finiq_dim_entity, finiq_economic_cell, finiq_financial_cons, "
         "finiq_financial_replan, finiq_financial_replan_cons, finiq_item_composite_item",
         "Major expansion — all remaining dimensions, consolidated facts, and replan tables. Delta v3/v7 properties."],
        ["Jan 08-09, 2026", "finiq_rls_last_change, finiq_account_formula",
         "Security and formula metadata additions."],
        ["Mar 04, 2026", "finiq_vw_ncfo_entity",
         "First precomputed view — NCFO analytics."],
        ["Mar 16, 2026", "finiq_vw_pl_brand_product, finiq_vw_pl_entity",
         "P&L analytics views — completing the PES-equivalent analytical layer."],
        ["Mar 25, 2026", "(Documentation generated)",
         "This FinIQ UC Documentation generated by dipendra.das@effem.com."],
    ],
    col_widths=[1.0, 2.5, 3.0]
)

doc.add_heading("10.2 Naming Conventions", level=2)

add_bullet("All objects use the finiq_ prefix")
add_bullet("Dimension tables: finiq_dim_* (for hierarchy dimensions) or finiq_* (for flat dimensions)")
add_bullet("Fact tables: finiq_financial* (with suffixes: _base, _cons, _replan, _replan_cons)")
add_bullet("Views: finiq_vw_* (with descriptive suffixes: _ncfo_entity, _pl_entity, _pl_brand_product)")
add_bullet("Bridge/junction tables: finiq_*_map or finiq_*_composite_*")
add_bullet("System tables: finiq_rls_*")

doc.add_heading("10.3 Data Types Used", level=2)

add_table(
    ["Data Type", "Usage", "Count"],
    [
        ["string",        "Identifiers, names, codes, classifications", "Most common"],
        ["int",           "Numeric IDs, levels, sign conversion, Date_ID, Year", "Common"],
        ["float",         "Financial values (USD_Value, Local_Value) in fact tables", "Core metrics"],
        ["double",        "Financial values in view outputs (ROUND result type)", "View outputs"],
        ["array<string>", "Multi-value fields (Parent_Account, Statement, Components)", "Hierarchy arrays"],
        ["timestamp",     "RLS change tracking", "1 column"],
        ["bigint",        "RLS version counter", "1 column"],
    ],
    col_widths=[1.2, 3.5, 1.0]
)

doc.add_heading("10.4 View SQL Pattern Notes", level=2)

add_bullet("Date_Offset = 100 means Last Year; Date_Offset = 0 means Current Year")
add_bullet("View_ID = 1 means Periodic view; View_ID = 2 means YTD view")
add_bullet("Account S900077 receives special treatment: Date_Offset starts at 100 instead of 0, "
           "and uses +200 offset for its LY comparison")
add_bullet("Growth accounts (S900083, S900071, S900073, S900072) are derived via a "
           "numerator/denominator pattern, not stored directly")
add_bullet("COALESCE with sentinel values ('EMPTY BRAND', 'EMPTY PCAT', 'EMPTY PCONS') "
           "handles NULL product attributes in vw_pl_brand_product")
add_bullet("All views use LEFT JOIN for Dimensions_Account (some accounts may lack aliases) "
           "but INNER JOIN for Dimensions_Entity (all entities must resolve)")
add_bullet("All monetary values are ROUND'd to 4 decimal places in view outputs")
add_bullet("GROUP BY ALL is used throughout (Databricks SQL extension)")


# ═══════════════════════════════════════════════════════════════════════════════
# FOOTER
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("— End of Document —")
run.font.size = Pt(12)
run.font.color.rgb = NAVY
run.font.bold = True
run.font.name = 'Calibri'

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f"Generated {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')} by Amira Technologies (QDT)")
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.font.name = 'Calibri'


# ── Save ──
output_path = r"D:\Amira FinIQ\Matt's databricks schema\FinIQ Databricks Schema Reference.docx"
doc.save(output_path)
print(f"Document saved to: {output_path}")
