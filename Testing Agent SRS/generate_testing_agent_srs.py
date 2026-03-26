"""
Amira FinIQ Testing Agent — Software Requirements Specification (SRS) Generator
IEEE 830 / ISO/IEC/IEEE 29148 Format
Version 1.1 — Quantitative Evaluation Framework (Karpathy methodology)
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

# ── Branding colors ──
NAVY = RGBColor(0x0A, 0x1F, 0x44)       # Dark navy for titles
BLUE = RGBColor(0x1A, 0x56, 0xDB)       # Amira blue for headings
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)  # Body text
LIGHT_BLUE = RGBColor(0xE8, 0xF0, 0xFE) # Table header bg
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
TABLE_BORDER_COLOR = "CCCCCC"
ACCENT = "1A56DB"

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# ── Style definitions ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = DARK_GRAY

for level, (size, color) in enumerate([(22, NAVY), (16, BLUE), (13, BLUE), (11, DARK_GRAY)], 1):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Calibri'
    h.font.size = Pt(size)
    h.font.color.rgb = color
    h.font.bold = True
    if level <= 2:
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(8)

# ── Helper functions ──

def add_title_page():
    """Create a professional cover page."""
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Amira FinIQ Testing Agent")
    run.font.size = Pt(36)
    run.font.color.rgb = NAVY
    run.font.bold = True
    run.font.name = 'Calibri'

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Software Requirements Specification")
    run.font.size = Pt(18)
    run.font.color.rgb = BLUE
    run.font.name = 'Calibri'

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("IEEE 830 / ISO/IEC/IEEE 29148 Format")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.font.name = 'Calibri'

    doc.add_paragraph()
    doc.add_paragraph()

    # Metadata table
    meta = [
        ("Document Version", "1.1"),
        ("Date", datetime.date.today().strftime("%B %d, %Y")),
        ("Classification", "Confidential"),
        ("Prepared For", "Mars, Incorporated"),
        ("Prepared By", "Amira Technologies"),
        ("Standard", "IEEE 830-1998 / ISO/IEC/IEEE 29148:2018"),
    ]
    table = doc.add_table(rows=len(meta), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(meta):
        for j, text in enumerate([k, v]):
            cell = table.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            run.font.size = Pt(10)
            run.font.color.rgb = DARK_GRAY
            run.font.name = 'Calibri'
            if j == 0:
                run.font.bold = True

    doc.add_page_break()


def add_para(text, bold=False, italic=False, size=None, color=None, space_after=6):
    """Add a paragraph with formatting."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.font.bold = True
    if italic:
        run.font.italic = True
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_bullet(text, level=0):
    """Add a bullet point."""
    p = doc.add_paragraph(text, style='List Bullet')
    if level > 0:
        p.paragraph_format.left_indent = Inches(0.5 * (level + 1))
    p.paragraph_format.space_after = Pt(2)
    return p


def add_numbered(text):
    """Add a numbered item."""
    p = doc.add_paragraph(text, style='List Number')
    p.paragraph_format.space_after = Pt(2)
    return p


def add_table(headers, rows, col_widths=None):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE
        run.font.name = 'Calibri'
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{ACCENT}"/>')
        cell._element.get_or_add_tcPr().append(shading)

    # Data rows
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.cell(i + 1, j)
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.name = 'Calibri'
            run.font.color.rgb = DARK_GRAY
            if i % 2 == 1:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F6FC"/>')
                cell._element.get_or_add_tcPr().append(shading)

    # Borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="{TABLE_BORDER_COLOR}"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="{TABLE_BORDER_COLOR}"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{TABLE_BORDER_COLOR}"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="{TABLE_BORDER_COLOR}"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{TABLE_BORDER_COLOR}"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="{TABLE_BORDER_COLOR}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

    if col_widths:
        for i_row in range(len(rows) + 1):
            for j, w in enumerate(col_widths):
                table.cell(i_row, j).width = Inches(w)

    doc.add_paragraph()  # spacing
    return table


def add_req_table(req_id, title, description, priority="High", category="Functional"):
    """Add a requirement in structured format."""
    p = doc.add_paragraph()
    run = p.add_run(f"{req_id}: {title}")
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = BLUE
    p.paragraph_format.space_after = Pt(4)
    add_para(description, space_after=2)
    p2 = doc.add_paragraph()
    run2 = p2.add_run(f"Priority: {priority}  |  Category: {category}")
    run2.font.size = Pt(9)
    run2.font.italic = True
    run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    p2.paragraph_format.space_after = Pt(10)


# ══════════════════════════════════════════════════════════════════
#  BUILD THE DOCUMENT (IEEE 830 Format)
# ══════════════════════════════════════════════════════════════════

add_title_page()

# ── TABLE OF CONTENTS placeholder ──
doc.add_heading("Table of Contents", level=1)
add_para(
    "[Table of Contents will auto-generate when opened in Word: References > Update Table]",
    italic=True, color=RGBColor(0x99, 0x99, 0x99)
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
#  1. INTRODUCTION (IEEE 830 Section 1)
# ══════════════════════════════════════════════════════════════════
doc.add_heading("1. Introduction", level=1)

# 1.1 Purpose
doc.add_heading("1.1 Purpose", level=2)
add_para(
    "This Software Requirements Specification (SRS) defines the functional and non-functional "
    "requirements for the Amira FinIQ Testing Agent, a specialized AI agent that validates the "
    "Amira FinIQ platform by executing test scenarios against both simulated (synthetic) and "
    "real (production) data. The Testing Agent ensures that FinIQ's analytics, reporting, and "
    "query capabilities produce correct, consistent results before and after deployment."
)
add_para(
    "This SRS is separate from the main Amira FinIQ SRS (v3.0) and defines the Testing Agent's "
    "own requirements. It is intended for use by Amira Technologies' QA engineers, developers, "
    "product owners, and Mars's technical stakeholders who oversee platform quality."
)
add_para(
    "This document follows the IEEE 830-1998 and ISO/IEC/IEEE 29148:2018 standards for software "
    "requirements specifications."
)

# 1.2 Scope
doc.add_heading("1.2 Scope", level=2)
add_para(
    "The FinIQ Testing Agent is a purpose-built validation component that operates within the "
    "Amira FinIQ agent ecosystem. Its primary mission is to verify correctness of all platform "
    "capabilities across FinIQ's 50 functional requirements."
)

doc.add_heading("1.2.1 In Scope", level=3)
add_bullet("Automated test execution against all FinIQ functional requirement groups (FR1\u2013FR8)")
add_bullet("Dual-mode operation: simulated data mode (synthetic/deterministic) and real data mode (production Databricks)")
add_bullet("Validation of Period End Summary generation, including all 6 KPIs, sub-unit rankings, and trend taglines")
add_bullet("Validation of budget variance reporting (actual vs. replan vs. forecast)")
add_bullet("Natural language query accuracy and source attribution testing")
add_bullet("Competitive intelligence module validation (themed summaries, P2P benchmarking, cross-reference)")
add_bullet("Job board lifecycle, SLA compliance, and scheduling tests")
add_bullet("Prompt-response pair validation for real data (user-provided expected results)")
add_bullet("Regression detection by comparing current outputs against stored baselines")
add_bullet("Test coverage reporting mapped to all 50 main SRS functional requirements")
add_bullet("Test report generation in JSON, HTML, and PDF formats")
add_bullet("Integration with CI/CD pipeline (GitHub Actions webhooks)")
add_bullet("On-demand test execution via the FinIQ job board")

doc.add_heading("1.2.2 Out of Scope", level=3)
add_bullet("Performance and load testing (separate concern, not covered by this agent)")
add_bullet("Security penetration testing (separate concern)")
add_bullet("UI/UX visual testing and accessibility auditing (separate concern)")
add_bullet("Test case authoring tools (test cases are managed externally and provided to the agent)")

# 1.3 Definitions
doc.add_heading("1.3 Definitions, Acronyms, and Abbreviations", level=2)
add_table(
    ["Term", "Definition"],
    [
        ["Testing Agent", "The specialized AI agent defined by this SRS that validates FinIQ platform correctness"],
        ["Test Case", "A single unit of validation consisting of inputs, execution steps, and expected results"],
        ["Test Suite", "An ordered collection of test cases grouped by functional area"],
        ["Expected Result", "The correct output for a given test case, either computed from synthetic data or provided by the user"],
        ["Actual Result", "The output produced by FinIQ during test execution"],
        ["Simulated Mode", "Testing mode using synthetic data (finiq_synthetic.db or Databricks workspace.default) with deterministic expected values"],
        ["Real Mode", "Testing mode using production Databricks data with user-provided prompt-response pairs as expected results"],
        ["Regression", "A defect where previously correct behavior becomes incorrect after a code or data change"],
        ["Coverage", "The percentage of main SRS functional requirements that have at least one mapped test case"],
        ["Prompt-Response Pair", "A JSON structure pairing a natural language prompt with its expected output, used in real data mode"],
        ["Golden Dataset", "A versioned, approved set of expected outputs used as the baseline for regression detection"],
        ["Tolerance Threshold", "The maximum acceptable deviation between expected and actual numeric values (default: 0.01%)"],
        ["PES", "Period End Summary \u2014 AI-generated executive performance narratives"],
        ["KPI", "Key Performance Indicator"],
        ["P2P", "Peer-to-Peer benchmarking (competitive comparison tables)"],
        ["CI/CD", "Continuous Integration / Continuous Deployment"],
        ["SLA", "Service Level Agreement"],
        ["FR", "Functional Requirement (from main FinIQ SRS)"],
        ["TR", "Test Requirement (defined in this SRS)"],
        ["Scalar Metric", "A single numeric score (0\u2013100%) that summarizes pass rate for a capability"],
        ["Eval Harness", "The immutable scoring infrastructure that the agent under test cannot modify"],
        ["Keep-or-Revert", "Evaluation loop: if score improves after a code change, commit; if score drops, revert"],
        ["Binary Criterion", "A pass/fail check with no subjective judgment \u2014 outcome is 0 or 1"],
    ],
    col_widths=[2.0, 4.5]
)

# 1.4 References
doc.add_heading("1.4 References", level=2)
add_table(
    ["Reference", "Description"],
    [
        ["Amira FinIQ SRS v3.0", "Base functional requirements specification (50 FRs across FR1\u2013FR8)"],
        ["FinIQ Databricks Schema Reference", "Comprehensive reference for all 20 finiq_ tables/views"],
        ["FinIQ SRS Addendum A", "Databricks Integration addendum (4 amended + 4 new requirements)"],
        ["IEEE 829-2008", "Standard for Software and System Test Documentation"],
        ["IEEE 830-1998", "Recommended Practice for Software Requirements Specifications"],
        ["ISO/IEC/IEEE 29148:2018", "Systems and software engineering \u2014 Life cycle processes \u2014 Requirements engineering"],
        ["finiq_synthetic.db", "SQLite synthetic data file (local development mode)"],
        ["Databricks workspace.default", "Synthetic data catalog in Databricks (cloud development mode)"],
        ["Karpathy AutoResearch", "Andrej Karpathy, AutoResearch pattern for AI agent evaluation: scalar metrics, immutable eval harness, binary pass/fail, time-boxed cycles, keep-or-revert loop (2025)"],
    ],
    col_widths=[2.5, 4.0]
)

# 1.5 Document Overview
doc.add_heading("1.5 Document Overview", level=2)
add_para(
    "This document is organized into eight sections plus appendices. Section 1 introduces the "
    "Testing Agent's purpose and scope. Section 2 describes the overall product context, user "
    "characteristics, and constraints. Section 3 specifies the detailed functional and non-functional "
    "requirements (TR1\u2013TR9, 31 test requirements). Section 4 describes the test data architecture "
    "for both simulated and real modes. Section 5 defines the Quantitative Evaluation Framework "
    "based on Karpathy's AutoResearch methodology, including scalar metrics, binary pass/fail "
    "criteria, and the keep-or-revert evaluation loop. Section 6 provides the test case matrix "
    "mapping main SRS requirements to test requirements. Section 7 provides binary pass/fail "
    "criteria for each test requirement group. Section 8 defines quantitative acceptance criteria "
    "for the Testing Agent itself. Appendices provide the prompt-response pair template and "
    "simulated data summary."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
#  2. OVERALL DESCRIPTION (IEEE 830 Section 2)
# ══════════════════════════════════════════════════════════════════
doc.add_heading("2. Overall Description", level=1)

# 2.1 Product Perspective
doc.add_heading("2.1 Product Perspective", level=2)
add_para(
    "The Testing Agent is one of the specialized agents in FinIQ's agent pool, alongside the "
    "PES Agent, CI Agent, NL Query Agent, and others. Unlike those agents, which serve user-facing "
    "queries, the Testing Agent validates the system itself. It connects to the same data sources "
    "as FinIQ (Databricks, Azure Blob Storage, Azure AI Search) but compares outputs against known "
    "expected results rather than delivering them to end users."
)
add_para(
    "The Testing Agent operates within FinIQ's existing architecture. It submits requests through "
    "the same API endpoints that users and other agents use, then intercepts and evaluates the "
    "responses. This approach ensures that tests exercise the full production code path, not a "
    "separate test harness."
)
add_para(
    "In simulated mode, the Testing Agent targets synthetic data with pre-computed expected values. "
    "In real mode, it targets production data using user-provided prompt-response pairs as ground "
    "truth. This dual-mode design mirrors FinIQ's own dual-mode architecture (development vs. "
    "production data sources)."
)

# 2.2 Product Functions
doc.add_heading("2.2 Product Functions", level=2)
add_para("The Testing Agent provides the following core capabilities:")
add_bullet("Automated test execution against all FR groups (FR1\u2013FR8, 50 functional requirements)")
add_bullet("Dual-mode operation: simulated data (deterministic) and real data (user-validated)")
add_bullet("Prompt-response validation for real data \u2014 users provide prompts with expected results, the agent compares actual output")
add_bullet("Regression detection \u2014 compares current test outputs against stored baseline (golden dataset)")
add_bullet("Test coverage reporting \u2014 maps every test case to its corresponding main SRS requirement")
add_bullet("CI/CD integration \u2014 full regression suite runs automatically on every deployment to staging")
add_bullet("On-demand test execution \u2014 individual tests or suites can be triggered via the FinIQ job board")

# 2.3 User Characteristics
doc.add_heading("2.3 User Characteristics", level=2)
add_para("Three primary user types interact with the Testing Agent:")

add_para("QA Engineer", bold=True, space_after=2)
add_para(
    "Designs test cases, manages test suites, reviews test results, and maintains golden datasets. "
    "Has deep knowledge of FinIQ's functional requirements and expected behavior. Responsible for "
    "ensuring adequate test coverage across all FR groups."
)

add_para("Developer", bold=True, space_after=2)
add_para(
    "Runs tests during development to verify changes, investigates test failures, and uses test "
    "results to guide debugging. Interacts with the Testing Agent primarily through CI/CD pipeline "
    "feedback and on-demand test execution."
)

add_para("Product Owner", bold=True, space_after=2)
add_para(
    "Reviews coverage reports and test trend data to make release decisions. Uses test results as "
    "evidence that the platform meets acceptance criteria. Does not design or execute tests directly."
)

# 2.4 Constraints
doc.add_heading("2.4 Constraints", level=2)
add_bullet("Must not modify production data \u2014 all operations are read-only against Databricks and other data stores")
add_bullet("Must use read-only access credentials for all Databricks connections")
add_bullet("Test execution time budget: full regression suite must complete within 30 minutes")
add_bullet("Must support the same dual-mode toggle (simulated vs. real) as the main FinIQ application")
add_bullet("Must not interfere with concurrent user operations on the FinIQ platform")
add_bullet("Test reports must not include raw production data values \u2014 only pass/fail status and deviation metrics")

# 2.5 Assumptions
doc.add_heading("2.5 Assumptions", level=2)
add_bullet("Synthetic data (finiq_synthetic.db for local, workspace.default for Databricks) is available and stable")
add_bullet("Real data prompt-response pairs will be provided by the client/team \u2014 placeholder sections are included in this SRS")
add_bullet("FinIQ API endpoints are accessible for integration testing in both development and staging environments")
add_bullet("The FinIQ agent pool infrastructure is operational and can schedule Testing Agent jobs")
add_bullet("Golden dataset baselines will be established during the initial test suite setup and updated through a controlled review process")
add_bullet("CI/CD pipeline (GitHub Actions) is configured and accessible for webhook integration")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
#  3. SPECIFIC REQUIREMENTS (IEEE 830 Section 3)
# ══════════════════════════════════════════════════════════════════
doc.add_heading("3. Specific Requirements", level=1)

# 3.1 External Interface Requirements
doc.add_heading("3.1 External Interface Requirements", level=2)

doc.add_heading("3.1.1 FinIQ API Interface", level=3)
add_para(
    "The Testing Agent connects to all FinIQ API endpoints defined in the main SRS Section 3.1.4 "
    "(Communications Interfaces). It issues requests identical to those made by end users and other "
    "agents, then captures and evaluates the responses."
)

doc.add_heading("3.1.2 Databricks Interface", level=3)
add_para(
    "The Testing Agent connects to the same Databricks workspace as FinIQ. In simulated mode, it "
    "targets workspace.default (synthetic data). In real mode, it targets "
    "corporate_finance_analytics_dev.finsight_core_model_mvp3. All connections use read-only "
    "credentials."
)

doc.add_heading("3.1.3 Test Report Output", level=3)
add_para(
    "The Testing Agent produces test reports in three formats: JSON (machine-readable, for CI/CD "
    "integration), HTML (human-readable, for browser viewing), and PDF (archival, for stakeholder "
    "distribution). All formats contain identical data: pass/fail results, execution times, error "
    "details, coverage metrics, and regression indicators."
)

doc.add_heading("3.1.4 CI/CD Integration", level=3)
add_para(
    "The Testing Agent integrates with GitHub Actions via webhooks. On each deployment to the "
    "staging environment, the CI/CD pipeline triggers a full regression suite execution. Results "
    "are reported back to the pipeline as pass/fail status with a link to the full HTML report."
)

# 3.2 Functional Requirements
doc.add_heading("3.2 Functional Requirements", level=2)
add_para(
    "This section defines 31 test requirements organized into 9 groups (TR1\u2013TR9). Each "
    "requirement specifies a testable capability of the Testing Agent.",
    bold=True
)

# ── TR1: Test Data Management ──
doc.add_heading("3.2.1 TR1: Test Data Management", level=3)
add_para("4 requirements governing data source selection and golden dataset management.")

add_req_table(
    "TR1.1", "Simulated Data Mode",
    "The Testing Agent shall connect to synthetic Databricks data (workspace.default catalog) or "
    "local SQLite (finiq_synthetic.db) based on environment configuration. All test cases in simulated "
    "mode shall execute against known synthetic data with deterministic expected values that are "
    "pre-computed from the synthetic dataset.",
    priority="Critical", category="Functional"
)
add_req_table(
    "TR1.2", "Real Data Mode",
    "The Testing Agent shall connect to real Databricks "
    "(corporate_finance_analytics_dev.finsight_core_model_mvp3) when configured for real data mode. "
    "Test cases in real mode shall use user-provided prompt-response pairs as expected results. "
    "Expected results are provided as input, not computed by the agent.",
    priority="Critical", category="Functional"
)
add_req_table(
    "TR1.3", "Mode Toggle",
    "The Testing Agent shall provide a single configuration switch to toggle between simulated and "
    "real data modes. The same test framework, test runner, and reporting infrastructure shall be "
    "used in both modes \u2014 only the data source and expected result source differ.",
    priority="High", category="Functional"
)
add_req_table(
    "TR1.4", "Golden Dataset Management",
    "The Testing Agent shall maintain versioned golden datasets containing expected outputs for each "
    "test case. Golden datasets shall be updatable through a controlled review workflow when "
    "requirements change. Full version history shall be preserved to support regression analysis "
    "across golden dataset versions.",
    priority="High", category="Functional"
)

# ── TR2: Period End Summary Testing ──
doc.add_heading("3.2.2 TR2: Period End Summary Testing", level=3)
add_para("5 requirements validating PES generation correctness.")

add_req_table(
    "TR2.1", "PES Generation Validation",
    "For each of the 6 KPIs (Organic Growth, MAC Shape %, A&CP Shape %, CE Shape %, "
    "Controllable Overhead Shape %, NCFO), the Testing Agent shall verify that generated narratives "
    "contain correct numerical values from the source data. Validation shall cover all three output "
    "formats: Summary, What's Working Well (WWW), and What's Not Working Well (WNWW).",
    priority="Critical", category="Functional"
)
add_req_table(
    "TR2.2", "KPI Accuracy",
    "The Testing Agent shall validate that each KPI value matches the source data within a "
    "configurable tolerance threshold (default: 0.01%). The tolerance shall be adjustable per test "
    "case to accommodate KPIs with different precision characteristics.",
    priority="Critical", category="Functional"
)
add_req_table(
    "TR2.3", "Sub-Unit Ranking Validation",
    "The Testing Agent shall verify that RANK 1, TOP 3, and BOTTOM 3 rankings for each KPI are "
    "correct by independently calculating rankings from the source data and comparing against "
    "the PES-generated rankings. Ties shall be handled according to the same ranking logic used "
    "by the PES Agent.",
    priority="High", category="Functional"
)
add_req_table(
    "TR2.4", "Trend Tagline Validation",
    "The Testing Agent shall verify that trend taglines accurately reflect multi-period data trends. "
    "Validation shall confirm that directional indicators (improving, declining, stable) match the "
    "underlying data trajectory across the relevant time periods.",
    priority="High", category="Functional"
)
add_req_table(
    "TR2.5", "PES Prompt-Response Testing (Real Mode)",
    "In real data mode, the Testing Agent shall execute user-provided prompts (e.g., 'Generate "
    "Petcare Period End Summary for P6 2025') and compare the actual output against user-provided "
    "expected results. Comparison shall account for the configured tolerance threshold and format "
    "variations. [PLACEHOLDER: Expected results to be provided by Rajiv]",
    priority="High", category="Functional"
)

# ── TR3: Budget Variance Testing ──
doc.add_heading("3.2.3 TR3: Budget Variance Testing", level=3)
add_para("3 requirements validating budget variance calculations from finiq_financial_replan data.")

add_req_table(
    "TR3.1", "Variance Calculation Accuracy",
    "The Testing Agent shall verify that Actual vs. Replan variance calculations match independent "
    "computation from finiq_financial_replan data within the configured tolerance threshold. "
    "Validation shall cover both absolute and percentage variance figures.",
    priority="Critical", category="Functional"
)
add_req_table(
    "TR3.2", "Three-Way Comparison",
    "The Testing Agent shall validate that the Actual vs. Replan vs. Forecast comparison renders "
    "correctly when all three data sources are available. It shall also verify graceful handling "
    "when one or more data sources are missing (e.g., forecast not yet available for a period).",
    priority="High", category="Functional"
)
add_req_table(
    "TR3.3", "Currency Handling",
    "The Testing Agent shall verify that USD and Local currency values are correctly displayed and "
    "converted. Tests shall validate that currency indicators are present, that conversion rates "
    "are applied correctly, and that the user's configured currency preference is respected.",
    priority="High", category="Functional"
)

# ── TR4: Natural Language Query Testing ──
doc.add_heading("3.2.4 TR4: Natural Language Query Testing", level=3)
add_para("4 requirements validating the NL query interface.")

add_req_table(
    "TR4.1", "Query Accuracy",
    "The Testing Agent shall execute a suite of test queries covering all supported query types "
    "(financial lookups, trend questions, comparison queries, what-if scenarios) and validate "
    "factual correctness of responses against source data. Each query shall have a defined expected "
    "answer or answer range.",
    priority="Critical", category="Functional"
)
add_req_table(
    "TR4.2", "Source Attribution",
    "The Testing Agent shall verify that every answer includes correct source references linking "
    "back to the specific data tables, document sections, or calculations that produced the result. "
    "Attribution shall be validated against the known data lineage for each test query.",
    priority="High", category="Functional"
)
add_req_table(
    "TR4.3", "Multi-Turn Context",
    "The Testing Agent shall validate that conversation context is maintained across follow-up "
    "queries within the same session. Test sequences shall include pronoun resolution (e.g., "
    "'What about last year?' following a specific-period query) and topic continuation.",
    priority="High", category="Functional"
)
add_req_table(
    "TR4.4", "Intent Classification",
    "The Testing Agent shall verify that queries are routed to the correct processing pipeline "
    "(PES generation, CI lookup, financial query, job submission, etc.). Misclassification shall "
    "be detected by comparing the pipeline used against the expected pipeline for each test query.",
    priority="Medium", category="Functional"
)

# ── TR5: Competitive Intelligence Testing ──
doc.add_heading("3.2.5 TR5: Competitive Intelligence Testing", level=3)
add_para("3 requirements validating competitive intelligence capabilities.")

add_req_table(
    "TR5.1", "Themed Summary Validation",
    "The Testing Agent shall verify that competitor themed summaries (Organic Growth, Margins, "
    "Projections, Consumer Trends, Product Launches, Product Summary, Miscellaneous) contain "
    "accurate data with correct source references back to the ingested competitor documents.",
    priority="High", category="Functional"
)
add_req_table(
    "TR5.2", "P2P Benchmarking Accuracy",
    "The Testing Agent shall validate that quantitative peer-to-peer comparison tables (OG%, Price, "
    "Volume, Mix, Adj Core Operating Profit %) match the values in the source competitor documents. "
    "Both quarterly and YTD views shall be validated for each peer group.",
    priority="High", category="Functional"
)
add_req_table(
    "TR5.3", "Internal-External Cross-Reference",
    "The Testing Agent shall verify that queries combining internal financial data with competitor "
    "data return correct results from both sources. Cross-reference accuracy shall be validated "
    "by independently querying each source and comparing the combined result.",
    priority="High", category="Functional"
)

# ── TR6: Data Ingestion Testing ──
doc.add_heading("3.2.6 TR6: Data Ingestion Testing", level=3)
add_para("3 requirements validating data connectivity and ingestion.")

add_req_table(
    "TR6.1", "Databricks Connectivity",
    "The Testing Agent shall verify connectivity to all 20 finiq_ objects (17 tables + 3 views) "
    "in the configured Databricks catalog. Validation shall include schema match (column names and "
    "types match expected), row count verification (within expected ranges), and query execution "
    "against each object.",
    priority="Critical", category="Functional"
)
add_req_table(
    "TR6.2", "Excel Fallback",
    "The Testing Agent shall verify that the system falls back to Excel ingestion within 30 seconds "
    "of a Databricks connection failure. The test shall simulate a Databricks timeout and confirm "
    "that the fallback path produces equivalent results from the Excel data source.",
    priority="High", category="Functional"
)
add_req_table(
    "TR6.3", "Data Freshness",
    "The Testing Agent shall verify that stale data alerts fire correctly when expected data is "
    "missing or outdated. Tests shall simulate scenarios where the latest period's data has not "
    "been loaded and confirm that appropriate warnings are surfaced to users.",
    priority="Medium", category="Functional"
)

# ── TR7: Job Board Testing ──
doc.add_heading("3.2.7 TR7: Job Board Testing", level=3)
add_para("3 requirements validating job board operations.")

add_req_table(
    "TR7.1", "Job Lifecycle",
    "The Testing Agent shall submit test jobs and verify they progress through all lifecycle "
    "stages: Submitted \u2192 Queued \u2192 Assigned \u2192 In Progress \u2192 Completed (or Failed). "
    "Each state transition shall be verified against expected timing and status codes.",
    priority="High", category="Functional"
)
add_req_table(
    "TR7.2", "SLA Compliance",
    "The Testing Agent shall verify that jobs complete within SLA targets for each priority level "
    "(Critical, High, Medium, Low). Tests shall include jobs at each priority level and measure "
    "actual completion time against the defined SLA thresholds.",
    priority="High", category="Functional"
)
add_req_table(
    "TR7.3", "Scheduled Jobs",
    "The Testing Agent shall verify that recurring job scheduling works correctly. Tests shall "
    "create scheduled jobs with defined recurrence patterns and confirm that jobs are triggered "
    "at the expected times with the expected parameters.",
    priority="Medium", category="Functional"
)

# ── TR8: Regression Testing ──
doc.add_heading("3.2.8 TR8: Regression Testing", level=3)
add_para("3 requirements for regression detection and reporting.")

add_req_table(
    "TR8.1", "Baseline Comparison",
    "The Testing Agent shall compare current test outputs against stored baseline outputs (golden "
    "dataset) for every test case. Any difference exceeding the configured tolerance threshold "
    "shall be flagged as a potential regression with details of the expected vs. actual values.",
    priority="Critical", category="Functional"
)
add_req_table(
    "TR8.2", "Automated Regression Suite",
    "A full regression suite shall run automatically on every deployment to the staging environment. "
    "The suite shall execute all test cases in both simulated mode (full suite) and real mode "
    "(all available prompt-response pairs). Deployment to production shall be blocked if critical "
    "test failures are detected.",
    priority="Critical", category="Functional"
)
add_req_table(
    "TR8.3", "Regression Report",
    "The Testing Agent shall generate a detailed diff report showing what changed between the "
    "baseline and current output. The report shall include: changed values with before/after "
    "comparison, newly failing test cases, newly passing test cases, and overall regression "
    "risk assessment.",
    priority="High", category="Functional"
)

# ── TR9: Test Reporting ──
doc.add_heading("3.2.9 TR9: Test Reporting", level=3)
add_para("3 requirements for test result reporting and analytics.")

add_req_table(
    "TR9.1", "Test Execution Report",
    "The Testing Agent shall generate a comprehensive test execution report showing pass/fail "
    "status for each test case, execution time, error details for failures, and summary statistics "
    "(total passed, total failed, total skipped, execution duration). Reports shall be available "
    "in JSON, HTML, and PDF formats.",
    priority="High", category="Functional"
)
add_req_table(
    "TR9.2", "Coverage Report",
    "The Testing Agent shall generate a coverage report mapping every test case to its corresponding "
    "main SRS functional requirement(s). The report shall show coverage percentage per FR group "
    "(FR1\u2013FR8), identify any FRs without test coverage, and provide an overall coverage score.",
    priority="High", category="Functional"
)
add_req_table(
    "TR9.3", "Trend Report",
    "The Testing Agent shall track test results over time and generate trend reports showing "
    "improvement or regression trends across test runs. The report shall include pass rate history, "
    "flaky test identification (tests that intermittently pass/fail), and execution time trends.",
    priority="Medium", category="Functional"
)

# 3.3 Performance Requirements
doc.add_heading("3.3 Performance Requirements", level=2)
add_table(
    ["Metric", "Target", "Notes"],
    [
        ["Full simulated test suite", "< 15 minutes", "All TR1\u2013TR9 tests against synthetic data"],
        ["Full real data test suite", "< 30 minutes", "All available prompt-response pairs against production data"],
        ["Individual test case", "< 30 seconds", "Any single test case in isolation"],
        ["Report generation", "< 60 seconds", "Any format (JSON, HTML, PDF)"],
        ["Regression diff computation", "< 120 seconds", "Full golden dataset comparison"],
    ],
    col_widths=[2.0, 1.5, 3.0]
)

# 3.4 Design Constraints
doc.add_heading("3.4 Design Constraints", level=2)
add_bullet("Must use the same technology stack as FinIQ: Python, LangChain, Azure OpenAI Foundry")
add_bullet("Must run in both local (development) and cloud (CI/CD) environments without modification")
add_bullet("Must not require manual intervention during test execution \u2014 fully automated from trigger to report")
add_bullet("Must use the same Databricks connection libraries and authentication mechanisms as the main FinIQ platform")
add_bullet("Test case definitions shall be stored in JSON format for portability and version control")
add_bullet("Must integrate with the existing FinIQ agent pool infrastructure (LangGraph agent runtime)")

# 3.5 Software System Attributes
doc.add_heading("3.5 Software System Attributes", level=2)

add_para("Reliability", bold=True, space_after=2)
add_para(
    "Tests must be deterministic \u2014 the same inputs must produce the same pass/fail results on "
    "every execution. Non-deterministic behavior (e.g., from LLM temperature variations) shall be "
    "accounted for through tolerance thresholds and semantic comparison rather than exact string "
    "matching."
)

add_para("Maintainability", bold=True, space_after=2)
add_para(
    "Adding a new test case shall require no more than 15 minutes of effort. Test cases are defined "
    "as JSON configurations that are loaded by the test runner \u2014 no code changes required for "
    "standard test additions."
)

add_para("Portability", bold=True, space_after=2)
add_para(
    "The Testing Agent shall run on developer machines (using SQLite/finiq_synthetic.db) and in "
    "the cloud (using Databricks) with no code changes \u2014 only configuration differences. The "
    "same test definitions shall work in both environments."
)

add_para("Traceability", bold=True, space_after=2)
add_para(
    "Every test case shall be traceable to one or more main SRS functional requirements. "
    "Every test execution shall be logged with a unique run ID, timestamp, environment details, "
    "and complete results for audit purposes."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
#  4. TEST DATA ARCHITECTURE
# ══════════════════════════════════════════════════════════════════
doc.add_heading("4. Test Data Architecture", level=1)

# 4.1 Simulated Data
doc.add_heading("4.1 Simulated Data", level=2)
add_para(
    "Simulated mode uses a synthetic dataset designed to exercise all FinIQ functional requirements "
    "with known, deterministic values. The synthetic data is generated with a fixed random seed "
    "(seed 42) to ensure reproducibility across environments."
)

add_table(
    ["Property", "Value"],
    [
        ["Local source", "finiq_synthetic.db (SQLite)"],
        ["Cloud source", "Databricks workspace.default catalog"],
        ["Table count", "17 tables + 3 views (mirrors production schema)"],
        ["Total rows", "165,000+ across all tables"],
        ["Random seed", "42 (fixed for reproducibility)"],
        ["Expected results", "Pre-computed from synthetic data, stored in golden dataset"],
        ["Org units", "Synthetic hierarchy: Mars Inc > 3 GBUs > 9 Divisions > 27 Regions"],
        ["Time periods", "P1\u2013P13 for fiscal years 2024 and 2025"],
    ],
    col_widths=[2.0, 4.5]
)

# 4.2 Real Data (Placeholder)
doc.add_heading("4.2 Real Data (Placeholder)", level=2)
add_para(
    "Real data mode targets Mars's production Databricks environment. Expected results are not "
    "computed by the Testing Agent \u2014 they are provided by the client/team as prompt-response pairs.",
    bold=True
)

add_table(
    ["Property", "Value"],
    [
        ["Source", "corporate_finance_analytics_dev.finsight_core_model_mvp3"],
        ["Access", "Read-only credentials"],
        ["Expected results", "User-provided prompt-response pairs (JSON format)"],
        ["Tolerance", "Configurable per test case (default: 0.01%)"],
    ],
    col_widths=[2.0, 4.5]
)

add_para("Prompt-Response Pair Format:", bold=True, space_after=4)
add_para(
    "Each real data test case is defined as a JSON object with the following structure:",
    space_after=4
)

# Show JSON format as a code-like block
json_example = (
    '{\n'
    '    "test_id": "PES_001",\n'
    '    "prompt": "Generate Petcare Period End Summary for P6 2025",\n'
    '    "expected_output": {\n'
    '        "format": "Summary",\n'
    '        "kpis": ["Organic Growth", "MAC Shape %", ...],\n'
    '        "values": { ... }\n'
    '    },\n'
    '    "tolerance": 0.01,\n'
    '    "mode": "real",\n'
    '    "category": "PES"\n'
    '}'
)
p = doc.add_paragraph()
run = p.add_run(json_example)
run.font.name = 'Consolas'
run.font.size = Pt(9)
run.font.color.rgb = DARK_GRAY
p.paragraph_format.space_after = Pt(8)

add_para(
    "[PLACEHOLDER: Rajiv to provide example prompts and expected results for real data validation. "
    "This section will be updated when prompt-response pairs are finalized.]",
    italic=True, color=RGBColor(0x99, 0x66, 0x00)
)

# 4.3 Golden Dataset Management
doc.add_heading("4.3 Golden Dataset Management", level=2)
add_para(
    "Golden datasets serve as the authoritative expected outputs for regression testing. They are "
    "versioned and managed through a controlled update process."
)
add_bullet("Versioned storage: Each golden dataset version is stored with a timestamp and change description")
add_bullet("Diff tool: Built-in comparison tool highlights differences between golden dataset versions")
add_bullet("Update workflow: A new golden dataset version requires review and explicit approval before it replaces the current baseline")
add_bullet("Rollback: Previous golden dataset versions can be restored if a new version is found to be incorrect")
add_bullet("Format: Golden datasets are stored as JSON files alongside the test case definitions")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
#  5. QUANTITATIVE EVALUATION FRAMEWORK (Karpathy methodology)
# ══════════════════════════════════════════════════════════════════
doc.add_heading("5. Quantitative Evaluation Framework", level=1)
add_para(
    "This section defines the quantitative evaluation methodology for the FinIQ Testing Agent, "
    "adapted from Andrej Karpathy's AutoResearch pattern for AI agent evaluation. The framework "
    "ensures that every capability is measured by a single scalar metric, scored by an immutable "
    "eval harness, and evaluated through binary pass/fail criteria with no subjective judgment."
)

# 5.1 Core Principles
doc.add_heading("5.1 Core Principles", level=2)

add_para("Single Scalar Metric per Capability", bold=True, space_after=2)
add_para(
    "Each FinIQ capability is reduced to one aggregate pass rate: the percentage of test cases "
    "that pass across all scenarios for that capability. This scalar collapses multiple dimensions "
    "(accuracy, completeness, timeliness) into one number that can be tracked over time and compared "
    "across builds."
)

add_para("Immutable Eval Harness", bold=True, space_after=2)
add_para(
    "The scoring infrastructure is separated from the agent under test. The FinIQ application "
    "cannot modify, influence, or access its own scoring logic. The eval harness runs independently, "
    "reads outputs produced by FinIQ, and compares them against pre-computed ground truth. This "
    "separation prevents gaming or circular validation."
)

add_para("Binary Pass/Fail Criteria", bold=True, space_after=2)
add_para(
    "Every test case resolves to a binary outcome: pass (1) or fail (0). There is no partial "
    "credit, no subjective scoring, and no human judgment during execution. Each test requirement "
    "group defines 3\u20136 specific binary criteria. A test case passes only when all applicable "
    "criteria are satisfied."
)

add_para("Time-Boxed Test Cycles", bold=True, space_after=2)
add_para(
    "Each test cycle operates within a fixed compute budget. The full simulated suite must complete "
    "within 15 minutes; the full real-data suite within 30 minutes. If a test case exceeds its "
    "individual time budget (30 seconds), it is marked as failed regardless of output correctness. "
    "This prevents runaway execution and ensures predictable CI/CD cycle times."
)

add_para("Keep-or-Revert Evaluation Loop", bold=True, space_after=2)
add_para(
    "After every code change, the full test suite runs automatically. If the aggregate score "
    "improves or remains stable (within 2% tolerance), the change is committed. If the score drops "
    "by more than 2%, the change is reverted. This loop ensures that the codebase only moves forward "
    "in measurable quality. The 2% threshold is configurable per deployment environment."
)

add_para("Ground Truth from Synthetic Data", bold=True, space_after=2)
add_para(
    "In simulated mode, all expected values are pre-computed deterministically from the synthetic "
    "dataset (finiq_synthetic.db / Databricks workspace.default). The synthetic data is generated "
    "with a fixed random seed (seed 42), making every expected value reproducible across environments. "
    "This eliminates ambiguity about what constitutes a correct answer."
)

# 5.2 Scalar Metrics Summary Table
doc.add_heading("5.2 Scalar Metrics Summary", level=2)
add_para(
    "The following table maps each FinIQ capability to its single scalar metric, target threshold, "
    "and tolerance. These metrics form the quantitative scorecard for every test run."
)

add_table(
    ["Capability", "Scalar Metric", "Target", "Tolerance", "Source View / Table"],
    [
        ["PES Reports", "KPI exact-match %", "\u226595%", "\u00b10.01%", "finiq_vw_pl_entity, finiq_vw_pl_brand_product, finiq_vw_ncfo_entity"],
        ["NL Queries", "Query accuracy rate", "\u226585%", "Exact match", "All finiq_ tables/views"],
        ["Budget Variance", "Value match %", "\u226595%", "\u00b10.01%", "finiq_financial_replan, finiq_financial_replan_cons"],
        ["CI Summaries", "Factual accuracy", "\u226580%", "N/A", "Ingested competitor documents"],
        ["Job Board", "SLA compliance %", "\u226590%", "N/A", "Job lifecycle timestamps"],
        ["Data Ingestion", "Checksum match %", "100%", "Exact", "All 20 finiq_ objects"],
        ["Regression", "Score stability", "\u22642% drop", "Per build", "Golden dataset comparison"],
    ],
    col_widths=[1.2, 1.3, 0.7, 0.8, 2.5]
)

# 5.3 Quantitative Thresholds per TR Group
doc.add_heading("5.3 Quantitative Thresholds per Test Requirement Group", level=2)
add_para(
    "Each test requirement group has specific numeric thresholds that convert qualitative expectations "
    "into measurable pass/fail criteria."
)

add_para("TR2 \u2014 PES Testing", bold=True, space_after=2)
add_bullet("KPI values must match expected values within \u00b10.01% tolerance")
add_bullet("Expected values are pre-computed from finiq_vw_pl_entity, finiq_vw_pl_brand_product, and finiq_vw_ncfo_entity")
add_bullet("Rankings (RANK 1, TOP 3, BOTTOM 3) must match exactly \u2014 no tolerance on ordinal positions")
add_bullet("All 6 KPIs must be present in output; missing KPI = automatic failure")

add_para("TR3 \u2014 Budget Variance Testing", bold=True, space_after=2)
add_bullet("Actual vs. replan values must match finiq_financial_replan within \u00b10.01%")
add_bullet("Variance percentage must be computed correctly: (Actual \u2212 Replan) / Replan \u00d7 100")
add_bullet("Currency values must match within \u00b10.01 of source (USD and local currency)")

add_para("TR4 \u2014 NL Query Testing", bold=True, space_after=2)
add_bullet("Generated SQL must execute without error (binary: executes = pass, error = fail)")
add_bullet("Result row count must match expected row count exactly")
add_bullet("Numeric values in results must match within \u00b10.01% tolerance")
add_bullet("Response must be returned within the SLA time budget (30 seconds per query)")

add_para("TR5 \u2014 CI Testing", bold=True, space_after=2)
add_bullet("Factual accuracy score = verifiable claims / total claims (target \u226580%)")
add_bullet("Every factual claim must be traceable to a specific source document section")
add_bullet("P2P quantitative values must match source documents exactly")

add_para("TR6 \u2014 Data Ingestion Testing", bold=True, space_after=2)
add_bullet("Row count for each of the 20 finiq_ objects must match known synthetic data counts")
add_bullet("Schema checksum (column names + types) must match expected schema exactly")
add_bullet("Sample data spot-checks: 10 random rows per table verified against source")

add_para("TR7 \u2014 Job Board Testing", bold=True, space_after=2)
add_bullet("Percentage of jobs completed within SLA threshold (target \u226590%)")
add_bullet("All lifecycle state transitions must occur in correct order")
add_bullet("Scheduling accuracy: jobs trigger within \u00b160 seconds of scheduled time")

add_para("TR8 \u2014 Regression Testing", bold=True, space_after=2)
add_bullet("Aggregate score must not regress more than 2% between builds")
add_bullet("No individual capability may drop more than 5% between builds")
add_bullet("Regression detection must complete within 5 minutes of code change")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
#  6. TEST CASE MATRIX (renumbered from 5)
# ══════════════════════════════════════════════════════════════════
doc.add_heading("6. Test Case Matrix", level=1)
add_para(
    "The following matrix maps each main SRS functional requirement to the corresponding test "
    "requirements and test cases defined in this SRS. This mapping ensures complete traceability "
    "between the base FinIQ requirements and the Testing Agent's validation coverage."
)

add_table(
    ["Main SRS Req", "Description", "Test Req", "Test Cases", "Mode"],
    [
        ["FR1.1", "Excel Data Ingestion", "TR6.1, TR6.2", "Connectivity, schema, fallback", "Both"],
        ["FR1.2", "Competitor PDF Ingestion", "TR5.1", "Document parsing, summary accuracy", "Both"],
        ["FR1.3", "Third-Party Connectors", "TR6.1", "Connector availability, data freshness", "Both"],
        ["FR1.4", "Data Lineage", "TR9.2", "Lineage tracking in test reports", "Both"],
        ["FR1.5", "Scheduling", "TR7.3", "Scheduled ingestion triggers", "Both"],
        ["FR1.6", "Databricks Connectivity", "TR6.1", "All 20 finiq_ objects, row counts", "Both"],
        ["FR2.1", "PES Generation", "TR2.1\u2013TR2.5", "KPI accuracy, rankings, formats", "Both"],
        ["FR2.2", "Configurable KPIs", "TR2.2", "KPI tolerance, dynamic KPI sets", "Both"],
        ["FR2.3", "Rankings", "TR2.3", "RANK 1, TOP 3, BOTTOM 3 validation", "Both"],
        ["FR2.4", "Interactive Tables", "TR2.1", "Table rendering, data accuracy", "Both"],
        ["FR2.5", "Custom Report Builder", "TR4.1", "Query-driven report generation", "Both"],
        ["FR2.6", "Export", "TR9.1", "JSON, HTML, PDF export correctness", "Both"],
        ["FR2.7", "Budget Variance", "TR3.1\u2013TR3.3", "Variance calcs, three-way, currency", "Both"],
        ["FR3.1", "Themed Summaries", "TR5.1", "7 theme types, source attribution", "Both"],
        ["FR3.2", "P2P Benchmarking", "TR5.2", "Quantitative tables, peer groups", "Both"],
        ["FR3.3", "Cross-Reference", "TR5.3", "Internal + external data queries", "Both"],
        ["FR3.4", "Monitoring", "TR5.1", "New document detection, alerts", "Both"],
        ["FR4.1", "Query Engine", "TR4.1\u2013TR4.4", "Accuracy, attribution, multi-turn", "Both"],
        ["FR4.2", "Multi-Turn", "TR4.3", "Context maintenance, pronoun resolution", "Both"],
        ["FR4.3", "Intent Classification", "TR4.4", "Pipeline routing accuracy", "Both"],
        ["FR4.4", "Source Attribution", "TR4.2", "Citation correctness, link validity", "Both"],
        ["FR5.1", "Job Submission", "TR7.1", "Lifecycle from submit to complete", "Both"],
        ["FR5.2", "Agent Pool", "TR7.1", "Agent assignment, capacity", "Both"],
        ["FR5.3", "SLA Routing", "TR7.2", "Priority-based SLA compliance", "Both"],
        ["FR5.4", "Lifecycle", "TR7.1", "All state transitions verified", "Both"],
        ["FR5.5", "Dashboard", "TR7.1", "Job status visibility", "Both"],
        ["FR5.6", "Scheduling", "TR7.3", "Recurring job execution", "Both"],
        ["FR5.7", "Review", "TR7.1", "Job result review workflow", "Both"],
        ["FR6.1", "Forecasting API", "TR4.1", "Forecast data in query responses", "Real"],
        ["FR6.2", "Marketing API", "TR4.1", "Marketing data in query responses", "Real"],
        ["FR6.3", "Recommendation Engine", "TR4.1", "Recommendation accuracy", "Real"],
        ["FR6.4", "External Gateway", "TR6.1", "Third-party data connectivity", "Both"],
        ["FR6.5", "Databricks Direct", "TR6.1", "Direct Databricks query execution", "Both"],
        ["FR7.1", "Templates", "TR2.1", "Template-driven output format", "Both"],
        ["FR7.2", "Org Hierarchy", "TR2.3", "Hierarchy-based rankings", "Both"],
        ["FR7.3", "Peer Groups", "TR5.2", "Peer group configuration", "Both"],
        ["FR7.4", "Prompt Management", "TR4.1", "Prompt registry validation", "Both"],
        ["FR7.5", "RBAC", "TR7.1", "Role-based access in test context", "Both"],
        ["FR7.6", "Databricks Config", "TR6.1", "Connection configuration", "Both"],
        ["FR8.1\u2013FR8.11", "Dynamic UI (11 reqs)", "TR4.1, TR9.1", "UI-driven query and report output", "Both"],
    ],
    col_widths=[0.8, 1.5, 1.2, 1.8, 0.6]
)

add_para(
    "Total coverage: All 50 main SRS functional requirements are mapped to at least one test "
    "requirement. The Testing Agent's 31 test requirements (TR1.1\u2013TR9.3) collectively provide "
    "full functional coverage.",
    bold=True
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
#  7. BINARY PASS/FAIL CRITERIA PER TR GROUP
# ══════════════════════════════════════════════════════════════════
doc.add_heading("7. Binary Pass/Fail Criteria", level=1)
add_para(
    "Following Karpathy's recommendation, each test requirement group defines 3\u20136 binary pass/fail "
    "criteria. Each criterion resolves to exactly 0 (fail) or 1 (pass) with no partial credit. A "
    "test case passes only when all applicable criteria for its group are satisfied."
)

doc.add_heading("7.1 TR2: PES Testing Criteria", level=2)
add_table(
    ["ID", "Binary Criterion", "Pass Condition"],
    [
        ["PES-C1", "KPI value accuracy", "Each KPI numeric value matches expected within \u00b10.01%"],
        ["PES-C2", "KPI completeness", "All 6 KPIs present in output (missing = fail)"],
        ["PES-C3", "Ranking correctness", "RANK 1, TOP 3, BOTTOM 3 match expected order exactly"],
        ["PES-C4", "Format compliance", "Output contains all required sections (Summary / WWW / WNWW)"],
        ["PES-C5", "Trend direction", "Trend tagline direction (up/down/stable) matches data trajectory"],
    ],
    col_widths=[0.8, 2.0, 3.7]
)

doc.add_heading("7.2 TR3: Budget Variance Criteria", level=2)
add_table(
    ["ID", "Binary Criterion", "Pass Condition"],
    [
        ["BV-C1", "Actual value match", "Actual amount matches finiq_financial_replan within \u00b10.01%"],
        ["BV-C2", "Replan value match", "Replan amount matches finiq_financial_replan within \u00b10.01%"],
        ["BV-C3", "Variance computation", "Variance = (Actual \u2212 Replan) / Replan \u00d7 100, within \u00b10.01%"],
        ["BV-C4", "Currency correctness", "Correct currency indicator displayed; conversion rate applied accurately"],
        ["BV-C5", "Missing data handling", "Graceful handling when forecast unavailable (no error, clear message)"],
    ],
    col_widths=[0.8, 2.0, 3.7]
)

doc.add_heading("7.3 TR4: NL Query Criteria", level=2)
add_table(
    ["ID", "Binary Criterion", "Pass Condition"],
    [
        ["NL-C1", "SQL execution", "Generated SQL executes without error on target database"],
        ["NL-C2", "Row count match", "Result set row count matches expected count exactly"],
        ["NL-C3", "Value accuracy", "All numeric values in result match expected within \u00b10.01%"],
        ["NL-C4", "Time budget", "Response returned within 30-second SLA"],
        ["NL-C5", "Source attribution", "At least one valid source reference present in response"],
        ["NL-C6", "Intent routing", "Query routed to correct processing pipeline"],
    ],
    col_widths=[0.8, 2.0, 3.7]
)

doc.add_heading("7.4 TR5: CI Testing Criteria", level=2)
add_table(
    ["ID", "Binary Criterion", "Pass Condition"],
    [
        ["CI-C1", "Factual accuracy", "Verifiable claims / total claims \u226580%"],
        ["CI-C2", "Source traceability", "Every factual claim links to a specific document section"],
        ["CI-C3", "P2P value match", "Quantitative P2P table values match source documents exactly"],
        ["CI-C4", "Theme completeness", "All 7 required themes present in summary output"],
    ],
    col_widths=[0.8, 2.0, 3.7]
)

doc.add_heading("7.5 TR6: Data Ingestion Criteria", level=2)
add_table(
    ["ID", "Binary Criterion", "Pass Condition"],
    [
        ["DI-C1", "Object accessibility", "All 20 finiq_ objects queryable without error"],
        ["DI-C2", "Row count match", "Row count per table matches known synthetic data count"],
        ["DI-C3", "Schema checksum", "Column names and types match expected schema exactly"],
        ["DI-C4", "Spot-check accuracy", "10 random rows per table match source values"],
    ],
    col_widths=[0.8, 2.0, 3.7]
)

doc.add_heading("7.6 TR7: Job Board Criteria", level=2)
add_table(
    ["ID", "Binary Criterion", "Pass Condition"],
    [
        ["JB-C1", "Lifecycle completion", "Job progresses through all states in correct order"],
        ["JB-C2", "SLA compliance", "Job completes within SLA threshold for its priority level"],
        ["JB-C3", "Schedule accuracy", "Scheduled job triggers within \u00b160 seconds of target time"],
    ],
    col_widths=[0.8, 2.0, 3.7]
)

doc.add_heading("7.7 TR8: Regression Criteria", level=2)
add_table(
    ["ID", "Binary Criterion", "Pass Condition"],
    [
        ["RG-C1", "Aggregate stability", "Overall pass rate does not drop more than 2% vs. baseline"],
        ["RG-C2", "Capability floor", "No single capability drops more than 5% vs. baseline"],
        ["RG-C3", "Detection speed", "Regression detected and reported within 5 minutes of code change"],
        ["RG-C4", "Zero critical failures", "No data integrity or security test failures (zero tolerance)"],
    ],
    col_widths=[0.8, 2.0, 3.7]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
#  8. ACCEPTANCE CRITERIA (renumbered from 6, now quantitative)
# ══════════════════════════════════════════════════════════════════
doc.add_heading("8. Acceptance Criteria", level=1)
add_para(
    "The following acceptance criteria define the quantitative conditions under which the Testing "
    "Agent itself is considered complete and ready for deployment. All thresholds are derived from "
    "the Quantitative Evaluation Framework (Section 5)."
)

add_table(
    ["ID", "Criterion", "Quantitative Target", "Verification Method"],
    [
        ["AC-1", "Overall pass rate across all test cases", "\u226585%", "Execute full suite; compute aggregate pass rate"],
        ["AC-2", "No single capability below floor", "\u226570% per capability", "Check per-capability scalar metrics from Section 5.2"],
        ["AC-3", "Zero critical test failures", "0 failures in data integrity and security tests", "Execute critical test subset; verify zero failures"],
        ["AC-4", "Regression detection speed", "Within 5 minutes of code change", "Introduce known change; measure time to detection"],
        ["AC-5", "Coverage of main SRS FRs", "100% (all 50 FRs mapped)", "Generate coverage report; verify no unmapped FRs"],
        ["AC-6", "Real data mode executes all prompt-response pairs", "100% of provided pairs processed", "Load prompt-response JSON; execute in real mode"],
        ["AC-7", "Test reports in all three formats", "JSON, HTML, PDF all valid", "Generate each format; verify structure and content"],
        ["AC-8", "CI/CD integration operational", "Auto-trigger on staging deploy", "Deploy to staging; verify webhook fires and suite runs"],
        ["AC-9", "Individual test case time budget", "\u226430 seconds each", "Execute each test case; measure wall-clock time"],
        ["AC-10", "Full simulated suite time budget", "\u226415 minutes", "Execute full simulated suite; measure total duration"],
        ["AC-11", "Full real data suite time budget", "\u226430 minutes", "Execute full real suite; measure total duration"],
        ["AC-12", "Keep-or-revert loop functional", "Score drop >2% triggers revert", "Introduce regressive change; verify revert occurs"],
        ["AC-13", "Adding test case requires no code changes", "JSON-only addition", "Add test via JSON; execute without code modifications"],
        ["AC-14", "Mode toggle without restart", "Switch in <5 seconds", "Toggle mode; verify next run uses correct data source"],
        ["AC-15", "Golden dataset versioning", "\u22653 versions retrievable", "Create 3+ versions; verify all diffable"],
    ],
    col_widths=[0.5, 2.0, 1.7, 2.3]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
#  APPENDIX A: PROMPT-RESPONSE PAIR TEMPLATE
# ══════════════════════════════════════════════════════════════════
doc.add_heading("Appendix A: Prompt-Response Pair Template", level=1)
add_para(
    "This appendix provides the JSON template for submitting test cases with expected results "
    "in real data mode. Each prompt-response pair defines a single test case.",
    space_after=8
)

add_para("[PLACEHOLDER \u2014 To be populated by Rajiv with actual prompt-response pairs]",
         italic=True, bold=True, color=RGBColor(0x99, 0x66, 0x00))

add_para("Template Structure:", bold=True, space_after=4)

template_json = (
    '[\n'
    '  {\n'
    '    "test_id": "PES_001",\n'
    '    "category": "PES",\n'
    '    "description": "Petcare Period End Summary for P6 2025",\n'
    '    "prompt": "Generate Petcare Period End Summary for P6 2025",\n'
    '    "expected_output": {\n'
    '      "format": "Summary",\n'
    '      "unit": "Petcare",\n'
    '      "period": "P6",\n'
    '      "year": "2025",\n'
    '      "kpis": {\n'
    '        "Organic Growth": { "value": "X.X%", "direction": "up" },\n'
    '        "MAC Shape %": { "value": "X.X%", "direction": "down" },\n'
    '        "...": "..."\n'
    '      },\n'
    '      "rankings": {\n'
    '        "RANK_1": "Sub-Unit A",\n'
    '        "TOP_3": ["Sub-Unit A", "Sub-Unit B", "Sub-Unit C"],\n'
    '        "BOTTOM_3": ["Sub-Unit X", "Sub-Unit Y", "Sub-Unit Z"]\n'
    '      }\n'
    '    },\n'
    '    "tolerance": 0.01,\n'
    '    "mode": "real",\n'
    '    "tags": ["PES", "Petcare", "P6", "2025"]\n'
    '  },\n'
    '  {\n'
    '    "test_id": "NLQ_001",\n'
    '    "category": "NL Query",\n'
    '    "description": "Simple financial lookup query",\n'
    '    "prompt": "What was Mars Inc organic growth for P6 2025?",\n'
    '    "expected_output": {\n'
    '      "answer_contains": "X.X%",\n'
    '      "source_table": "finiq_vw_pl_entity",\n'
    '      "source_attribution_required": true\n'
    '    },\n'
    '    "tolerance": 0.01,\n'
    '    "mode": "real",\n'
    '    "tags": ["NLQ", "Financial", "Organic Growth"]\n'
    '  },\n'
    '  {\n'
    '    "test_id": "BV_001",\n'
    '    "category": "Budget Variance",\n'
    '    "description": "Actual vs Replan variance for Mars Inc P6 2025",\n'
    '    "prompt": "Show budget variance for Mars Inc P6 2025",\n'
    '    "expected_output": {\n'
    '      "actual": "X.X",\n'
    '      "replan": "X.X",\n'
    '      "variance_abs": "X.X",\n'
    '      "variance_pct": "X.X%"\n'
    '    },\n'
    '    "tolerance": 0.01,\n'
    '    "mode": "real",\n'
    '    "tags": ["Budget", "Variance", "Replan"]\n'
    '  }\n'
    ']'
)

p = doc.add_paragraph()
run = p.add_run(template_json)
run.font.name = 'Consolas'
run.font.size = Pt(8)
run.font.color.rgb = DARK_GRAY
p.paragraph_format.space_after = Pt(8)

add_para("Field Reference:", bold=True, space_after=4)
add_table(
    ["Field", "Type", "Required", "Description"],
    [
        ["test_id", "String", "Yes", "Unique identifier for the test case (e.g., PES_001, NLQ_001)"],
        ["category", "String", "Yes", "Test category: PES, NL Query, Budget Variance, CI, Job Board"],
        ["description", "String", "Yes", "Human-readable description of what the test validates"],
        ["prompt", "String", "Yes", "The exact prompt to submit to FinIQ"],
        ["expected_output", "Object", "Yes", "The expected response structure and values"],
        ["tolerance", "Float", "No", "Numeric tolerance (default: 0.01 = 0.01%)"],
        ["mode", "String", "Yes", "Must be 'real' for prompt-response pairs"],
        ["tags", "Array", "No", "Tags for filtering and grouping test cases"],
    ],
    col_widths=[1.2, 0.8, 0.8, 3.7]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
#  APPENDIX B: SIMULATED DATA SUMMARY
# ══════════════════════════════════════════════════════════════════
doc.add_heading("Appendix B: Simulated Data Summary", level=1)
add_para(
    "This appendix provides a quick reference for the synthetic data available in simulated mode. "
    "The synthetic dataset mirrors the production Databricks schema with deterministic, reproducible "
    "values."
)

add_para("Dimension Tables (11):", bold=True, space_after=4)
add_table(
    ["Table", "Rows", "Key Characteristics"],
    [
        ["finiq_date", "730", "2 fiscal years (2024\u20132025), P1\u2013P13 per year"],
        ["finiq_dim_entity", "150+", "Synthetic org hierarchy: Mars Inc > GBUs > Divisions > Regions > Sub-units"],
        ["finiq_dim_account", "200+", "Account master with Sign_Conversion and parent IDs (array)"],
        ["finiq_account_formula", "50+", "KPI calculation logic (numerator/denominator pairs)"],
        ["finiq_account_input", "100+", "Account input configurations"],
        ["finiq_composite_item", "500+", "12-column product master (brands, categories)"],
        ["finiq_item", "1,000+", "15-column granular product detail"],
        ["finiq_item_composite_item", "1,500+", "Bridge table linking items to composite items"],
        ["finiq_customer", "300+", "11-column customer master"],
        ["finiq_customer_map", "600+", "Customer hierarchy mapping"],
        ["finiq_economic_cell", "100+", "Economic cell definitions"],
    ],
    col_widths=[2.0, 0.8, 3.7]
)

add_para("Fact Tables (5):", bold=True, space_after=4)
add_table(
    ["Table", "Rows", "Key Characteristics"],
    [
        ["finiq_financial", "50,000+", "39-column denormalized wide table (primary fact table)"],
        ["finiq_financial_base", "30,000+", "7-column normalized base financials"],
        ["finiq_financial_cons", "20,000+", "9-column with currency (used by views)"],
        ["finiq_financial_replan", "15,000+", "18-column actual vs. replan (budget variance source)"],
        ["finiq_financial_replan_cons", "10,000+", "6-column consolidated replan"],
    ],
    col_widths=[2.5, 1.0, 3.0]
)

add_para("Views (3):", bold=True, space_after=4)
add_table(
    ["View", "Maps To", "Output Columns"],
    [
        ["finiq_vw_pl_entity", "PES P&L sheet", "YTD_LY, YTD_CY, Periodic_LY, Periodic_CY by entity"],
        ["finiq_vw_pl_brand_product", "PES Product/Brand sheets", "YTD_LY, YTD_CY, Periodic_LY, Periodic_CY by brand/product (3-way UNION ALL)"],
        ["finiq_vw_ncfo_entity", "PES NCFO sheet", "YTD_LY, YTD_CY, Periodic_LY, Periodic_CY by entity"],
    ],
    col_widths=[2.2, 1.5, 2.8]
)

add_para(
    "Note: All synthetic data is generated with seed 42 for full reproducibility. Row counts are "
    "approximate and may vary slightly based on the data generation script version.",
    italic=True, space_after=8
)

# ── END OF DOCUMENT ──
add_para("\u2014 End of Document \u2014", bold=True, size=11, color=RGBColor(0x99, 0x99, 0x99))

# ══════════════════════════════════════════════════════════════════
#  SAVE
# ══════════════════════════════════════════════════════════════════
output_path = r"D:\Amira FinIQ\Testing Agent SRS\FinIQ Testing Agent SRS v1.1.docx"
doc.save(output_path)
print(f"Document saved to: {output_path}")
print(f"Generated on: {datetime.date.today().strftime('%B %d, %Y')}")
print(f"Test Requirements: 31 (TR1.1 through TR9.3)")
print(f"Acceptance Criteria: 15 (AC-1 through AC-15)")
print(f"Binary Pass/Fail Criteria: 31 across 7 TR groups")
print(f"New in v1.1: Quantitative Evaluation Framework (Karpathy methodology)")
