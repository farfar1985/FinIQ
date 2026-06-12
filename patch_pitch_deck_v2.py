#!/usr/bin/env python3
"""
Patch AMIRA_PITCH_DECK_V2.docx (Rajiv's version, on Desktop) with the
QDL/QML/Q Marketing terminology fixes — applies the same changes we made
to our project markdown, but to Rajiv's docx in place.

Reads:  C:\\Users\\farza\\Desktop\\AMIRA_PITCH_DECK_V2.docx
Writes: C:\\Users\\farza\\Desktop\\AMIRA_PITCH_DECK_V3.docx
"""

import sys
import copy
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

DESKTOP = Path("C:/Users/farza/Desktop")
SRC = DESKTOP / "AMIRA_PITCH_DECK_V2.docx"
DST = DESKTOP / "AMIRA_PITCH_DECK_V3.docx"


# Plain text replacements (applied to every paragraph in body + every cell paragraph)
# Order matters: longer, more-specific phrases first.
REPLACEMENTS = [
    # Slide 10 — main framing line (longer phrase first)
    (
        "Mars's proprietary skills are first-class platform primitives, not afterthoughts:",
        "Proprietary APIs are pre-integrated as first-class skills, accessible to Mars associates through the platform — not afterthoughts, not re-built per project:",
    ),
    # Slide 10 — definition of skill (broaden to include APIs)
    (
        "A skill is a reusable platform capability — bound to a data source, an analytical method, or a service",
        "The platform exposes capabilities to AI agents through skills — reusable primitives that wrap a data source, an API, an analytical method, or a service",
    ),
    # Slide 10 — QML row (specific phrase)
    (
        "Machine-learning models trained on Mars data",
        "ML model APIs — train and deploy machine-learning models on enterprise data",
    ),
    # Slide 10 — Q Marketing row
    (
        "Marketing analytics, campaign intelligence",
        "Marketing analytics and campaign intelligence APIs",
    ),
    # Slide 10 — Replit/Cursor line
    (
        "they live in your environment",
        "they don't have integrations with the proprietary APIs that power Mars's analytical workflows",
    ),
    # Slide 18 — architecture line (replace; we'll insert the new "Proprietary APIs" line separately)
    (
        "Mars data lake (QDL, QML, Q Marketing)",
        "Mars data lake (Databricks)",
    ),
    # Slide 10 title / general
    (
        "Proprietary Skills Layer",
        "Proprietary APIs",
    ),
    # Slide 15 (Rajiv's feature taxonomy)
    (
        "Proprietary Skills Integration",
        "Proprietary APIs Integration",
    ),
    # Slide 3 / Slide 10 / fallback
    (
        "Mars's proprietary skills",
        "Proprietary APIs",
    ),
    # Generic Title-Case fallback
    (
        "Proprietary Skills",
        "Proprietary APIs",
    ),
    # Generic lowercase fallback
    (
        "proprietary skills",
        "proprietary APIs",
    ),
]

# Table-cell-only replacements (Slide 10 differentiator table)
TABLE_HEADER_REPLACEMENTS = {
    # exact cell text → new cell text
    "Skill": "API",
    "Capability": "Capability (accessed through the platform)",
}


def replace_in_paragraph(paragraph, old, new):
    """Replace old→new in a paragraph, preserving formatting where possible."""
    if old not in paragraph.text:
        return False
    # Try simple per-run replacement first (preserves run-level formatting)
    for run in paragraph.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            return True
    # Fall back: text spans runs. Concatenate all run text, do the replace,
    # put the result into run[0], blank the rest.
    if not paragraph.runs:
        return False
    full_text = "".join(run.text for run in paragraph.runs)
    if old not in full_text:
        return False
    new_text = full_text.replace(old, new)
    paragraph.runs[0].text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""
    return True


def apply_text_replacements(paragraph):
    changed = 0
    for old, new in REPLACEMENTS:
        if replace_in_paragraph(paragraph, old, new):
            changed += 1
    return changed


def walk_document_paragraphs(doc):
    """Yield all paragraphs in the body + all paragraphs inside tables."""
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def apply_table_header_fixes(doc):
    """Slide 10 table: change header row from 'Skill / Capability' to 'API / Capability (...)'.

    Heuristic: any table whose header row first cell text == 'Skill' (exact, after strip)
    is the differentiator table.
    """
    fixed = 0
    for table in doc.tables:
        if not table.rows:
            continue
        first_row = table.rows[0]
        if not first_row.cells:
            continue
        first_cell_text = first_row.cells[0].paragraphs[0].text.strip() if first_row.cells[0].paragraphs else ""
        if first_cell_text == "Skill":
            # rewrite header cells
            for cell in first_row.cells:
                cell_text = cell.paragraphs[0].text.strip() if cell.paragraphs else ""
                if cell_text in TABLE_HEADER_REPLACEMENTS:
                    new_val = TABLE_HEADER_REPLACEMENTS[cell_text]
                    # Clear cell and write fresh while preserving the first run's formatting
                    p = cell.paragraphs[0]
                    if p.runs:
                        p.runs[0].text = new_val
                        for r in p.runs[1:]:
                            r.text = ""
                    else:
                        p.add_run(new_val)
                    fixed += 1
    return fixed


def insert_paragraph_after(paragraph, text):
    """Insert a new paragraph immediately after the given paragraph,
    copying its style and prepending a bullet-style dash.
    """
    new_p = copy.deepcopy(paragraph._p)
    # Wipe runs in the copy
    for r in new_p.findall(qn("w:r")):
        new_p.remove(r)
    paragraph._p.addnext(new_p)
    # Wrap in Paragraph object so we can use add_run
    from docx.text.paragraph import Paragraph
    new_para = Paragraph(new_p, paragraph._parent)
    run = new_para.add_run(text)
    return new_para


def insert_proprietary_apis_line_after_databricks(doc):
    """After 'Mars data lake (Databricks)' bullet, insert a new bullet for 'Proprietary APIs'."""
    target_text = "Mars data lake (Databricks)"
    new_line = "Proprietary APIs (QDL, QML, Q Marketing) — accessed via the platform's pre-wired skill layer"
    inserted = 0
    for p in list(doc.paragraphs):
        if target_text in p.text and "Proprietary APIs (QDL" not in p.text:
            # Avoid double-insertion: check the next paragraph
            try:
                # Look at sibling element in XML
                next_elem = p._p.getnext()
                if next_elem is not None and next_elem.tag == qn("w:p"):
                    # Check if it already has the proprietary APIs text
                    next_text = "".join(t.text or "" for t in next_elem.iter(qn("w:t")))
                    if "Proprietary APIs (QDL" in next_text:
                        continue
            except Exception:
                pass
            insert_paragraph_after(p, new_line)
            inserted += 1
    return inserted


def main():
    if not SRC.exists():
        print(f"ERROR: source file not found: {SRC}")
        sys.exit(1)

    doc = Document(str(SRC))

    # 1. Plain text replacements across all paragraphs (body + tables)
    total_changed = 0
    for p in walk_document_paragraphs(doc):
        total_changed += apply_text_replacements(p)
    print(f"Plain text replacements: {total_changed}")

    # 2. Slide 10 table header fix
    header_fixed = apply_table_header_fixes(doc)
    print(f"Table header cells fixed: {header_fixed}")

    # 3. Insert the new 'Proprietary APIs' paragraph after the Databricks line on Slide 18
    inserted = insert_proprietary_apis_line_after_databricks(doc)
    print(f"New 'Proprietary APIs' bullets inserted: {inserted}")

    doc.save(str(DST))
    print(f"\nWrote {DST}")
    print(f"Size: {DST.stat().st_size:,} bytes")


if __name__ == "__main__":
    main()
