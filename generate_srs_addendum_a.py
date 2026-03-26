"""
Amira FinIQ — SRS Addendum A: Databricks / FinSight Data Layer Integration
Generates a separate addendum document that references SRS v2.1 (base) and
specifies amended, extended, and new requirements based on the FinIQ UC
Documentation (Matt Hutton's Databricks schema, 46 pages, 20 objects).

Process: "Amending specs with incremental requirements" — per Rajiv Chandrasekaran
Base Document: FinIQ SRS IEEE Format v2.1 Merged.docx (frozen, not modified)
Source Material: FinIQ UC Documentation (corporate_finance_analytics_dev.finsight_core_model_mvp3)
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

# ── Branding colors (matching base SRS) ──
NAVY = RGBColor(0x0A, 0x1F, 0x44)
BLUE = RGBColor(0x1A, 0x56, 0xDB)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
LIGHT_BLUE = RGBColor(0xE8, 0xF0, 0xFE)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
TABLE_BORDER_COLOR = "CCCCCC"
ACCENT = "1A56DB"
AMEND_COLOR = RGBColor(0xB4, 0x5F, 0x06)  # Orange-brown for amendment markers

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# ── Style definitions ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = DARK_GRAY

for level, (size, color) in enumerate([(22, NAVY), (16, BLUE), (13, BLUE), (12, BLUE)], 1):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Calibri'
    h.font.size = Pt(size)
    h.font.color.rgb = color
    h.font.bold = True
    if level <= 2:
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(8)
    else:
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(6)

# ── Helper functions ──

def add_para(text, bold=False, italic=False, size=None, color=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold: run.font.bold = True
    if italic: run.font.italic = True
    if size: run.font.size = Pt(size)
    if color: run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    if level > 0:
        p.paragraph_format.left_indent = Inches(0.5 * (level + 1))
    p.paragraph_format.space_after = Pt(2)
    return p

def add_numbered(text):
    p = doc.add_paragraph(text, style='List Number')
    p.paragraph_format.space_after = Pt(2)
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE
        run.font.name = 'Calibri'
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{ACCENT}"/>')
        cell._element.get_or_add_tcPr().append(shading)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.cell(i + 1, j)
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.name = 'Calibri'
            run.font.color.rgb = DARK_GRAY
            if i % 2 == 1:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F6FC"/>')
                cell._element.get_or_add_tcPr().append(shading)
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="{TABLE_BORDER_COLOR}"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="{TABLE_BORDER_COLOR}"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{TABLE_BORDER_COLOR}"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="{TABLE_BORDER_COLOR}"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{TABLE_BORDER_COLOR}"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="{TABLE_BORDER_COLOR}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)
    if col_widths:
        for i_row in range(len(rows) + 1):
            for j, w in enumerate(col_widths):
                table.cell(i_row, j).width = Inches(w)
    doc.add_paragraph()
    return table

def add_req_table(req_id, title, description, priority="High", category="Functional", change_type="AMENDED"):
    """Add a requirement with change-type marker."""
    p = doc.add_paragraph()
    # Change type tag
    run = p.add_run(f"[{change_type}] ")
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = AMEND_COLOR
    # Req ID and title
    run = p.add_run(f"{req_id}: {title}")
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = BLUE
    p.paragraph_format.space_after = Pt(4)
    add_para(description, space_after=2)
    p2 = doc.add_paragraph()
    run2 = p2.add_run(f"Priority: {priority}  |  Category: {category}  |  Change: {change_type}")
    run2.font.size = Pt(9)
    run2.font.italic = True
    run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    p2.paragraph_format.space_after = Pt(10)


# ═══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════════════════════════

for _ in range(5):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Amira FinIQ")
run.font.size = Pt(36)
run.font.color.rgb = NAVY
run.font.bold = True
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Unified Financial Analytics Hub")
run.font.size = Pt(18)
run.font.color.rgb = BLUE
run.font.name = 'Calibri'

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("SRS Addendum A")
run.font.size = Pt(20)
run.font.color.rgb = NAVY
run.font.bold = True
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Databricks / FinSight Data Layer Integration")
run.font.size = Pt(16)
run.font.color.rgb = BLUE
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Amending Specs with Incremental Requirements")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run.font.italic = True
run.font.name = 'Calibri'

doc.add_paragraph()
doc.add_paragraph()

meta = [
    ("Document Type", "SRS Addendum (Incremental Amendment)"),
    ("Addendum ID", "A"),
    ("Base Document", "Amira FinIQ SRS v2.1 (IEEE 830 Format)"),
    ("Base Document Date", "March 25, 2026"),
    ("Source Material", "FinIQ UC Documentation (Databricks schema, 46 pages)"),
    ("Source Author", "dipendra.das@effem.com (Mars / Effem)"),
    ("Addendum Date", datetime.date.today().strftime("%B %d, %Y")),
    ("Classification", "Confidential"),
    ("Prepared For", "Mars, Incorporated"),
    ("Prepared By", "Amira Technologies"),
]
table = doc.add_table(rows=len(meta), cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (k, v) in enumerate(meta):
    for j, text in enumerate([k, v]):
        cell = table.cell(i, j)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.color.rgb = DARK_GRAY
        run.font.name = 'Calibri'
        if j == 0:
            run.font.bold = True

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 1. ADDENDUM PURPOSE & SCOPE
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading("1. Addendum Purpose and Scope", level=1)

doc.add_heading("1.1 Purpose", level=2)
add_para(
    "This addendum amends the Amira FinIQ Software Requirements Specification v2.1 (the 'Base SRS') "
    "based on new source material: the FinIQ UC Documentation provided by Matt Hutton's team at Mars. "
    "This documentation describes the Databricks / FinSight data layer (20 tables and views in the "
    "finsight_core_model_mvp3 schema) that will serve as FinIQ's primary data source for internal "
    "financial analytics."
)
add_para(
    "This addendum follows the process of 'Amending specs with incremental requirements' \u2014 "
    "the Base SRS remains frozen and unchanged; this document specifies only the amendments, "
    "extensions, and new requirements derived from the Databricks schema analysis."
)

doc.add_heading("1.2 Scope of Changes", level=2)
add_para("This addendum affects the following sections of the Base SRS:")
add_table(
    ["Base SRS Section", "Change Type", "Summary"],
    [
        ["1.2.2 Out of Scope", "AMENDED", "Remove 'no direct SAP/ERP integration' \u2014 Databricks connectivity is now in scope"],
        ["1.3 Definitions", "EXTENDED", "Add Databricks/FinSight-specific terminology"],
        ["1.4 References", "EXTENDED", "Add FinIQ UC Documentation as a reference"],
        ["2.1 Product Perspective", "EXTENDED", "Add FinSight data layer as current-state context"],
        ["2.4 Constraints", "AMENDED", "Update data source constraint from Excel-only to Databricks-primary"],
        ["2.5 Assumptions", "AMENDED", "Update data availability assumptions"],
        ["3.2.1 FR1: Data Ingestion", "AMENDED", "FR1.1 amended from Excel to Databricks; new FR1.6 added"],
        ["3.2.2 FR2: Analytics", "EXTENDED", "Precomputed views enable direct analytics; new FR2.7 added"],
        ["3.2.6 FR6: Integration", "EXTENDED", "FR6.1 strengthened with replan data; new FR6.5 added"],
        ["3.2.7 FR7: Admin", "EXTENDED", "New FR7.6 for Databricks connection management"],
        ["4. Data Model", "AMENDED", "Replace conceptual entities with actual Databricks schema"],
        ["5. System Architecture", "AMENDED", "Add Databricks connector to Layer 4; amend data flows"],
        ["8. Acceptance Criteria", "EXTENDED", "Add criteria for Databricks integration"],
    ],
    col_widths=[2.0, 1.2, 3.3]
)

doc.add_heading("1.3 Change Type Legend", level=2)
add_table(
    ["Marker", "Meaning"],
    [
        ["[AMENDED]", "An existing base SRS requirement is modified. The addendum text supersedes the base text for that requirement."],
        ["[EXTENDED]", "An existing base SRS section receives additional content. The base text remains valid; this addendum adds to it."],
        ["[NEW]", "A new requirement not present in the base SRS. Adds net-new functionality or specifications."],
    ],
    col_widths=[1.5, 5.0]
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 2. SOURCE MATERIAL SUMMARY
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading("2. Source Material: FinSight Databricks Schema", level=1)

doc.add_heading("2.1 Schema Identity", level=2)
add_table(
    ["Property", "Value"],
    [
        ["Databricks Catalog", "corporate_finance_analytics_dev"],
        ["Schema (Database)", "finsight_core_model_mvp3"],
        ["Table/View Prefix", "finiq"],
        ["Storage Format", "Delta Lake"],
        ["Storage Location", "abfss://output@finsightmvp31218devsa.dfs.core.windows.net/"],
        ["Owner Group", "Finsight-Group-Mvp3"],
        ["Documentation Source", "FinIQ UC Documentation, generated 2026-03-25 by dipendra.das@effem.com"],
        ["Total Objects", "20 (17 tables + 3 views)"],
    ],
    col_widths=[2.0, 4.5]
)

doc.add_heading("2.2 Object Inventory", level=2)
add_table(
    ["#", "Object Name", "Type", "Cols", "Category"],
    [
        ["1",  "finiq_account_formula",        "TABLE", "4",  "Dimension \u2014 Account Metadata"],
        ["2",  "finiq_account_input",          "TABLE", "3",  "Dimension \u2014 Account Metadata"],
        ["3",  "finiq_composite_item",         "TABLE", "12", "Dimension \u2014 Product Taxonomy"],
        ["4",  "finiq_customer",               "TABLE", "11", "Dimension \u2014 Customer"],
        ["5",  "finiq_customer_map",           "TABLE", "5",  "Dimension \u2014 Customer Hierarchy"],
        ["6",  "finiq_date",                   "TABLE", "4",  "Dimension \u2014 Time"],
        ["7",  "finiq_dim_account",            "TABLE", "6",  "Dimension \u2014 Account Hierarchy"],
        ["8",  "finiq_dim_entity",             "TABLE", "5",  "Dimension \u2014 Org Hierarchy"],
        ["9",  "finiq_economic_cell",          "TABLE", "3",  "Dimension \u2014 Economic Cell"],
        ["10", "finiq_financial",              "TABLE", "39", "Fact \u2014 Denormalized (Wide)"],
        ["11", "finiq_financial_base",         "TABLE", "7",  "Fact \u2014 Normalized (Base)"],
        ["12", "finiq_financial_cons",         "TABLE", "9",  "Fact \u2014 Consolidated"],
        ["13", "finiq_financial_replan",       "TABLE", "18", "Fact \u2014 Actual vs. Replan"],
        ["14", "finiq_financial_replan_cons",  "TABLE", "6",  "Fact \u2014 Consolidated Replan"],
        ["15", "finiq_item",                   "TABLE", "15", "Dimension \u2014 Product (Granular)"],
        ["16", "finiq_item_composite_item",    "TABLE", "3",  "Dimension \u2014 Item Mapping"],
        ["17", "finiq_rls_last_change",        "TABLE", "2",  "System \u2014 Row-Level Security"],
        ["18", "finiq_vw_ncfo_entity",         "VIEW",  "7",  "View \u2014 NCFO by Entity"],
        ["19", "finiq_vw_pl_brand_product",    "VIEW",  "8",  "View \u2014 P&L by Brand/Product"],
        ["20", "finiq_vw_pl_entity",           "VIEW",  "7",  "View \u2014 P&L by Entity"],
    ],
    col_widths=[0.3, 2.3, 0.6, 0.5, 2.8]
)

doc.add_heading("2.3 Key Findings", level=2)
add_bullet(
    "Three precomputed views (vw_pl_entity, vw_pl_brand_product, vw_ncfo_entity) directly correspond "
    "to 3 of the 4 PES Excel input sheets (P&L, Product/Brand, NCFO). FinIQ can query these views "
    "instead of processing Excel files, eliminating the preprocessing pipeline."
)
add_bullet(
    "The finiq_financial_replan table contains Actual vs. Replan (budget) values side-by-side "
    "(Actual_USD_Value, Actual_Local_Value, Replan_USD_Value, Replan_Local_Value), enabling "
    "variance analysis that the current PES system does not support."
)
add_bullet(
    "The finiq_financial table (39 columns) is a fully denormalized wide fact table containing "
    "all dimension attributes inlined, suitable for direct analytical queries without joins."
)
add_bullet(
    "The finiq_account_formula table provides programmatic access to KPI calculation logic "
    "(Account, Formula, Components array), enabling dynamic KPI configuration without hardcoding."
)
add_bullet(
    "The finiq_dim_entity table encodes the full Mars organizational hierarchy (150+ units) with "
    "Parent/Child entity IDs and Entity_Level, directly mapping to the hierarchy the base SRS describes."
)
add_bullet(
    "Multi-currency support is native: finiq_financial_cons includes both USD_Value and Local_Value "
    "with a Currency_ID column."
)
add_bullet(
    "Row-level security tracking (finiq_rls_last_change) is already present in the schema, "
    "aligning with the base SRS's RBAC requirements."
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 3. AMENDED & EXTENDED REQUIREMENTS
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading("3. Amended and Extended Requirements", level=1)

# ── 3.1 Section 1 Amendments ──
doc.add_heading("3.1 Amendments to Section 1: Introduction", level=2)

doc.add_heading("3.1.1 Section 1.2.2 Out of Scope [AMENDED]", level=3)
add_para("The following item in the base SRS Section 1.2.2 is amended:", italic=True, space_after=4)
add_para(
    "Base text: 'Direct SAP/ERP integration (data provided via preprocessed Excel exports)'\n\n"
    "Amended to: 'Direct SAP/ERP integration (financial data provided via Databricks/FinSight "
    "data layer; Excel exports retained as fallback during transition period)'"
)

doc.add_heading("3.1.2 Section 1.3 Definitions [EXTENDED]", level=3)
add_para("The following terms are added to the glossary:", italic=True, space_after=4)
add_table(
    ["Term", "Definition"],
    [
        ["FinSight", "Mars's internal Databricks-based financial data platform (finsight_core_model_mvp3 schema)"],
        ["Delta Lake", "Open-source storage layer on top of data lake storage providing ACID transactions, schema enforcement, and time travel"],
        ["Databricks", "Unified analytics platform used by Mars for data engineering, data science, and business intelligence"],
        ["ABFSS", "Azure Blob File System Secure \u2014 the protocol used to access Azure Data Lake Storage Gen2"],
        ["Composite Item", "A product classification entity in Mars's taxonomy combining brand, segment, technology, and category attributes"],
        ["Economic Cell", "A classification unit representing a distinct economic grouping in Mars's financial structure"],
        ["Replan", "Mars's budget/forecast revision process; the replan value represents the most recent approved budget figure"],
        ["Date_Offset", "A view-level calculation pattern: offset of 100 represents Last Year, offset of 0 represents Current Year"],
        ["View_ID", "A parameter in FinSight views: 1 = Periodic (single period) values, 2 = Year-to-Date cumulative values"],
    ],
    col_widths=[1.5, 5.0]
)

doc.add_heading("3.1.3 Section 1.4 References [EXTENDED]", level=3)
add_para("The following references are added:", italic=True, space_after=4)
add_bullet("FinIQ UC Documentation, finsight_core_model_mvp3 schema (Mars internal, 46 pages, 2026-03-25)")
add_bullet("FinIQ Databricks Schema Reference (Amira internal, compiled 2026-03-26)")

doc.add_page_break()

# ── 3.2 Section 2 Amendments ──
doc.add_heading("3.2 Amendments to Section 2: Overall Description", level=2)

doc.add_heading("3.2.1 Section 2.1 Product Perspective [EXTENDED]", level=3)
add_para("The following subsection is added after Section 2.1.2:", italic=True, space_after=4)

add_para("2.1.3a Current System: FinSight Data Layer", bold=True)
add_bullet("Function: Mars's centralized Databricks-based financial data platform, providing the upstream data layer that feeds both PES and other analytical tools")
add_bullet("Schema: finsight_core_model_mvp3 in the corporate_finance_analytics_dev catalog")
add_bullet("Structure: 20 objects \u2014 11 dimension tables, 5 fact tables, 3 precomputed analytics views, 1 system table")
add_bullet("Fact tables: finiq_financial (39-col denormalized), finiq_financial_base (7-col normalized), finiq_financial_cons (9-col consolidated with multi-currency), finiq_financial_replan (18-col actual vs. budget), finiq_financial_replan_cons (6-col consolidated replan)")
add_bullet("Precomputed views: P&L by entity, P&L by brand/product, NCFO by entity \u2014 each providing YTD and Periodic comparisons vs. Last Year")
add_bullet("Storage: Delta Lake format on Azure Blob Storage (ABFSS protocol)")
add_bullet("Significance: This data layer is the direct source of the Excel files currently used by PES. FinIQ will connect directly to these Databricks tables/views, bypassing the Excel export step.")

doc.add_heading("3.2.2 Section 2.1.3 Gap Analysis [EXTENDED]", level=3)
add_para("The following row is added to the Gap Analysis table:", italic=True, space_after=4)
add_table(
    ["Gap", "Impact", "FinIQ Resolution"],
    [
        ["Excel intermediary", "PES reads preprocessed Excel exports of Databricks data, adding latency and a manual step", "Direct Databricks/FinSight connectivity via precomputed views and normalized fact tables"],
        ["No budget variance", "Current PES has no access to replan/forecast data for variance analysis", "Direct query of finiq_financial_replan for actual-vs-budget comparisons"],
        ["Limited product breakdown", "PES has basic Brand/Product from 4 Excel sheets", "Full 27-attribute product taxonomy via finiq_composite_item + finiq_item"],
        ["No customer dimension", "PES has no customer-level analytics", "11-attribute customer dimension with 3-level hierarchy via finiq_customer"],
    ],
    col_widths=[1.8, 2.2, 2.5]
)

doc.add_heading("3.2.3 Section 2.4 Constraints [AMENDED]", level=3)
add_para("The following constraint in the base SRS is amended:", italic=True, space_after=4)
add_para(
    "Base text: 'Initial data ingestion limited to preprocessed Excel exports (no direct SAP/ERP integration)'\n\n"
    "Amended to: 'Primary data ingestion via direct Databricks/FinSight connectivity "
    "(finsight_core_model_mvp3 schema). Excel ingestion retained as fallback path during "
    "transition and for any data not yet available in Databricks. No direct SAP/ERP integration "
    "(SAP data flows through Databricks).'"
)

doc.add_heading("3.2.4 Section 2.5 Assumptions [AMENDED]", level=3)
add_para("The following assumption in the base SRS is amended:", italic=True, space_after=4)
add_para(
    "Base text: 'Preprocessed Excel workbooks will continue to follow the existing naming convention and schema'\n\n"
    "Amended to: 'Mars will provide and maintain Databricks/FinSight connectivity credentials "
    "(service principal or managed identity) with read access to the finsight_core_model_mvp3 schema. "
    "The schema structure (20 tables/views) will remain stable or changes will be communicated in advance. "
    "Excel workbook ingestion is retained as a fallback and will follow the existing naming convention.'"
)
add_para("The following assumption is added:", italic=True, space_after=4)
add_bullet(
    "The external dimension tables referenced by FinSight views (Dimensions_View_Date_Map, "
    "Dimensions_Date, Dimensions_Entity, Dimensions_Account) will be accessible to FinIQ "
    "with the same credentials as the finiq_ prefixed tables."
)

doc.add_page_break()

# ── 3.3 FR1 Amendments ──
doc.add_heading("3.3 Amendments to FR1: Unified Data Ingestion Engine", level=2)

add_req_table("FR1.1", "Internal Financial Data Ingestion",
    "The system shall ingest internal financial data via direct connectivity to Mars's "
    "Databricks/FinSight data layer (catalog: corporate_finance_analytics_dev, schema: "
    "finsight_core_model_mvp3). The primary ingestion path shall query the 3 precomputed views "
    "(finiq_vw_pl_entity, finiq_vw_pl_brand_product, finiq_vw_ncfo_entity) for Period End Summary "
    "analytics, and the normalized fact tables (finiq_financial_cons, finiq_financial_base) for "
    "ad-hoc queries. The 11 dimension tables shall be synchronized to the FinIQ data layer for "
    "entity hierarchy resolution, account structure, product taxonomy, and customer classification. "
    "Column mapping, KPI resolution, unit filtering (150+ organizational units via finiq_dim_entity), "
    "and derived metric calculations shall be applied automatically. Excel workbook ingestion "
    "(preprocessed_output_{Period}_{YearShort}.xlsx) shall be retained as a fallback path.",
    priority="Critical", change_type="AMENDED")

add_req_table("FR1.6", "Databricks Connection Management",
    "The system shall maintain a configurable connection to the Databricks/FinSight environment "
    "with the following specifications: (a) Authentication via Azure service principal or managed "
    "identity with least-privilege read access, (b) Connection parameters (catalog, schema, "
    "warehouse/cluster endpoint) configurable via admin interface without code deployment, "
    "(c) Connection health monitoring with automatic retry (3 attempts, exponential backoff), "
    "(d) Query timeout of 60 seconds with graceful fallback to cached data, (e) Connection "
    "pooling with a maximum of 10 concurrent Databricks sessions, (f) All queries logged with "
    "execution time, row count, and data freshness timestamp for audit purposes.",
    priority="High", change_type="NEW")

# ── 3.4 FR2 Amendments ──
doc.add_heading("3.4 Amendments to FR2: Financial Analytics & Reporting Engine", level=2)

add_para(
    "The availability of precomputed views in Databricks changes the analytics pipeline "
    "significantly. Instead of processing raw Excel data, FinIQ can query views that already "
    "contain YTD and Periodic comparisons versus Last Year.",
    italic=True, space_after=6
)

add_req_table("FR2.1", "Period End Summary Generation",
    "The base requirement is extended: when Databricks connectivity is available, the system "
    "shall generate Period End Summaries by querying the precomputed FinSight views directly: "
    "finiq_vw_pl_entity for P&L by entity (maps to P&L sheet), finiq_vw_pl_brand_product for "
    "P&L by brand/product (maps to Product and Brand sheets), and finiq_vw_ncfo_entity for NCFO "
    "(maps to NCFO sheet). Each view provides Date_ID, Entity_Alias, Account_Alias, and four "
    "value columns (YTD_LY_Value, YTD_CY_Value, Periodic_LY_Value, Periodic_CY_Value) that "
    "directly feed the KPI narrative generation pipeline. This eliminates the Excel preprocessing "
    "step (upload \u2192 retrieval \u2192 preprocessing \u2192 markdown conversion) from the current 10-step pipeline, "
    "reducing it to: view query \u2192 6 parallel KPI generators \u2192 trend analysis \u2192 combination \u2192 caching \u2192 SSE delivery.",
    priority="Critical", change_type="EXTENDED")

add_req_table("FR2.7", "Budget Variance Reporting",
    "The system shall generate actual-vs-budget (replan) variance reports by querying the "
    "finiq_financial_replan table. For each organizational unit and time period, the system shall "
    "display: Actual_USD_Value, Replan_USD_Value, variance (Actual - Replan), and variance "
    "percentage. Both USD and Local currency values shall be supported. Variance reports shall "
    "be available as a standalone report type and as an enrichment layer on Period End Summaries "
    "(e.g., 'Petcare Organic Growth was 6.3% vs. a replan target of 5.5%, exceeding budget by "
    "80 basis points'). The Submission_Type_ID column shall be used to distinguish between "
    "different forecast revisions.",
    priority="High", change_type="NEW")

# ── 3.5 FR6 Amendments ──
doc.add_heading("3.5 Amendments to FR6: Integration Layer", level=2)

add_req_table("FR6.1", "Amira Financial Forecasting Integration",
    "The base requirement is strengthened: in addition to calling the Amira Forecasting API for "
    "forward-looking projections, the system shall also leverage the finiq_financial_replan table "
    "for Mars's internal replan/budget data. This enables a three-way comparison: Actual (from "
    "finiq_financial_cons) vs. Replan (from finiq_financial_replan) vs. Amira Forecast (from "
    "Forecasting API). Example output: 'Q2 Organic Growth was 6.3% (Actual) vs. 5.5% (Mars Replan) "
    "vs. 6.1% (Amira Forecast). Actual exceeded both targets. Based on current trends, Amira "
    "projects 5.8% for Q3.' Confidence intervals shall be displayed for the Amira Forecast; "
    "replan figures shall be labeled as Mars internal budget targets.",
    priority="Critical", change_type="EXTENDED")

add_req_table("FR6.5", "Databricks Data Freshness Monitoring",
    "The system shall monitor data freshness in the Databricks/FinSight schema by tracking: "
    "(a) the most recent Date_ID available in each fact table and view, (b) the Last_Change "
    "timestamp from finiq_rls_last_change for security policy currency, (c) row counts per "
    "entity and period for completeness validation. Data freshness indicators shall be displayed "
    "on all reports (e.g., 'Data as of Period 6, 2026'). Stale data alerts shall be triggered "
    "when expected period data is not available within a configurable window (default: 5 business "
    "days after period close).",
    priority="High", change_type="NEW")

# ── 3.6 FR7 Amendments ──
doc.add_heading("3.6 Amendments to FR7: Administration & Configuration", level=2)

add_req_table("FR7.6", "Databricks Connection Administration",
    "Administrators shall manage Databricks/FinSight connectivity through a dedicated admin "
    "interface. Configuration shall include: (a) Databricks workspace URL and warehouse/cluster "
    "endpoint, (b) Authentication credentials (service principal client ID/secret or managed "
    "identity), (c) Catalog and schema selection with object discovery (list available tables/views), "
    "(d) View-to-report mapping configuration (which views feed which report types), (e) Dimension "
    "table sync schedule (frequency for refreshing local copies of dimension data), (f) Connection "
    "test functionality with diagnostic output, (g) Fallback toggle: switch between Databricks "
    "primary mode and Excel fallback mode without code deployment.",
    priority="High", change_type="NEW")

doc.add_page_break()

# ── 3.7 Data Model Amendment ──
doc.add_heading("3.7 Amendments to Section 4: Data Model", level=2)

add_para(
    "The base SRS Section 4.1 defines 14 conceptual entities. This addendum maps those entities "
    "to the actual Databricks tables and identifies new entities not present in the base model.",
    space_after=6
)

doc.add_heading("3.7.1 Entity Mapping: Base SRS to Databricks", level=3)
add_table(
    ["Base SRS Entity", "Databricks Table(s)", "Status", "Notes"],
    [
        ["OrganizationalUnit", "finiq_dim_entity", "MAPPED", "Parent/Child hierarchy with Entity_Level; 150+ units"],
        ["FinancialPeriod", "finiq_date", "MAPPED", "Date_ID (int), Year, Period, Quarter"],
        ["KPIDefinition", "finiq_account_formula + finiq_dim_account", "MAPPED", "Account hierarchy + formula definitions provide KPI logic"],
        ["FinancialFact", "finiq_financial_cons (primary) + finiq_financial + finiq_financial_base", "MAPPED", "Three representations: denormalized (39-col), normalized (7-col), consolidated (9-col)"],
        ["CompetitorCompany", "\u2014", "UNCHANGED", "Not in Databricks schema (CI data separate)"],
        ["CompetitorDocument", "\u2014", "UNCHANGED", "Not in Databricks schema (CI data separate)"],
        ["DocumentChunk", "\u2014", "UNCHANGED", "Not in Databricks schema (CI data separate)"],
        ["GeneratedReport", "\u2014", "UNCHANGED", "FinIQ application entity (Redis/Cosmos)"],
        ["Job", "\u2014", "UNCHANGED", "FinIQ application entity (Service Bus/SQL)"],
        ["Agent", "\u2014", "UNCHANGED", "FinIQ application entity"],
        ["PeerGroup", "\u2014", "UNCHANGED", "FinIQ application entity"],
        ["PromptTemplate", "\u2014", "UNCHANGED", "FinIQ application entity (Cosmos DB)"],
        ["AuditLog", "\u2014", "UNCHANGED", "FinIQ application entity (Cosmos DB)"],
        ["DashboardConfig", "\u2014", "UNCHANGED", "FinIQ application entity"],
    ],
    col_widths=[1.5, 2.3, 1.0, 1.7]
)

doc.add_heading("3.7.2 New Entities from Databricks (Not in Base SRS)", level=3)
add_table(
    ["Databricks Table", "Entity Purpose", "Key Attributes"],
    [
        ["finiq_composite_item", "Product master taxonomy (12 attributes)", "Composite_Item_ID, EC_Group, Technology, Supply_Tech, Segment, Business_Segment, Market_Segment, Brand_ID, Brand, Consumer_Pack_Format, Product_Consolidation, Product_Category"],
        ["finiq_item", "Granular product dimension (15 attributes, ID+Alias pairs)", "Item_ID, EC_Group_ID/Alias, Brand_Flag_ID/Alias, Financial_Product_Segment_ID/Alias, Market_Segment_ID/Alias, Supply_Tech_ID/Alias, Business_Segment_ID/Alias, Product_Category_Consolidation_ID/Alias"],
        ["finiq_item_composite_item", "Bridge table: item \u2192 composite item", "Item_ID, Composite_Item_ID, IT_EC_Group_ID"],
        ["finiq_customer", "Customer dimension (11 attributes)", "Entity_Customer_ID, Customer_ID, Country, Customer_Name, SCM_ID, Customer_Level_1/2/3, Customer_Channel, Customer_Format, Customer_Subformat"],
        ["finiq_customer_map", "Customer hierarchy mapping", "Child_Entity_ID, Child_Customer_ID, Parent_Entity_ID, Parent_Customer_ID, Entity_Customer_ID"],
        ["finiq_economic_cell", "Economic cell classification", "Economic_Cell_ID, Economic_Cell, Archetype"],
        ["finiq_account_input", "Account-to-statement mapping", "Statement, Account_ID, Generation"],
        ["finiq_financial_replan", "Actual vs. Replan fact (18 cols)", "Submission_Type_ID, Date_ID, Year, Quarter, Entity, Account_KPI, Actual_USD/Local_Value, Replan_USD/Local_Value"],
        ["finiq_financial_replan_cons", "Consolidated replan fact (6 cols)", "Date_ID, Entity_ID, Account_ID, Currency_ID, USD_Value, Local_Value"],
        ["finiq_rls_last_change", "Row-level security tracking", "Last_Change (timestamp), Version (bigint)"],
    ],
    col_widths=[1.5, 2.0, 3.0]
)

doc.add_page_break()

# ── 3.8 Architecture Amendments ──
doc.add_heading("3.8 Amendments to Section 5: System Architecture", level=2)

doc.add_heading("3.8.1 Layer 4: Data Layer [AMENDED]", level=3)
add_para("The base SRS Layer 4 description is amended to include Databricks connectivity:", italic=True, space_after=4)

add_para("Layer 4: Data Layer (Amended)", bold=True)
add_bullet("Primary Data Source: Databricks/FinSight (finsight_core_model_mvp3 schema via Databricks SQL Connector or Unity Catalog REST API)")
add_bullet("Unified Data Store: Azure SQL Database for FinIQ application data (jobs, configs, audit logs) + synchronized dimension data from Databricks")
add_bullet("Document Store: Azure Blob Storage (raw documents) + Azure AI Search (indexed chunks) \u2014 unchanged")
add_bullet("Cache Layer: Azure Redis Cache for generated reports, Databricks query results, and KPI summaries \u2014 unchanged")
add_bullet("Metadata Catalog: Azure Cosmos DB for data lineage, audit logs, and configuration \u2014 unchanged")
add_bullet("Fallback Data Source: Azure Blob Storage for Excel workbook ingestion when Databricks is unavailable")

doc.add_heading("3.8.2 Component Topology [EXTENDED]", level=3)
add_para("The following component is added to the base SRS component topology table:", italic=True, space_after=4)
add_table(
    ["Component", "Technology", "Purpose"],
    [
        ["Databricks Connector", "Databricks SQL Connector for Python", "Read access to FinSight schema (views, facts, dimensions)"],
        ["Dimension Sync Service", "Python + Azure Functions", "Periodic sync of Databricks dimension tables to local Azure SQL for low-latency lookups"],
        ["Data Freshness Monitor", "Python + Azure Functions", "Monitors Databricks data currency and triggers stale-data alerts"],
    ],
    col_widths=[1.8, 2.5, 2.2]
)

doc.add_heading("3.8.3 Data Flow 1: Period End Summary [AMENDED]", level=3)
add_para("The base SRS Data Flow 1 (10 steps) is amended to reflect Databricks-primary ingestion:", italic=True, space_after=4)

add_para("Flow 1: Period End Summary Generation (Amended)", bold=True)
add_numbered("User selects organizational unit, period, year, and output format")
add_numbered("System checks Redis cache for existing summary")
add_numbered("If cache miss: queries Databricks views (finiq_vw_pl_entity, finiq_vw_pl_brand_product, finiq_vw_ncfo_entity) filtered by Date_ID and Entity_Alias")
add_numbered("If Databricks unavailable: falls back to Excel workbook from Blob Storage (original 10-step pipeline)")
add_numbered("View results (YTD_LY, YTD_CY, Periodic_LY, Periodic_CY per Account) converted to structured format")
add_numbered("Optionally: finiq_financial_replan queried for budget variance enrichment")
add_numbered("Six parallel LLM calls (one per KPI) via LangGraph, each with tailored prompt")
add_numbered("Trend analysis: multi-period data from Databricks (multiple Date_IDs) for trend taglines")
add_numbered("Six KPI narratives + trend taglines + optional variance data combined into unified Executive Summary")
add_numbered("HTML KPI tables appended for data verification")
add_numbered("Result cached in Redis (JSON) and streamed to frontend via SSE")

add_para(
    "Net effect: Steps 3-4 of the original pipeline (Excel retrieval + preprocessing) are "
    "consolidated into a single Databricks view query, reducing latency and eliminating the "
    "Excel intermediary. The pipeline goes from 10 steps to 11 steps but the critical path "
    "is shorter because view queries return pre-aggregated data.",
    italic=True, size=10
)

doc.add_page_break()

# ── 3.9 Acceptance Criteria ──
doc.add_heading("3.9 Amendments to Section 8: Acceptance Criteria", level=2)

add_para("The following acceptance criteria are added:", italic=True, space_after=4)
add_table(
    ["Requirement", "Acceptance Criterion", "Verification Method"],
    [
        ["FR1.1 (Amended)", "System generates Period End Summary from Databricks views in < 12 seconds (faster than Excel path)", "Automated comparison test: Databricks path vs. Excel path timing"],
        ["FR1.6 Databricks Connection", "System connects to Databricks, queries all 20 objects, and returns results within 10 seconds", "Integration test against Databricks dev environment"],
        ["FR2.7 Budget Variance", "Variance report shows Actual vs. Replan with < 0.01% calculation error", "Cross-validation against manual Excel variance calculation"],
        ["FR6.5 Data Freshness", "Stale data alert fires within 1 hour of expected data not appearing", "Automated test with delayed data injection"],
        ["FR7.6 Admin Connection", "Administrator can configure, test, and switch Databricks connection without code deployment", "UI walkthrough test of admin interface"],
        ["Databricks Fallback", "System automatically falls back to Excel ingestion within 30 seconds of Databricks connection failure", "Chaos test: terminate Databricks connectivity mid-operation"],
    ],
    col_widths=[1.5, 3.0, 2.0]
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 4. PRECOMPUTED VIEW MAPPING
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading("4. Precomputed View Mapping to PES", level=1)

add_para(
    "This section provides the detailed mapping between the three FinSight precomputed views and "
    "the current PES system's Excel input sheets and KPI structure."
)

doc.add_heading("4.1 View-to-Excel Sheet Mapping", level=2)
add_table(
    ["PES Excel Sheet", "Databricks View", "Output Columns", "Value Columns"],
    [
        ["P&L (Entity level)", "finiq_vw_pl_entity",
         "Date_ID, Entity_Alias, Account_Alias",
         "YTD_LY_Value, YTD_CY_Value, Periodic_LY_Value, Periodic_CY_Value"],
        ["Product / Brand", "finiq_vw_pl_brand_product",
         "Date_ID, Entity_Alias, Account_Alias, Item",
         "YTD_LY_Value, YTD_CY_Value, Periodic_LY_Value, Periodic_CY_Value"],
        ["NCFO", "finiq_vw_ncfo_entity",
         "Date_ID, Entity_Alias, Account_Alias",
         "YTD_LY_Value, YTD_CY_Value, Periodic_LY_Value, Periodic_CY_Value"],
    ],
    col_widths=[1.3, 1.8, 1.8, 1.6]
)

doc.add_heading("4.2 PES KPI-to-Account Code Mapping", level=2)
add_para(
    "The following maps the 6 PES KPIs to the Databricks account codes filtered in the views. "
    "This mapping enables FinIQ to extract KPI-specific data from view results by filtering on "
    "Account_Alias (resolved from Account_ID via the Dimensions_Account external table)."
)
add_table(
    ["PES KPI", "View", "Key Account Codes", "Derivation"],
    [
        ["Organic Growth", "vw_pl_entity", "S900083, S900227, S900067, S900070, S900069, S900077",
         "Derived: growth parent-child numerator/denominator pattern"],
        ["MAC Shape %", "vw_pl_entity", "FR4100, FR4200, FR4300, FR4000, MR5200, MR5100",
         "Direct: revenue and margin accounts"],
        ["A&CP Shape %", "vw_pl_entity", "SR5101, SR5103, MR8005",
         "Direct: advertising & consumer promotion accounts"],
        ["CE Shape %", "vw_pl_entity", "SR6102, SR6153, MR8004",
         "Direct: controllable expense accounts"],
        ["Overhead Shape %", "vw_pl_entity", "MR6300, MR6359, MR8003",
         "Direct: overhead margin accounts"],
        ["NCFO", "vw_ncfo_entity", "CF8129, CF8133, MC8136, MC8149, CF8147, MC8913, S900147, MC8100, MC8902",
         "Direct: cash flow and management control accounts"],
    ],
    col_widths=[1.3, 1.2, 2.2, 1.8]
)

doc.add_heading("4.3 View SQL Computation Patterns", level=2)
add_para(
    "The views follow specific computation patterns that FinIQ must understand when interpreting "
    "query results:", space_after=4
)
add_bullet("Date_Offset = 100 represents Last Year values; Date_Offset = 0 represents Current Year values")
add_bullet("View_ID = 1 represents Periodic (single period) values; View_ID = 2 represents YTD (year-to-date) values")
add_bullet("Account S900077 receives special treatment: its Date_Offset starts at 100 (not 0) and uses a +200 offset for Last Year comparison")
add_bullet("Growth KPIs (S900083, S900071, S900073, S900072) are derived from child accounts via a numerator/denominator pattern, not stored directly as values")
add_bullet("The vw_pl_brand_product view creates 3 parallel item breakdowns via UNION ALL: by Brand (COALESCE with 'EMPTY BRAND'), by Product_Category ('EMPTY PCAT'), and by Product_Consolidation ('EMPTY PCONS')")
add_bullet("All monetary values are ROUND'd to 4 decimal places in view outputs")
add_bullet("Views use GROUP BY ALL (Databricks SQL extension) for aggregation")
add_bullet("Views LEFT JOIN to Dimensions_Account (some accounts may lack aliases) but INNER JOIN to Dimensions_Entity (all entities must resolve)")


doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 5. SIMULATED DATA STRATEGY
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading("5. Simulated Data Strategy", level=1)

add_para(
    "Per Rajiv Chandrasekaran's direction: FinIQ development shall proceed using simulated data "
    "that mirrors the Databricks schema structure while Mars Datalake access is being provisioned. "
    "The simulation shall be comprehensive enough to validate all Databricks-dependent features "
    "end-to-end."
)

doc.add_heading("5.1 Simulation Requirements", level=2)
add_bullet("Generate synthetic data for all 20 Databricks objects matching exact column names, data types, and array structures")
add_bullet("Organizational hierarchy: simulate 150+ units mirroring the Mars structure (Mars Inc > 4 GBUs > Divisions > Regions > Sub-units)")
add_bullet("Financial data: generate realistic P&L and NCFO values with period-over-period trends, seasonal patterns, and variance ranges typical of FMCG companies")
add_bullet("Replan data: generate budget targets that are 90-110% of actuals to create realistic variance scenarios")
add_bullet("Product taxonomy: generate 100+ composite items across multiple brands, segments, and categories")
add_bullet("Customer data: generate 50+ customers with hierarchical relationships and channel/format classifications")
add_bullet("Date dimension: populate for FY2024 and FY2025 with Mars fiscal calendar periods")
add_bullet("View outputs: simulate the precomputed view outputs (YTD_LY, YTD_CY, Periodic_LY, Periodic_CY) with mathematically consistent values")

doc.add_heading("5.2 Simulation Delivery", level=2)
add_bullet("Simulated data shall be loadable into a local SQLite or PostgreSQL database for development and testing")
add_bullet("A Databricks-compatible API mock shall be provided that returns simulated data in the same format as the real Databricks SQL connector")
add_bullet("Integration tests shall run against both simulated and real Databricks environments with a single configuration toggle")


# ═══════════════════════════════════════════════════════════════════════════════
# 6. CHANGE SUMMARY
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_page_break()
doc.add_heading("6. Change Summary", level=1)

add_para("This addendum introduces the following changes to the base SRS v2.1:")

add_table(
    ["Change ID", "Type", "Requirement", "Description"],
    [
        ["A-001", "AMENDED",  "FR1.1", "Data ingestion: Excel \u2192 Databricks/FinSight primary with Excel fallback"],
        ["A-002", "NEW",      "FR1.6", "Databricks connection management (auth, pooling, retry, logging)"],
        ["A-003", "EXTENDED", "FR2.1", "PES generation from precomputed Databricks views instead of Excel"],
        ["A-004", "NEW",      "FR2.7", "Budget variance reporting (actual vs. replan from finiq_financial_replan)"],
        ["A-005", "EXTENDED", "FR6.1", "Three-way comparison: Actual vs. Replan vs. Amira Forecast"],
        ["A-006", "NEW",      "FR6.5", "Databricks data freshness monitoring and stale-data alerts"],
        ["A-007", "NEW",      "FR7.6", "Databricks connection administration UI"],
        ["A-008", "AMENDED",  "Sec 1.2.2", "Out of Scope updated: Databricks connectivity now in scope"],
        ["A-009", "EXTENDED", "Sec 1.3", "9 new glossary terms for Databricks/FinSight concepts"],
        ["A-010", "EXTENDED", "Sec 1.4", "2 new references (FinIQ UC Documentation, Schema Reference)"],
        ["A-011", "EXTENDED", "Sec 2.1", "FinSight data layer added to current-state description"],
        ["A-012", "EXTENDED", "Sec 2.1.3", "4 new rows in Gap Analysis table"],
        ["A-013", "AMENDED",  "Sec 2.4", "Constraint: Excel-only \u2192 Databricks-primary with Excel fallback"],
        ["A-014", "AMENDED",  "Sec 2.5", "Assumption: Excel schema stability \u2192 Databricks schema stability"],
        ["A-015", "AMENDED",  "Sec 4", "Data model: conceptual entities mapped to actual Databricks tables; 10 new entities"],
        ["A-016", "AMENDED",  "Sec 5", "Architecture Layer 4 amended; 3 new components; Data Flow 1 amended"],
        ["A-017", "EXTENDED", "Sec 8", "6 new acceptance criteria for Databricks integration"],
    ],
    col_widths=[0.8, 1.0, 1.0, 3.7]
)

add_para(
    "Total: 4 amended requirements, 4 new requirements, 9 extended sections, yielding a net addition "
    "of 4 new functional requirements (FR1.6, FR2.7, FR6.5, FR7.6) bringing the effective total from "
    "46 to 50 functional requirements when this addendum is applied to the base SRS.",
    bold=True, space_after=10
)


# ═══════════════════════════════════════════════════════════════════════════════
# FOOTER
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("--- End of Addendum A ---")
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.font.italic = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Amira FinIQ SRS Addendum A | Confidential | Prepared for Mars, Incorporated")
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)


# ═══════════════════════════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════════════════════════

output_path = r"D:\Amira FinIQ\FinIQ SRS Addendum A - Databricks Integration.docx"
doc.save(output_path)
print(f"Addendum A saved to: {output_path}")
