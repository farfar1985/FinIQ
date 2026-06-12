#!/usr/bin/env python3
"""
Generate SPEC_AGENT_DESIGN.docx from SPEC_AGENT_DESIGN.md
Professional IEEE-style Word document.
"""

import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


ROOT = Path(__file__).resolve().parent
MD_PATH = ROOT / "SPEC_AGENT_DESIGN.md"
# Optional CLI arg: output filename (without path). Default = SPEC_AGENT_DESIGN.docx
_docx_name = sys.argv[1] if len(sys.argv) > 1 else "SPEC_AGENT_DESIGN.docx"
DOCX_PATH = ROOT / _docx_name


# ---------- Helpers ----------

def set_cell_shading(cell, color_hex):
    """Shade a table cell with a hex color (e.g. '1F3864')."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_page_number(doc):
    """Add simple page numbering in footer."""
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.text = "PAGE"
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr)
    run._r.append(fld_char_2)


def add_formatted_runs(paragraph, text):
    """Parse inline markdown (**bold**, *italic*, `code`) into runs."""
    # Escape &amp; etc are not a concern for plain-text runs.
    # Tokenize: find the nearest special marker each step.
    i = 0
    n = len(text)
    while i < n:
        # Bold **...**
        if text[i:i + 2] == "**":
            end = text.find("**", i + 2)
            if end > i + 2:
                run = paragraph.add_run(text[i + 2:end])
                run.bold = True
                i = end + 2
                continue
        # Inline code `...`
        if text[i] == "`":
            end = text.find("`", i + 1)
            if end > i:
                run = paragraph.add_run(text[i + 1:end])
                run.font.name = "Consolas"
                run.font.size = Pt(10)
                i = end + 1
                continue
        # Italic *...*  (avoid matching bullet markers)
        if (text[i] == "*"
                and (i + 1 < n)
                and text[i + 1] != " "
                and text[i + 1] != "*"):
            end = text.find("*", i + 1)
            if end > i + 1 and (end + 1 >= n or text[end + 1] != "*"):
                run = paragraph.add_run(text[i + 1:end])
                run.italic = True
                i = end + 1
                continue
        # Plain character run — accumulate until next marker
        j = i
        while j < n:
            if text[j] == "`":
                break
            if text[j:j + 2] == "**":
                break
            if text[j] == "*" and (j + 1 < n) and text[j + 1] != " ":
                break
            j += 1
        paragraph.add_run(text[i:j])
        i = j


def add_code_block(doc, code_text):
    """Add a monospace code / diagram block with subtle shading."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def add_table_from_rows(doc, header, data_rows):
    """Add a styled table. Header row shaded dark blue with white text."""
    if not header:
        return
    cols = len(header)
    table = doc.add_table(rows=1, cols=cols)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header
    hdr_cells = table.rows[0].cells
    for idx, h in enumerate(header):
        hdr_cells[idx].text = ""
        p = hdr_cells[idx].paragraphs[0]
        run = p.add_run(h.strip())
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(hdr_cells[idx], "1F3864")

    # Data rows
    for row in data_rows:
        cells = table.add_row().cells
        for idx in range(cols):
            cell_text = row[idx].strip() if idx < len(row) else ""
            cells[idx].text = ""
            p = cells[idx].paragraphs[0]
            add_formatted_runs(p, cell_text)

    # Small space after the table
    doc.add_paragraph()


def add_blockquote(doc, lines):
    """Render blockquote as indented italic paragraph."""
    text = "\n".join(lines)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.right_indent = Cm(0.8)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    add_formatted_runs(p, text)
    # Italicize whole paragraph
    for run in p.runs:
        run.italic = True


# ---------- Main MD-to-docx renderer ----------

def render_md(md_text, doc):
    lines = md_text.split("\n")
    i = 0
    n = len(lines)

    # Skip the title block (first H1) — handled by title page separately.
    title_handled = False

    while i < n:
        line = lines[i]

        # --- Horizontal rule ---
        if line.strip() == "---":
            # Treat as a page break for section boundary
            doc.add_paragraph()
            i += 1
            continue

        # --- Headings ---
        if line.startswith("# "):
            text = line[2:].strip()
            if not title_handled:
                # The very first H1 is the title — skip (added on title page).
                title_handled = True
                i += 1
                continue
            doc.add_heading(text, 0)
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), 1)
            i += 1
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), 2)
            i += 1
            continue
        if line.startswith("#### "):
            doc.add_heading(line[5:].strip(), 3)
            i += 1
            continue

        # --- Code block / ASCII diagram ---
        if line.startswith("```"):
            code_lines = []
            i += 1
            while i < n and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            add_code_block(doc, "\n".join(code_lines))
            i += 1  # skip closing fence
            continue

        # --- Table ---
        if line.strip().startswith("|") and "|" in line.strip()[1:]:
            # Collect consecutive pipe-rows
            tbl_lines = []
            while i < n and lines[i].strip().startswith("|"):
                tbl_lines.append(lines[i])
                i += 1
            # Parse header | separator | rows
            if len(tbl_lines) >= 2 and re.search(r"\|[\s:\-|]+\|", tbl_lines[1]):
                header = [c.strip() for c in tbl_lines[0].strip().strip("|").split("|")]
                data_rows = []
                for row_line in tbl_lines[2:]:
                    cells = [c.strip() for c in row_line.strip().strip("|").split("|")]
                    data_rows.append(cells)
                add_table_from_rows(doc, header, data_rows)
            else:
                # Not a real table; emit as paragraphs
                for tl in tbl_lines:
                    p = doc.add_paragraph()
                    add_formatted_runs(p, tl)
            continue

        # --- Blockquote ---
        if line.startswith("> "):
            quote_lines = []
            while i < n and (lines[i].startswith("> ") or lines[i].startswith(">")):
                stripped = lines[i][2:] if lines[i].startswith("> ") else lines[i][1:]
                quote_lines.append(stripped)
                i += 1
            add_blockquote(doc, quote_lines)
            continue

        # --- Bullet list ---
        if line.startswith("- "):
            while i < n and lines[i].startswith("- "):
                p = doc.add_paragraph(style="List Bullet")
                add_formatted_runs(p, lines[i][2:])
                i += 1
                # Continuation lines (indented 2 spaces)
                while i < n and lines[i].startswith("  ") and lines[i].strip() and not lines[i].startswith("  -"):
                    p.add_run("\n" + lines[i].strip())
                    i += 1
            continue

        # --- Numbered list ---
        if re.match(r"^\d+\.\s", line):
            while i < n and re.match(r"^\d+\.\s", lines[i]):
                text = re.sub(r"^\d+\.\s", "", lines[i])
                p = doc.add_paragraph(style="List Number")
                add_formatted_runs(p, text)
                i += 1
            continue

        # --- Empty line ---
        if not line.strip():
            i += 1
            continue

        # --- Regular paragraph ---
        p = doc.add_paragraph()
        add_formatted_runs(p, line)
        # Continuation lines (no blank line) → join into same paragraph
        i += 1
        while i < n and lines[i].strip() and not (
            lines[i].startswith("#")
            or lines[i].startswith("- ")
            or lines[i].startswith("> ")
            or lines[i].startswith("|")
            or lines[i].startswith("```")
            or re.match(r"^\d+\.\s", lines[i])
            or lines[i].strip() == "---"
        ):
            p.add_run(" ")
            add_formatted_runs(p, lines[i])
            i += 1


def build_title_page(doc):
    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Spec Agent")
    run.font.size = Pt(36)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Design Document")
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

    # Spacing
    for _ in range(6):
        doc.add_paragraph()

    # Metadata block
    meta = [
        ("Version", "0.4 draft"),
        ("Status", "FOR TEAM REVIEW — not approved for build"),
        ("Audience", "QDT / Amira platform team"),
        ("Approval gate", "No code until team signs off on scope, architecture, and Phase 1 MVP requirements"),
    ]
    for label, value in meta:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{label}: ")
        run.bold = True
        run.font.size = Pt(12)
        run2 = p.add_run(value)
        run2.font.size = Pt(12)

    doc.add_page_break()


def configure_styles(doc):
    # Normal body
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # Headings
    for level, size in [(0, 24), (1, 18), (2, 14), (3, 12)]:
        name = "Title" if level == 0 else f"Heading {level}"
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)


def main():
    md_text = MD_PATH.read_text(encoding="utf-8")

    doc = Document()

    # Page margins
    section = doc.sections[0]
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    configure_styles(doc)
    build_title_page(doc)
    render_md(md_text, doc)
    add_page_number(doc)

    doc.save(DOCX_PATH)
    print(f"Wrote {DOCX_PATH}")
    print(f"Size: {DOCX_PATH.stat().st_size:,} bytes")


if __name__ == "__main__":
    main()
