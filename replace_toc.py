"""
Replace the manual hand-typed TOC at the top of each docx with a real Word TOC field.

The TOC field auto-populates when the document is opened in Word (updateFields=true
in settings.xml triggers silent refresh on open) — so when Farzaneh or Rajiv opens
either file, the correct page numbers and heading list appear automatically.

Operates on:
  C:\\Users\\farza\\Desktop\\Amira_Proposal_for_Mars_2026-04-26_Polished_INLINE.docx
  C:\\Users\\farza\\Desktop\\Amira_Proposal_for_Mars_2026-04-26_Polished_APPENDIX.docx

Overwrites in place.
"""

from docx import Document
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

FILES = [
    r"C:\Users\farza\Desktop\Amira_Proposal_for_Mars_2026-04-26_Polished_INLINE.docx",
    r"C:\Users\farza\Desktop\Amira_Proposal_for_Mars_2026-04-26_Polished_APPENDIX.docx",
]


def replace_manual_toc(doc):
    paras = doc.paragraphs

    # 1. Find the old "TABLE OF CONTENTS" header paragraph
    toc_start_i = None
    for i, p in enumerate(paras):
        if p.text.strip().upper() == "TABLE OF CONTENTS":
            toc_start_i = i
            break
    if toc_start_i is None:
        raise LookupError("Manual TOC header not found")

    # 2. Walk forward to find the first heading paragraph (where TOC ends)
    toc_end_i = None
    for i in range(toc_start_i + 1, len(paras)):
        if paras[i].style.name.startswith("Heading"):
            toc_end_i = i
            break
    if toc_end_i is None:
        raise LookupError("End of manual TOC not found (no heading after TOC)")

    next_heading_el = paras[toc_end_i]._element

    # 3. Collect old TOC paragraph elements for deletion
    old_toc_elements = [paras[i]._element for i in range(toc_start_i, toc_end_i)]

    # 4. Build new "Table of Contents" header paragraph (Normal style, bold, 14pt — mimics Rajiv's typography)
    new_header = doc.add_paragraph()
    run = new_header.add_run("Table of Contents")
    run.bold = True
    run.font.size = Pt(14)
    next_heading_el.addprevious(new_header._element)

    # 5. Build the TOC field paragraph as raw OOXML
    #    - <w:fldChar begin> with w:dirty="true" tells Word to refresh on open
    #    - instrText: TOC \o "1-3" \h \z \u  =  pick Heading 1-3, hyperlinked, hide tab leaders, use outline
    #    - separator + cached placeholder text + end
    toc_field_xml = (
        '<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'xmlns:xml="http://www.w3.org/XML/1998/namespace">'
        '<w:r><w:fldChar w:fldCharType="begin" w:dirty="true"/></w:r>'
        '<w:r><w:instrText xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText></w:r>'
        '<w:r><w:fldChar w:fldCharType="separate"/></w:r>'
        '<w:r><w:t>The table of contents will populate when this document is opened in Word. '
        'If it does not, right-click anywhere on this line and select Update Field.</w:t></w:r>'
        '<w:r><w:fldChar w:fldCharType="end"/></w:r>'
        '</w:p>'
    )
    toc_field_el = parse_xml(toc_field_xml)
    next_heading_el.addprevious(toc_field_el)

    # 6. Delete the old hand-typed TOC paragraphs
    for el in old_toc_elements:
        el.getparent().remove(el)

    # 7. Set settings.xml updateFields=true so Word silently refreshes the TOC field on open
    settings_el = doc.settings.element
    upd = settings_el.find(qn("w:updateFields"))
    if upd is None:
        upd = OxmlElement("w:updateFields")
        upd.set(qn("w:val"), "true")
        settings_el.append(upd)
    else:
        upd.set(qn("w:val"), "true")


def main():
    for path in FILES:
        print(f"Processing: {path}")
        doc = Document(path)
        replace_manual_toc(doc)
        doc.save(path)
        print(f"  saved.")
    print("Done.")


if __name__ == "__main__":
    main()
