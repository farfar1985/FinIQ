"""
Builds two versions of Rajiv's polished proposal with Cesar's 28 demo screenshots:

  Version A (inline)   — screenshots placed at the relevant sections in Rajiv's doc.
  Version B (appendix) — Rajiv's doc untouched + Appendix A "Workflow Walkthrough"
                          with all 28 screenshots in demo-step order.

Both versions are saved to C:\\Users\\farza\\Desktop\\.

After opening each in Word, right-click the TOC -> Update Field -> Update entire table
to refresh page numbers and pick up the appendix heading (Version B only).
"""

import os
from copy import deepcopy
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = r"C:\Users\farza\Desktop\Amira_Proposal_for_Mars_2026-04-26_Polished.docx"
SHOTS_DIR = r"C:\Users\farza\Desktop\demo-screenshots-extracted\demo-screenshots"
OUT_INLINE = r"C:\Users\farza\Desktop\Amira_Proposal_for_Mars_2026-04-26_Polished_INLINE.docx"
OUT_APPENDIX = r"C:\Users\farza\Desktop\Amira_Proposal_for_Mars_2026-04-26_Polished_APPENDIX.docx"

IMG_WIDTH = Inches(5.8)


# ----- Caption helpers ---------------------------------------------------------

def add_image_paragraph(doc, img_filename, width=IMG_WIDTH):
    """Append an image paragraph to the end of doc body. Returns the paragraph element."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_path = os.path.join(SHOTS_DIR, img_filename)
    if not os.path.exists(img_path):
        raise FileNotFoundError(f"Missing screenshot: {img_path}")
    p.add_run().add_picture(img_path, width=width)
    return p


def add_caption_paragraph(doc, caption_text):
    """Append a caption paragraph to the end of doc body. Returns the paragraph element."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(caption_text)
    run.italic = True
    run.font.size = Pt(9)
    return p


def insert_after(cursor_element, new_element):
    """Move new_element to immediately after cursor_element in the body XML."""
    cursor_element.addnext(new_element)
    return new_element


def insert_image_with_caption(doc, cursor_element, img_filename, caption_text):
    """Create image+caption paragraphs at end of doc, then move them right after cursor.
    Returns the caption element so the caller can chain (cursor = result)."""
    img_p = add_image_paragraph(doc, img_filename)
    cap_p = add_caption_paragraph(doc, caption_text)

    # add_image_paragraph and add_caption_paragraph both append to end. Move them.
    insert_after(cursor_element, img_p._element)
    insert_after(img_p._element, cap_p._element)
    return cap_p._element


def find_heading_index(doc, text):
    """Return the index in doc.paragraphs of the heading whose text starts with `text`."""
    for i, p in enumerate(doc.paragraphs):
        if p.style.name.startswith("Heading") and p.text.strip().startswith(text):
            return i
    raise LookupError(f"Heading not found: {text}")


# ----- Section -> screenshots mapping for INLINE version -----------------------
# Each entry: (section_heading_text, next_section_heading_text, [(img, caption), ...])
# Captions are figure-numbered in document order across all sections (1..28).
# The "anchor" for insertion is the last paragraph BEFORE next_section_heading_text.

INLINE_SECTIONS = [
    ("1.2 Amira at a Glance", "2. Platform Architecture", [
        ("step-01-home-portal.png",
         "Figure 1. The Amira home portal — projects, approvals inbox, and skills shelf in one workspace."),
    ]),
    ("3.1 Specifications Phase", "3.2 Development Phase (Canvas)", [
        ("step-02a-spec-new.png",
         "Figure 2. Spec Agent entry — natural-language intent with turn-budget controls."),
        ("step-02b-projects-new.png",
         "Figure 3. Structured project creation — from scratch, repo import, or fork."),
        ("step-02c-import-from-repo.png",
         "Figure 4. Repository import — auto-detected stack, agent extraction toggles, security badges."),
        ("step-03-spec-finiq.png",
         "Figure 5. Spec workspace with a Decision Point card — alternatives surfaced with trade-offs."),
    ]),
    ("3.2 Development Phase (Canvas)", "3.3 Artifacts Phase", [
        ("step-07-canvas-preview.png",
         "Figure 6. Canvas — three-panel workspace: AI chat, code editor, and live preview."),
        ("step-08a-enlarge-chart-clicked.png",
         "Figure 7. Build Agent edit (“enlarge chart”) rebuilds the preview in place."),
        ("step-08b-working-capital-added.png",
         "Figure 8. Build Agent adds a Working Capital KPI bound to live data."),
        ("step-10-compliance-matrix.png",
         "Figure 9. Compliance matrix — automated FR / NFR / AC scoring with status and evidence."),
    ]),
    ("3.3 Artifacts Phase", "3.4 Reversibility and Versioning", [
        ("step-14-project-finiq.png",
         "Figure 10. Project lineage — Spec → Build → Deploy events with companion-agent activity."),
    ]),
    ("3.4 Reversibility and Versioning", "4. Differentiators", [
        ("step-04-version-history.png",
         "Figure 11. Version history — every iteration is a separate, traceable spec version."),
    ]),
    ("4.1 Proprietary APIs", "4.2 How Skills Connect to Specifications", [
        ("step-12a-skills.png",
         "Figure 12. Skills marketplace — pre-integrated proprietary APIs and platform primitives."),
        ("step-12c-skills-bottom.png",
         "Figure 13. Skills marketplace (continued) — coverage of Mars-specific and general-purpose skills."),
        ("step-09b-add-resource-drawer.png",
         "Figure 14. Add-skills drawer — attach a skill to the running app."),
    ]),
    ("4.2 How Skills Connect to Specifications", "4.3 Apps Become Agents", [
        ("step-09a-resources-tab.png",
         "Figure 15. Resources tab — primitives bound to the app, role-scoped."),
    ]),
    ("4.3 Apps Become Agents", "5. Governance and Security", [
        ("step-12b-skills-app-agents.png",
         "Figure 16. Apps Become Agents — every shipped app surfaces as a callable companion."),
        ("step-09c-add-companion-agent.png",
         "Figure 17. Companion agents drawer — making other apps callable from this one."),
        ("step-13a-ask-amira-drawer.png",
         "Figure 18. Ask Amira — querying FinIQ Agent in-context from any surface."),
        ("step-13b-ask-amira-nestle.png",
         "Figure 19. Ask Amira — competitor comparison via the same agent, with FMP provenance."),
    ]),
    ("5.1 Human Governance and Audit", "5.2 Knowledge Base and Secret Vault", [
        ("step-05-route-esignature-modal.png",
         "Figure 20. Route-for-e-signature modal — approver locked by governance role."),
        ("step-06a-approve-page.png",
         "Figure 21. Approver review page — spec on left, signature block and validation matrix on right."),
        ("step-06b-signature-recorded.png",
         "Figure 22. Signature recorded — audit ID generated and bound to the spec version."),
    ]),
    ("5.2 Knowledge Base and Secret Vault", "6. Platform Features Summary", [
        ("step-09d-add-knowledge.png",
         "Figure 23. Knowledge tab — session-scoped uploads, private or team-shared."),
    ]),
    ("8.2 Deployment Options", "8.3 Authentication and API Key Strategy", [
        ("step-11a-deploy-publish.png",
         "Figure 24. Deploy step 1 — publish details and audience scoping."),
        ("step-11b-deploy-compliance.png",
         "Figure 25. Deploy step 2 — compliance standards and deploy-blocking thresholds."),
        ("step-11c-deploy-environment.png",
         "Figure 26. Deploy step 3 — environment, network policy, and secrets vault confirmation."),
        ("step-11d-deploy-approval.png",
         "Figure 27. Deploy step 4 — approval gate routed to the governance approver."),
        ("step-11e-deploy-progress.png",
         "Figure 28. Deploy in progress — build, scan, deploy, smoke test, and agent registration."),
    ]),
]


# ----- Appendix walkthrough steps (15 steps, 28 screenshots in demo order) -----

APPENDIX_STEPS = [
    ("Step 1 — Home portal",
     "The user lands on the Amira home portal: in-flight projects, approvals inbox, "
     "and a shelf of pre-integrated skills are all visible from a single surface.",
     [("step-01-home-portal.png",
       "The Amira home portal — projects, approvals, and skills in one view.")]),

    ("Step 2 — Three ways to start a project",
     "Amira supports three project-genesis surfaces: a natural-language entry point that hands "
     "directly to the Spec Agent; a structured creation page with a fork option for existing "
     "applications; and a guided import flow that reverse-engineers a specification from an "
     "existing repository.",
     [("step-02a-spec-new.png",
       "(2a) Natural-language entry — describe the application; the Spec Agent takes it from there."),
      ("step-02b-projects-new.png",
       "(2b) Structured project creation — from scratch, repo import, or fork."),
      ("step-02c-import-from-repo.png",
       "(2c) Repository import — auto-detected stack, agent extraction toggles, security badges.")]),

    ("Step 3 — Decision Point",
     "The user resumes an in-flight FinIQ specification. At an architectural decision point — "
     "the source of competitive-intelligence data — the Spec Agent surfaces three alternatives "
     "with trade-offs, recommends one, and records the user’s choice against a specific "
     "functional requirement.",
     [("step-03-spec-finiq.png",
       "Decision Point card — alternatives surfaced rather than silently picked.")]),

    ("Step 4 — Living specification, gaps tracker, version history",
     "The same workspace shows the IEEE-830 specification under iteration on the left, an open "
     "gaps-tracker on the right (warnings the agent has flagged for resolution), and a version-"
     "history dropdown listing every prior spec version with its approval state.",
     [("step-04-version-history.png",
       "Living spec on the left, gaps on the right, version history at the top.")]),

    ("Step 5 — Route for e-signature",
     "The user routes the iterating spec to the designated authorized approver. The approver "
     "is locked by governance role (CFO Office — Approvals), with pre-filled notes that "
     "summarize the spec version and current compliance score.",
     [("step-05-route-esignature-modal.png",
       "E-signature routing modal — approver locked by governance role.")]),

    ("Step 6 — Authorized approver view",
     "On the approver’s side, the spec is rendered alongside a digital-signature block. "
     "The approver reviews validation results, types their name to record intent, and the "
     "signature is captured with a permanent audit ID bound to that specific spec version.",
     [("step-06a-approve-page.png",
       "(6a) Pre-sign — spec on the left, signature block and validation matrix on the right."),
      ("step-06b-signature-recorded.png",
       "(6b) Post-sign — signature recorded with audit ID.")]),

    ("Step 7 — Open Canvas",
     "After approval, the user enters Canvas: a three-panel live workspace with the Build "
     "Agent chat on the left, the file tree, and a live preview of the running application on "
     "the right. The seeded build shows FinIQ rendering against real Mars financial data.",
     [("step-07-canvas-preview.png",
       "Canvas — chat, code, and live preview side by side.")]),

    ("Step 8 — Build Agent chat",
     "The user iterates conversationally. Two example edits are demonstrated: enlarging the "
     "quarterly chart, and adding a new Working Capital KPI tile bound to live data — each "
     "performed by typing in natural language; the Build Agent rebuilds the preview in place.",
     [("step-08a-enlarge-chart-clicked.png",
       "(8a) Enlarge chart — the Build Agent resizes the component and rebuilds the preview."),
      ("step-08b-working-capital-added.png",
       "(8b) Add Working Capital KPI — a new tile bound to live data.")]),

    ("Step 9 — Combine skills, knowledge, and companion agents",
     "Through the Resources tab and the Add drawer, the user attaches additional skills "
     "(QDL, QML, charting, voice), uploads knowledge files scoped to the application only, "
     "and adds companion agents from other applications — all reusable primitives that "
     "the Build Agent then composes into the running app.",
     [("step-09a-resources-tab.png",
       "(9a) Resources tab — primitives currently bound to the app, role-scoped."),
      ("step-09b-add-resource-drawer.png",
       "(9b) Skills tab — pre-integrated and marketplace skills available to add."),
      ("step-09c-add-companion-agent.png",
       "(9c) Companion Agents tab — cross-project agents available as dependencies."),
      ("step-09d-add-knowledge.png",
       "(9d) Knowledge tab — session-scoped uploads, private or team-shared.")]),

    ("Step 10 — Compliance matrix",
     "The Build Agent maintains a live compliance matrix that scores the running application "
     "against every functional and non-functional requirement and acceptance criterion in "
     "the spec, with evidence pointers to source files. Failing rows block deployment.",
     [("step-10-compliance-matrix.png",
       "Compliance matrix — automated FR / NFR / AC scoring with evidence and status.")]),

    ("Step 11 — Deploy modal",
     "The deploy flow is a four-step modal: publish details and audience scope; compliance "
     "standards and deploy-blocking thresholds; environment, network policy, and secrets "
     "resolution; and a final approval gate routed to the same governance role as the spec gate. "
     "On confirmation, build, scan, deploy, smoke test, and companion-agent registration "
     "execute sequentially.",
     [("step-11a-deploy-publish.png",
       "(11a) Step 1 — publish details and audience scoping."),
      ("step-11b-deploy-compliance.png",
       "(11b) Step 2 — compliance standards and deploy-blocking thresholds."),
      ("step-11c-deploy-environment.png",
       "(11c) Step 3 — environment, network policy, and secrets vault confirmation."),
      ("step-11d-deploy-approval.png",
       "(11d) Step 4 — approval gate routed to the governance approver."),
      ("step-11e-deploy-progress.png",
       "(11e) Deploy in progress — build, scan, deploy, smoke test, and agent registration.")]),

    ("Step 12 — Skills marketplace",
     "The Skills marketplace exposes the full catalogue of pre-integrated primitives — "
     "Mars proprietary APIs (QDL, QML, Q Marketing, Q Analytics) and general-purpose tools "
     "(charting, web search, market hours). At the top, an Apps Become Agents section lists "
     "companion agents auto-generated from every app already shipped on the platform.",
     [("step-12a-skills.png",
       "(12a) Full marketplace — skills catalog with role-scoped filtering."),
      ("step-12b-skills-app-agents.png",
       "(12b) Apps Become Agents — companion agents auto-generated from shipped apps."),
      ("step-12c-skills-bottom.png",
       "(12c) Marketplace tail — long-tail skills available to all projects.")]),

    ("Step 13 — Ask Amira",
     "From any surface, the user opens an Ask Amira drawer scoped to a chosen agent. The "
     "demo shows the FinIQ Agent answering a Q3 Petcare net-sales question with chart and "
     "drivers, then a follow-up cross-competitor comparison via FMP — same data, same "
     "permissions, same audit boundary as opening FinIQ directly.",
     [("step-13a-ask-amira-drawer.png",
       "(13a) FinIQ Agent answering a Q3 Petcare question, with chart and drivers."),
      ("step-13b-ask-amira-nestle.png",
       "(13b) Follow-up competitor comparison via the same agent, with FMP provenance.")]),

    ("Step 14 — Lineage",
     "The project page consolidates the full Spec → Build → Deploy lineage — every "
     "approved spec version, every build, every deployment, every approval. A right-rail card "
     "shows the companion FinIQ Agent’s recent activity and the audit-log preview.",
     [("step-14-project-finiq.png",
       "Project lineage — every spec, build, deploy, and approval, attributed and timestamped.")]),

    ("Step 15 — Wrap",
     "The walkthrough closes on the same project page — the user has gone from a vague "
     "intent through specification, build, governance, deploy, marketplace, in-context chat, "
     "and lineage, with the entire journey captured as one auditable artifact.",
     []),
]


# ====== VERSION A — INLINE ====================================================

def build_inline():
    print(f"Building INLINE: {OUT_INLINE}")
    doc = Document(SRC)

    for section_text, next_section_text, shots in INLINE_SECTIONS:
        # Anchor: the paragraph immediately BEFORE next_section_text
        try:
            next_idx = find_heading_index(doc, next_section_text)
        except LookupError:
            print(f"  WARN: next-section heading not found: {next_section_text}; skipping {section_text}")
            continue
        # Walk backwards from next_idx-1 — but body ordering at the XML level may include tables
        # that don't appear in doc.paragraphs. We use the next-heading element and addprevious
        # so we always insert *just before* the next heading, regardless of intervening tables.
        next_heading_element = doc.paragraphs[next_idx]._element

        # Insert each (image, caption) pair just before next_heading_element, in order.
        # Trick: addprevious() inserts immediately before the anchor; if we keep using the same
        # anchor, items end up in REVERSE order. So we use addprevious() but iterate in order,
        # and after each insertion the new caption becomes the previous-sibling of the heading
        # — i.e., the next addprevious puts the new one *between* the last caption and the heading.
        for img_filename, caption in shots:
            img_p = add_image_paragraph(doc, img_filename)
            cap_p = add_caption_paragraph(doc, caption)
            # add_*_paragraph appended at end of body; move both into place.
            next_heading_element.addprevious(img_p._element)
            next_heading_element.addprevious(cap_p._element)
            print(f"  + {section_text}  <-  {img_filename}")

    doc.save(OUT_INLINE)
    print(f"  saved: {OUT_INLINE}\n")


# ====== VERSION B — APPENDIX ==================================================

def build_appendix():
    print(f"Building APPENDIX: {OUT_APPENDIX}")
    doc = Document(SRC)

    # Append page break before appendix
    last_p = doc.paragraphs[-1]
    last_p.add_run().add_break()  # soft break

    # Appendix heading
    h = doc.add_heading("Appendix A — Workflow Walkthrough", level=1)

    intro = doc.add_paragraph(
        "The following walkthrough shows the platform end-to-end as a sequence of 15 steps, "
        "each captured directly from the running application. The flow goes from project "
        "genesis through specification, approval, Canvas build, governance, deployment, "
        "skills marketplace, in-context chat, and lineage — closing the loop on the "
        "three-step pipeline described above."
    )

    for step_title, step_blurb, shots in APPENDIX_STEPS:
        doc.add_heading(step_title, level=2)
        doc.add_paragraph(step_blurb)
        for img_filename, caption in shots:
            add_image_paragraph(doc, img_filename)
            add_caption_paragraph(doc, caption)
            print(f"  + {step_title}  <-  {img_filename}")

    doc.save(OUT_APPENDIX)
    print(f"  saved: {OUT_APPENDIX}\n")


# ====== Main ==================================================================

if __name__ == "__main__":
    if not os.path.exists(SRC):
        raise SystemExit(f"Source missing: {SRC}")
    if not os.path.isdir(SHOTS_DIR):
        raise SystemExit(f"Screenshots dir missing: {SHOTS_DIR}")

    expected = []
    for _, _, shots in INLINE_SECTIONS:
        expected += [s[0] for s in shots]
    missing = [f for f in expected if not os.path.exists(os.path.join(SHOTS_DIR, f))]
    if missing:
        raise SystemExit(f"Missing screenshot files: {missing}")

    build_inline()
    build_appendix()
    print("Done.")
