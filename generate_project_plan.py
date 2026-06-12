"""
Generate FinIQ Project Plan & Timeline document for Mars leadership visibility.
Covers: what's been done, day-by-day plan through April 21, 2026 MLT presentation.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from datetime import datetime

doc = Document()

# -- Page setup --
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# -- Styles --
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10.5)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

# Helper functions
def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)  # Dark navy
    return h

def add_table_with_style(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # Dark header background
        shading = cell._element.get_or_add_tcPr()
        shading_elm = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): '1B3A5C',
            qn('w:val'): 'clear'
        })
        shading.append(shading_elm)

    # Data rows
    for r, row_data in enumerate(rows):
        for c, cell_text in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(cell_text)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(width)

    return table

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10)
        p.add_run(text)
    else:
        p.add_run(text)
    for run in p.runs:
        run.font.size = Pt(10)
    return p

# ============================================================
# TITLE PAGE
# ============================================================
doc.add_paragraph()
doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('FinIQ')
run.bold = True
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Unified Financial Analytics Hub')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()

subtitle2 = doc.add_paragraph()
subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle2.add_run('Project Plan & Development Timeline')
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

subtitle3 = doc.add_paragraph()
subtitle3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle3.add_run('April 1 \u2013 April 21, 2026')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

doc.add_paragraph()
doc.add_paragraph()

# Meta info
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('Prepared for: ').font.size = Pt(11)
run = meta.add_run('Mars, Incorporated')
run.bold = True
run.font.size = Pt(11)

meta2 = doc.add_paragraph()
meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta2.add_run('Prepared by: ').font.size = Pt(11)
run = meta2.add_run('Amira Technologies / QDT')
run.bold = True
run.font.size = Pt(11)

meta3 = doc.add_paragraph()
meta3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta3.add_run(f'Date: April 1, 2026')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

meta4 = doc.add_paragraph()
meta4.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta4.add_run('Version 1.0')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

doc.add_page_break()

# ============================================================
# 1. EXECUTIVE SUMMARY
# ============================================================
add_heading_styled('1. Executive Summary', level=1)

doc.add_paragraph(
    'FinIQ is a Unified Financial Analytics Hub that consolidates and enhances '
    'Mars\'s existing AI analytics capabilities\u2014Period End Summary (PES) and '
    'Competitive Intelligence (CI)\u2014into a single intelligent platform with '
    'natural language querying, an enterprise agent job board, and real-time '
    'competitive market data.'
)

doc.add_paragraph(
    'The team has completed a rapid parallel-build phase where three independent '
    'implementations were developed simultaneously, each optimizing for different '
    'strengths. These have now been merged into a single unified application '
    'combining the best components from each build. The merged application is '
    'live on GitHub and has been tested against real Databricks production data.'
)

p = doc.add_paragraph()
run = p.add_run('Target milestone: ')
run.bold = True
p.add_run(
    'April 21, 2026 \u2014 Working demonstration for Mars Leadership Team (MLT) '
    'to showcase rapid development capability and secure further investment.'
)

# ============================================================
# 2. WHAT HAS BEEN ACCOMPLISHED
# ============================================================
add_heading_styled('2. Work Completed to Date', level=1)

add_heading_styled('2.1 Parallel Build Phase (March 25\u201331)', level=2)

doc.add_paragraph(
    'Three team members independently built full-stack implementations from the '
    'same SRS v3.1 specification (52 functional requirements), creating healthy '
    'competition and diverse approaches:'
)

add_table_with_style(
    ['Team Member', 'Focus Area', 'Key Deliverables', 'Strengths'],
    [
        ['Alessandro Savino\n(Atlas)', 'UI/UX & Design',
         'Next.js app, Data Explorer, OKLCH dark theme,\nDatabricks integration, FMP API, design system',
         'Bloomberg-quality UI,\ninteractive charts, design spec'],
        ['Rajiv Chandrasekaran\n(Asimov)', 'Competitive Intelligence',
         'CI module (10 tabs), Alerts system, ProvenanceBadge,\nClean header, SRS v3.1 authorship',
         'Richest CI module,\nself-contained, deployed live'],
        ['Farzaneh\n(Claude Code)', 'Backend Intelligence',
         'Anthropic LLM engine, Voice Agent, Job Board,\nXLSX export, rate limiting, schema discovery',
         '80/80 compliance,\nreal Databricks tested'],
        ['Cesar Flores', 'Architecture & Data',
         'Semantic layer (YAML), Databricks query tools,\nPlatform architecture, deployment infrastructure',
         'Production-proven queries,\nself-learning data layer'],
    ],
    col_widths=[1.4, 1.2, 2.5, 1.5]
)

doc.add_paragraph()
add_heading_styled('2.2 Three-Way Merge (March 31\u2013April 1)', level=2)

doc.add_paragraph(
    'The best components from all three builds were merged into a single unified '
    'application on Alessandro\'s repository (github.com/quantumdatatechnologies/fin_iq). '
    'Eight commits have been pushed to the merged branch:'
)

add_table_with_style(
    ['Phase', 'Description', 'Source', 'Status'],
    [
        ['Phase 1', 'Schema rename (Entity\u2192Unit, Account\u2192RL) to match production Databricks', 'New', '\u2705 Pushed'],
        ['Phase 2a', 'Anthropic LLM query engine + real Databricks schema context', 'Farzaneh', '\u2705 Pushed'],
        ['Phase 2b', 'Voice Agent (OpenAI Realtime API) + WebSocket proxy + UI', 'Farzaneh', '\u2705 Pushed'],
        ['Phase 3a', 'Rajiv\'s CI module (10 tabs, Alerts, ProvenanceBadge, charts)', 'Rajiv', '\u2705 Pushed'],
        ['Phase 3b', 'Job Board backend + XLSX export + rate limiting', 'Farzaneh', '\u2705 Pushed'],
        ['Phase 4', 'Compliance fixes (all FAILs resolved) + Cesar\'s semantic layer', 'All', '\u2705 Pushed'],
        ['Phase 4c', 'PR merge to main branch — team pulls main to get unified app', 'All', '\u2705 Pushed'],
    ],
    col_widths=[0.9, 3.0, 1.0, 0.9]
)

doc.add_paragraph()
add_heading_styled('2.3 Live Testing Results (March 31)', level=2)

doc.add_paragraph('The merged application was tested locally with real production connections:')

add_bullet('Real Databricks connected ', bold_prefix='\u2705 ')
doc.add_paragraph('     Production data flowing (P&L, MAC, revenue queries working)', style='Normal')
add_bullet('FMP API connected ', bold_prefix='\u2705 ')
doc.add_paragraph('     Real competitor stock prices and financials (10 competitors)', style='Normal')
add_bullet('Anthropic Claude LLM connected ', bold_prefix='\u2705 ')
doc.add_paragraph('     Natural language queries generating SQL against real Databricks', style='Normal')
add_bullet('OpenAI API configured ', bold_prefix='\u2705 ')
doc.add_paragraph('     Voice agent infrastructure ready', style='Normal')
add_bullet('Compliance: 68/80 items PASS (85%), 0 FAILs remaining', bold_prefix='\u2705 ')

doc.add_paragraph()
add_heading_styled('2.4 Specification & Documentation', level=2)

add_bullet('SRS v3.1 Final ', bold_prefix='')
doc.add_paragraph('     52 functional requirements, CI/FMP API integration, suggested prompt catalog', style='Normal')
add_bullet('Frontend Design Guideline v1.0 ', bold_prefix='')
doc.add_paragraph('     Bloomberg-inspired design system (OKLCH, IBM Plex Sans, shadcn/ui)', style='Normal')
add_bullet('Real Databricks Schema Reference ', bold_prefix='')
doc.add_paragraph('     21 objects documented, 5.7B+ row tables, all relationships mapped', style='Normal')
add_bullet('Semantic Layer (YAML) ', bold_prefix='')
doc.add_paragraph('     7 YAML files describing every table, column, relationship, and metric', style='Normal')

# ============================================================
# 3. CURRENT STATUS & KNOWN ISSUES
# ============================================================
doc.add_page_break()
add_heading_styled('3. Current Status', level=1)

add_heading_styled('3.1 Application Capabilities', level=2)

add_table_with_style(
    ['Capability', 'Status', 'Notes'],
    [
        ['Dashboard (6 KPIs, charts)', 'Working', 'Recharts area/bar charts, KPI cards'],
        ['Data Explorer (SQL builder)', 'Working', 'Interactive query builder, table/chart output'],
        ['PES Reports (narratives)', 'Working', 'AI-generated summaries, What\'s Working/Not Working'],
        ['NL Query Engine (Anthropic)', 'Working', 'Claude Haiku generating SQL from natural language'],
        ['Competitive Intelligence (10 tabs)', 'Working', 'FMP API, SWOT, Porter\'s, Alerts, ESG'],
        ['Job Board (submission + agents)', 'Working', 'SLA routing, lifecycle, WebSocket streaming'],
        ['XLSX Export', 'Working', 'Mars-branded spreadsheet output'],
        ['Voice Agent', 'Partial', 'WebSocket proxy working; browser mic needs HTTPS'],
        ['Real Databricks', 'Working', 'Production catalog connected, views queried'],
        ['Suggested Prompts', 'Working', '18 curated prompts with variable resolution'],
        ['Admin Panel', 'Working', 'Connection config, template management'],
    ],
    col_widths=[2.2, 0.8, 3.0]
)

doc.add_paragraph()
add_heading_styled('3.2 Items Requiring Resolution', level=2)

add_table_with_style(
    ['Issue', 'Impact', 'Resolution Path', 'Priority'],
    [
        ['Unit alias matching', 'Some NL queries fail\n(informal unit names)', 'Add fuzzy matching / lookup table\nfor 766 org units', 'High'],
        ['Voice agent mic', 'Voice input not capturing\nin browser', 'Requires HTTPS or\nAudioWorklet replacement', 'Medium'],
        ['Chart follow-ups', '"Show me a chart" runs\nnew query vs re-rendering', 'Add follow-up detection\nand data re-use logic', 'Medium'],
        ['Production build', 'Only dev mode tested', 'Run next build, fix any\nSSR/build errors', 'High'],
        ['Multi-turn context', 'LLM loses entity/period\nacross turns', 'Add persistent session\nmemory for queries', 'Medium'],
    ],
    col_widths=[1.3, 1.5, 1.8, 0.8]
)

# ============================================================
# 4. DAY-BY-DAY PLAN
# ============================================================
doc.add_page_break()
add_heading_styled('4. Development Timeline: April 1\u201321, 2026', level=1)

doc.add_paragraph(
    'The following day-by-day plan organizes remaining work into three phases, '
    'culminating in a polished demo for the April 21 MLT presentation.'
)

# WEEK 1
add_heading_styled('Week 1: Stabilization & Core Fixes (April 1\u20134)', level=2)

add_table_with_style(
    ['Date', 'Day', 'Workstream', 'Owner(s)', 'Deliverable'],
    [
        ['Apr 1\n(Tue)', 'Day 1', 'Project plan & alignment\nMerge branch cleanup',
         'Farzaneh\nCesar', 'This document delivered\nPR review for merged \u2192 main'],
        ['Apr 2\n(Wed)', 'Day 2', 'Unit alias resolution\nProduction build test',
         'Farzaneh\nAlessandro', 'Fuzzy matching for 766 org units\nnext build passing'],
        ['Apr 3\n(Thu)', 'Day 3', 'Multi-turn query context\nVoice agent HTTPS setup',
         'Farzaneh\nCesar', 'Session memory persists entity/period\nSSL cert for voice testing'],
        ['Apr 4\n(Fri)', 'Day 4', 'Chart follow-up detection\nCI module QA',
         'Farzaneh\nRajiv', 'Re-render existing data as chart\nAll 10 CI tabs verified with live data'],
    ],
    col_widths=[0.6, 0.5, 1.8, 1.0, 2.2]
)

p = doc.add_paragraph()
run = p.add_run('Week 1 goal: ')
run.bold = True
p.add_run('All core queries working reliably with real Databricks. Production build passing. Zero query failures on standard use cases.')

doc.add_paragraph()

# WEEK 2
add_heading_styled('Week 2: Enhancement & Integration (April 7\u201311)', level=2)

add_table_with_style(
    ['Date', 'Day', 'Workstream', 'Owner(s)', 'Deliverable'],
    [
        ['Apr 7\n(Mon)', 'Day 5', 'Cesar semantic layer integration\nA2A protocol review',
         'Cesar\nBill', 'YAML-driven query routing live\nMCP/A2A integration plan documented'],
        ['Apr 8\n(Tue)', 'Day 6', 'Voice agent end-to-end\nDashboard polish',
         'Farzaneh\nAlessandro', 'Voice \u2192 query \u2192 response working\nKPI cards, charts, responsive layout'],
        ['Apr 9\n(Wed)', 'Day 7', 'Job board UX improvements\nBudget variance queries',
         'Farzaneh\nCesar', 'Job submission \u2192 result flow polished\nActual vs replan queries working'],
        ['Apr 10\n(Thu)', 'Day 8', 'Compliance sweep\nGemini documentation review',
         'Farzaneh\nRajiv', 'Push compliance from 85% \u2192 90%+\nA2A toolkit assessment complete'],
        ['Apr 11\n(Fri)', 'Day 9', 'VM deployment test\nEnd-to-end integration test',
         'Cesar\nAll', 'App running on Azure VM\nAll pages tested with real data'],
    ],
    col_widths=[0.6, 0.5, 1.8, 1.0, 2.2]
)

p = doc.add_paragraph()
run = p.add_run('Week 2 goal: ')
run.bold = True
p.add_run('Application deployed on Azure VM. Voice agent functional. All compliance items at 90%+. Integration with Cesar\'s platform architecture validated.')

doc.add_paragraph()

# WEEK 3
add_heading_styled('Week 3: Demo Prep & Polish (April 14\u201321)', level=2)

add_table_with_style(
    ['Date', 'Day', 'Workstream', 'Owner(s)', 'Deliverable'],
    [
        ['Apr 14\n(Mon)', 'Day 10', 'Demo script creation\nUI/UX final polish',
         'Rajiv\nAlessandro', 'Step-by-step demo narrative\nAll pages pixel-perfect'],
        ['Apr 15\n(Tue)', 'Day 11', 'Demo rehearsal #1\nBug fixes from rehearsal',
         'All\nFarzaneh', 'First run-through recorded\nCritical issues identified & fixed'],
        ['Apr 16\n(Wed)', 'Day 12', 'Demo data curation\nPerformance optimization',
         'Cesar\nFarzaneh', 'Curated queries that showcase best\nQuery response times < 5 seconds'],
        ['Apr 17\n(Thu)', 'Day 13', 'Demo rehearsal #2\nFallback scenarios prepared',
         'All\nCesar', 'Smooth run-through achieved\nOffline backup if Databricks is slow'],
        ['Apr 18\n(Fri)', 'Day 14', 'Final bug fixes\nDocumentation package',
         'Farzaneh\nRajiv', 'All known issues resolved\nTechnical summary for Mars team'],
        ['Apr 21\n(Mon)', 'Demo\nDay', 'MLT Presentation\n& Live Demo',
         'All', 'Working demo of FinIQ:\n\u2022 NL queries on real Databricks\n\u2022 CI with live competitor data\n\u2022 Voice agent (if HTTPS ready)\n\u2022 Job board submission flow\n\u2022 Dashboard with KPI charts'],
    ],
    col_widths=[0.6, 0.5, 1.8, 1.0, 2.2]
)

p = doc.add_paragraph()
run = p.add_run('Week 3 goal: ')
run.bold = True
p.add_run('Polished, rehearsed demo ready for Mars leadership. Two full rehearsals completed. Fallback plan in place.')

# ============================================================
# 5. KEY MILESTONES
# ============================================================
doc.add_page_break()
add_heading_styled('5. Key Milestones', level=1)

add_table_with_style(
    ['Date', 'Milestone', 'Success Criteria', 'Owner'],
    [
        ['Apr 1', 'Project plan delivered', 'Plan reviewed and approved by team', 'Farzaneh'],
        ['Apr 4', 'Core stability achieved', 'All standard NL queries working on real Databricks\nZero failures on top 20 use cases', 'Farzaneh'],
        ['Apr 7', 'Semantic layer live', 'YAML-driven query routing active\nCesar\'s data agent integrated', 'Cesar'],
        ['Apr 11', 'Azure VM deployment', 'App running on provisioned VM\nAccessible to team for testing', 'Cesar'],
        ['Apr 11', '90%+ compliance', 'Compliance score 72+/80 items passing', 'Farzaneh'],
        ['Apr 15', 'Demo rehearsal #1', 'Full end-to-end demo completed\nGap list documented', 'All'],
        ['Apr 17', 'Demo rehearsal #2', 'Smooth run, < 3 issues remaining', 'All'],
        ['Apr 21', 'MLT Presentation', 'Live demo showcasing FinIQ capabilities\nto Mars leadership team', 'All'],
    ],
    col_widths=[0.7, 1.5, 2.5, 0.9]
)

# ============================================================
# 6. TEAM RESPONSIBILITIES
# ============================================================
add_heading_styled('6. Team Responsibilities', level=1)

add_table_with_style(
    ['Team Member', 'Primary Responsibility', 'Key Deliverables (Apr 1\u201321)'],
    [
        ['Cesar Flores', 'Architecture, Deployment,\nData Layer',
         '\u2022 Azure VM deployment & infrastructure\n\u2022 Semantic layer integration (YAML \u2192 query routing)\n\u2022 Databricks query optimization\n\u2022 A2A/MCP protocol decision\n\u2022 Demo environment setup'],
        ['Alessandro Savino', 'UI/UX, Design System,\nFrontend Polish',
         '\u2022 Production build fixes\n\u2022 Dashboard & chart polish\n\u2022 Responsive layout QA\n\u2022 Design system consistency check\n\u2022 Demo visual quality'],
        ['Rajiv Chandrasekaran', 'CI Module, QA,\nDemo Script',
         '\u2022 CI module QA with live FMP data\n\u2022 Demo script & narrative creation\n\u2022 Compliance verification\n\u2022 A2A/Gemini documentation review\n\u2022 Technical documentation'],
        ['Farzaneh', 'Backend Intelligence,\nIntegration, Coordination',
         '\u2022 Unit alias resolution & fuzzy matching\n\u2022 Multi-turn query context\n\u2022 Voice agent completion\n\u2022 Chart follow-up detection\n\u2022 Compliance score optimization\n\u2022 Day-to-day coordination'],
        ['Bill Dennis', 'Platform Strategy,\nGovernance',
         '\u2022 MCP/A2A protocol review\n\u2022 Human governance workflow integration\n\u2022 Enterprise scalability planning\n\u2022 Platform convergence coordination'],
        ['Atif Ishaq', 'Governance,\nStakeholder Management',
         '\u2022 MLT presentation preparation\n\u2022 Funding narrative\n\u2022 Mars stakeholder communication\n\u2022 Resource coordination (Kumar, Marc)'],
    ],
    col_widths=[1.3, 1.3, 3.5]
)

# ============================================================
# 7. DEPENDENCIES & RISKS
# ============================================================
add_heading_styled('7. Dependencies & Risks', level=1)

add_heading_styled('7.1 Dependencies', level=2)

add_table_with_style(
    ['Dependency', 'Owner', 'Status', 'Impact if Delayed'],
    [
        ['Azure VM provisioning', 'Kumar / Marc\n(Mars infra)', 'In progress', 'Cannot deploy demo externally;\nfallback to localhost demo'],
        ['Databricks production access\n(read permissions)', 'Matt Hutton\n(Mars data)', 'Active', 'Already connected; need continued\naccess through April 21'],
        ['HTTPS certificate\n(for voice agent)', 'Cesar', 'Not started', 'Voice demo runs without mic input;\ncan demo text-only fallback'],
        ['A2A / Gemini documentation', 'Mars team\n(shared via Teams)', 'Received', 'Review needed by Apr 10 to\nassess integration feasibility'],
    ],
    col_widths=[1.5, 1.2, 0.9, 2.2]
)

doc.add_paragraph()
add_heading_styled('7.2 Risks & Mitigations', level=2)

add_table_with_style(
    ['Risk', 'Likelihood', 'Mitigation'],
    [
        ['Databricks query latency\n(5.7B row tables)', 'Medium', 'Views pre-filter data; semantic layer optimizes queries;\nCesar\'s YAML routing avoids full table scans'],
        ['VM not provisioned in time', 'Low', 'Demo from local machine with real Databricks connection;\nscreen share during MLT presentation'],
        ['Voice agent not ready\nfor demo', 'Medium', 'Text-based NL queries are primary demo flow;\nvoice is a bonus feature, not critical path'],
        ['FMP API rate limits\nduring demo', 'Low', 'Cache competitor data before demo;\nfallback to pre-fetched snapshots'],
    ],
    col_widths=[1.5, 0.8, 3.5]
)

# ============================================================
# 8. TECHNICAL ARCHITECTURE SUMMARY
# ============================================================
add_heading_styled('8. Technical Architecture', level=1)

doc.add_paragraph(
    'The merged application uses a modern full-stack architecture designed for '
    'enterprise-grade financial analytics:'
)

add_table_with_style(
    ['Layer', 'Technology', 'Purpose'],
    [
        ['Frontend', 'Next.js 15, React 19, Tailwind CSS 4,\nshadcn/ui, Recharts', 'Bloomberg-inspired dark-first UI\nwith interactive charts and dashboards'],
        ['LLM Engine', 'Anthropic Claude (queries),\nOpenAI (voice agent)', 'Natural language \u2192 SQL generation,\nvoice interaction, AI narratives'],
        ['Data Layer', 'Databricks (production),\nFMP API (competitors)', 'Real financial data (P&L, MAC, NCFO)\n+ live competitor market data'],
        ['Semantic Layer', 'YAML definitions (7 files),\n3-tier schema index', 'Column/relationship/metric descriptions\nfor accurate LLM query generation'],
        ['Job Board', 'WebSocket + in-memory queue,\nSLA routing', 'Enterprise agent job board for\n100+ concurrent user queries'],
        ['Infrastructure', 'Azure VM (target),\nNode.js 20, Git/GitHub', 'Deployment target: Azure resource group\nEAA-CORPAIML-SANDBOX-EUS2-DEV-RG'],
    ],
    col_widths=[1.0, 2.0, 2.5]
)

# ============================================================
# 9. PLATFORM VISION — REUSABLE AGENT ARCHITECTURE
# ============================================================
add_heading_styled('9. Platform Vision: Reusable Agent Architecture', level=1)

doc.add_paragraph(
    'FinIQ is being developed as the first proof-of-concept within a broader '
    'Unified AI Data Platform strategy. The architecture is designed around four '
    'reusable agentic capabilities that can be applied across multiple business '
    'units and use cases beyond financial analytics.'
)

add_heading_styled('9.1 Reusable Technical Components', level=2)

add_table_with_style(
    ['Component', 'Capability', 'FinIQ Application', 'Cross-Use-Case Potential'],
    [
        ['Orchestration\nAgent', 'Routes queries to\nspecialized agents,\nmanages workflow', 'Job Board routes NL queries\nto PES, CI, or Databricks\nagents based on intent', 'Any multi-agent system\nneeding intelligent\nquery routing'],
        ['Data Fetcher\n(Text-to-SQL)', 'Converts natural\nlanguage to structured\ndata queries', 'Anthropic LLM generates SQL\nagainst Databricks views;\nsemantic layer guides accuracy', 'Health data, supply chain,\nmarketing analytics\u2014\nany SQL-backed domain'],
        ['Data Science\nAgent', 'Statistical analysis,\nanomalies, trend\ndetection', 'KPI trend analysis,\nbudget variance detection,\ncompetitor benchmarking', 'Chi-square tests, P-values,\nregression\u2014applicable\nacross domains'],
        ['Visualization\nAgent', 'Auto-selects chart\ntypes, renders\ninteractive visuals', 'Recharts area/bar/treemap,\nSimpleChart auto-detection,\nKPI sparklines', 'Reusable charting for any\ndata domain without\nper-use-case rebuilds'],
    ],
    col_widths=[1.1, 1.3, 1.8, 1.6]
)

doc.add_paragraph()
add_heading_styled('9.2 Integration Architecture', level=2)

doc.add_paragraph(
    'The platform supports multiple integration patterns to enable cross-team '
    'collaboration and component reuse:'
)

add_bullet('Direct API integration for synchronous queries and data retrieval', bold_prefix='APIs: ')
add_bullet('Agent-to-agent communication for autonomous task delegation between specialized agents', bold_prefix='A2A: ')
add_bullet('Model Context Protocol under evaluation for standardized cross-platform agent interoperability', bold_prefix='MCP: ')
add_bullet('Cesar\'s YAML semantic layer provides a declarative data access abstraction that decouples agents from raw schema details', bold_prefix='Semantic Layer: ')

doc.add_paragraph()
add_heading_styled('9.3 ROI Argument: Speed & Repeatability', level=2)

doc.add_paragraph(
    'The reusable component architecture directly addresses the key ROI criteria '
    'identified by Mars leadership:'
)

add_bullet('Previous internal builds took 12 weeks; FinIQ was built in under 2 weeks with minimal resources', bold_prefix='Development speed: ')
add_bullet('Visualization, text-to-SQL, and orchestration agents are built once, applied to every new use case', bold_prefix='Reduced duplication: ')
add_bullet('New domains (health data, supply chain) require only data ingestion and feature mapping\u2014not full rebuilds', bold_prefix='Faster onboarding: ')
add_bullet('FinIQ validates the framework; success here proves the model scales across Mars business units', bold_prefix='FinIQ as proof point: ')

# ============================================================
# 10. ROADMAP: AFTER APRIL 21
# ============================================================
add_heading_styled('10. Roadmap: Beyond April 21', level=1)

doc.add_paragraph(
    'Following the MLT demonstration, the platform roadmap focuses on three '
    'areas: deepening FinIQ capabilities, enabling cross-domain reuse, and '
    'adding advanced intelligence layers.'
)

add_table_with_style(
    ['Phase', 'Timeline', 'Focus Area', 'Key Deliverables'],
    [
        ['Phase 2:\nIntelligence', 'May 2026', 'Advanced analytics\n& integration',
         '\u2022 Embeddings & vectorization layer for semantic search\n'
         '\u2022 Amira Financial Forecasting API integration\n'
         '\u2022 Amira Marketing Analytics API integration\n'
         '\u2022 Three-way comparison: Actual vs Replan vs Forecast'],
        ['Phase 3:\nScale', 'Jun\u2013Jul 2026', 'Enterprise readiness\n& multi-tenancy',
         '\u2022 RBAC and role-based data access\n'
         '\u2022 Multi-tenant deployment (per-division agents)\n'
         '\u2022 A2A protocol finalization (MCP or custom)\n'
         '\u2022 Hierarchical agent system (division \u2192 leadership view)'],
        ['Phase 4:\nExpansion', 'Q3 2026', 'Cross-domain reuse\n& new use cases',
         '\u2022 Digital Health data agent (pet care, tooth scans)\n'
         '\u2022 Supply chain analytics agent\n'
         '\u2022 Commodity market data integration\n'
         '\u2022 Reusable component marketplace for Mars teams'],
    ],
    col_widths=[0.9, 0.9, 1.3, 3.2]
)

doc.add_paragraph()

add_heading_styled('10.1 Embeddings & Semantic Search (Next Priority)', level=2)

doc.add_paragraph(
    'A key enhancement identified during architecture discussions is adding an '
    'embeddings and vectorization layer to improve query precision and enable '
    'semantic retrieval across data assets:'
)

add_bullet('Vector embeddings map natural language queries to the most relevant data tables and columns, improving accuracy beyond keyword matching')
add_bullet('Retrieval-Augmented Generation (RAG) workflows ground LLM responses in actual data, reducing hallucination')
add_bullet('Semantic search enables users to find insights across documents, competitor filings, and structured data using natural language')
add_bullet('Implementation considerations include storage, indexing, update cadence, and query latency optimization')

doc.add_paragraph()

add_heading_styled('10.2 ETL & Precomputed Tables', level=2)

doc.add_paragraph(
    'For production scalability with billion-row tables, the team is evaluating '
    'precomputed/synthetic tables in Databricks to optimize heavy query workloads. '
    'This approach reduces query latency from minutes to seconds by pre-aggregating '
    'common analytical views, while maintaining the ability to drill down into raw '
    'data when needed.'
)

# ============================================================
# 11. COMMUNICATION PLAN
# ============================================================
add_heading_styled('11. Communication & Reporting', level=1)

add_bullet('Daily status updates via FinIQ GenAI WhatsApp group', bold_prefix='Daily: ')
add_bullet('Brief sync calls as needed (Cesar to coordinate)', bold_prefix='As needed: ')
add_bullet('Progress report against milestones shared with Atif', bold_prefix='Weekly: ')
add_bullet('Demo rehearsals (April 15 and April 17)', bold_prefix='Pre-demo: ')
add_bullet('MLT presentation and live demo', bold_prefix='April 21: ')

doc.add_paragraph()

doc.add_paragraph(
    'This plan will be updated as work progresses. Any changes to scope, '
    'timeline, or dependencies will be communicated immediately to the team.'
)

# ============================================================
# FOOTER
# ============================================================
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\u2014 End of Document \u2014')
run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
run.font.size = Pt(10)

# Save
output_path = r'C:\Users\farza\Desktop\FinIQ Project Plan - April 2026.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
